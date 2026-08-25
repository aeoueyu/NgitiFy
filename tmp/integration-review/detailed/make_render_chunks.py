from copy import deepcopy
from pathlib import Path

from docx import Document


SOURCE = Path(r"C:\Users\Administrator\Desktop\NGITIFY DENTIME\output\Ngitify_Integration_Testing_Valid_Inputs_Sample_Wording.docx")
OUTDIR = Path(r"C:\Users\Administrator\Desktop\NGITIFY DENTIME\tmp\integration-review\detailed\valid-wording-render-chunks")
NAMES = ["mobile-patient", "admin", "owner", "branch-manager", "dentist", "secretary", "web-patient"]
CHUNK_SIZE = 25


def main():
    OUTDIR.mkdir(parents=True, exist_ok=True)
    master = Document(SOURCE)
    for name, source_table in zip(NAMES, master.tables):
        total = len(source_table.rows) - 6
        for start in range(0, total, CHUNK_SIZE):
            stop = min(start + CHUNK_SIZE, total)
            document = Document()
            body = document._element.body
            section = deepcopy(master.sections[-1]._sectPr)
            for child in list(body):
                body.remove(child)
            table_xml = deepcopy(source_table._tbl)
            body.append(table_xml)
            body.append(section)
            table = document.tables[0]
            for offset, row in reversed(list(enumerate(list(table.rows[6:])))):
                if offset < start or offset >= stop:
                    table._tbl.remove(row._tr)
            output = OUTDIR / f"{name}-{start + 1:03d}-{stop:03d}.docx"
            document.save(output)
            print(output.name)


if __name__ == "__main__":
    main()
