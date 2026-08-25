import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


SOURCE = Path(r"C:\Users\Administrator\Desktop\NGITIFY DENTIME\output\Ngitify_Integration_Testing_All_Roles_Detailed.docx")
OUTPUT = Path(r"C:\Users\Administrator\Desktop\NGITIFY DENTIME\output\Ngitify_Integration_Testing_Valid_Inputs_Sample_Wording.docx")


INVALID_ACTION_PATTERNS = [
    r"\bblank\b", r"\binvalid\b", r"\bincorrect\b", r"\bexpired\b",
    r"\bincomplete\b", r"\bnoncompliant\b", r"\bmismatched\b",
    r"\bmalformed\b", r"\bweak\b", r"\bunsupported\b", r"\boversized\b",
    r"\bmissing\b", r"\bconflicting\b", r"\bunavailable\b", r"\bfails?\b",
    r"\bfailure\b", r"\berror\b", r"\bsimulated\b", r"\bineligible\b",
    r"\binsufficient\b", r"\bduplicate (?:identity|information|item)\b",
    r"\bwithout required\b", r"\bwithout choosing\b", r"\bwithout saving\b",
    r"\boutside (?:the )?.*scope\b", r"\bout-of-scope\b", r"\bunassigned\b",
    r"\banother role\b", r"\banother branch\b", r"\banother patient\b",
    r"\bcopied url\b", r"\baltered request\b", r"\balter a record\b",
    r"^Paste ", r"^Attempt ", r"^Remain inactive", r"^Let the session",
    r"^Trigger a loading", r"^Try to ", r"\bafter a load failure\b",
]


def visible_cells(row):
    result, seen = [], set()
    for cell in row.cells:
        key = id(cell._tc)
        if key not in seen:
            seen.add(key)
            result.append(cell)
    return result


def replace_cell_text(cell, value):
    paragraph = cell.paragraphs[0]
    for extra in cell.paragraphs[1:]:
        cell._tc.remove(extra._p)
    if paragraph.runs:
        paragraph.runs[0].text = value
        for extra in paragraph.runs[1:]:
            paragraph._p.remove(extra._r)
    else:
        paragraph.add_run(value)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def is_valid_case(action):
    text = action.lower()
    return not any(re.search(pattern, text, re.I) for pattern in INVALID_ACTION_PATTERNS)


def button_for(module, action):
    text = f"{module} {action}".lower()
    choices = [
        ("login", "Login"), ("send code", "Send Code"), ("otp", "Verify"),
        ("reset password", "Reset Password"), ("change password", "Save"),
        ("change email", "Submit"), ("booking", "Submit"),
        ("appointment", "Save"), ("schedule", "Save"), ("staff", "Save"),
        ("patient", "Save"), ("branch", "Save"), ("inventory", "Save"),
        ("material usage", "Save"), ("treatment", "Save"),
        ("medical and dental history", "Save"), ("odontogram", "Save"),
        ("radiograph", "Upload"), ("finding", "Save"),
        ("configuration", "Save"), ("permission", "Save"),
        ("notification", "Save"), ("profile", "Save"),
        ("oral care", "Save"), ("onboarding", "Continue"),
        ("consent", "Continue"),
    ]
    for needle, button in choices:
        if needle in text:
            return button
    return "Save"


def data_description(module, action):
    text = f"{module} {action}".lower()
    choices = [
        ("login", "the correct email and password"),
        ("otp", "the correct OTP"),
        ("reset password", "the correct new password and confirm new password"),
        ("change password", "the correct current password, new password, and confirm new password"),
        ("change email", "the correct password and new email address"),
        ("search", "a valid search keyword"),
        ("schedule", "all schedule fields completely and in the correct format"),
        ("appointment", "all appointment fields completely and in the correct format"),
        ("staff", "all staff fields completely and in the correct format"),
        ("patient", "all patient fields completely and in the correct format"),
        ("branch", "all branch fields completely and in the correct format"),
        ("inventory", "all inventory fields completely and in the correct format"),
        ("material usage", "all material usage fields completely and in the correct format"),
        ("treatment", "all treatment fields completely and in the correct format"),
        ("medical and dental history", "all medical and dental history fields completely and in the correct format"),
        ("odontogram", "all tooth information completely and in the correct format"),
        ("radiograph", "all radiograph information completely and in the correct format"),
        ("finding", "all dentist finding fields completely and in the correct format"),
        ("configuration", "all configuration fields completely and in the correct format"),
        ("profile", "all profile fields completely and in the correct format"),
        ("oral care", "all oral care fields completely and in the correct format"),
        ("onboarding", "all onboarding fields completely and in the correct format"),
    ]
    for needle, description in choices:
        if needle in text:
            return description
    return "all fields completely and in the correct format"


def sample_action(module, action):
    action = action.strip().rstrip(".")
    low = action.lower()
    if low == "open the ngitify web url":
        return "Input the correct Ngitify URL"
    if low.startswith("launch the ngitify mobile application"):
        return "Click the Ngitify mobile application icon"

    action = re.sub(r"^Tap\b", "Click", action, flags=re.I)
    action = re.sub(r"^Enter\b", "Input", action, flags=re.I)
    if action.startswith("Input"):
        action = re.sub(r"\band (?:tap|press)\b", "and click", action, flags=re.I)
        return action
    if action.startswith("Click"):
        action = re.sub(r"\bthen (?:tap|press)\b", "then click", action, flags=re.I)
        action = re.sub(r"\band (?:tap|press)\b", "and click", action, flags=re.I)
        return action

    if re.match(r"^(Search|Search,)", action, re.I):
        return "Input a valid search keyword, select the available filters, then click Search button"

    input_verbs = r"^(Submit|Save|Change|Edit|Upload|Add|Complete|Provide|Verify)\b"
    if re.match(input_verbs, action, re.I):
        return f"Input {data_description(module, action)}, then click {button_for(module, action)} button"

    click_replacements = {
        "Open": "Click",
        "Choose": "Click",
        "Select": "Click",
        "Use": "Click",
        "Switch": "Click",
        "Move": "Click",
        "Expand": "Click",
        "Review": "Click",
        "Run": "Click",
        "Refresh": "Click",
        "Apply": "Click",
        "Request": "Click",
        "Approve": "Click",
        "Confirm": "Click",
    }
    for verb, replacement in click_replacements.items():
        if re.match(rf"^{verb}\b", action, re.I):
            return re.sub(rf"^{verb}\b", replacement, action, flags=re.I)

    return f"Click the {module} button"


def finalize_action(module, original, action):
    action = re.sub(r"\benter\b", "input", action, flags=re.I)
    action = re.sub(r"\btap\b", "click", action, flags=re.I)
    action = re.sub(r"\bpress\b", "click", action, flags=re.I)
    low = original.lower()

    if "search" in low:
        return "Input a valid search keyword, select the available filters, then click Search button"

    if "forgot password" in low and "send code" in low:
        return "Input the correct registered email, then click Send Code button"
    if "valid active-account credentials" in low or "valid active patient credentials" in low:
        return "Input the correct email and password, then click Login button"
    if "valid otp" in low and "matching" in low:
        return "Input the correct OTP, new password, and confirm new password, then click Reset Password button"
    if "preferred name" in low:
        return "Input the preferred name, then click Continue button"
    if "cancel appointment" in low and ("reason" in low or "required information" in low):
        return "Input the cancellation reason, then click Confirm Cancellation button"
    if "reschedule" in low and "replacement slot" in low:
        return "Input the new appointment schedule, then click Confirm Reschedule button"
    if "transfer" in low and "reason" in low:
        return "Input the transfer details completely, then click Confirm Transfer button"
    if "add schedule entry" in low and "complete valid" in low:
        return "Click Add Schedule Entry button"
    if "privacy summary" in low:
        return "Click Privacy Summary button"
    if "symptom details" in low and ("severity" in low or "notes" in low):
        return "Input the symptom details completely, then click Done button"
    if "notification toggles" in low or "notification preferences" in low:
        return "Click the notification preference controls, then click Save button"
    if low.startswith("open ") and "review" in low:
        return f"Click the {module} on the sidebar"
    if low.startswith("review "):
        return f"Click the {module} section"
    if low.startswith("switch "):
        return f"Click each visible {module} tab"
    if low.startswith("expand "):
        return f"Click on selected {module} row"
    if low.startswith("move "):
        return "Click the previous or next calendar button"
    if low.startswith("open home"):
        return "Click Home button"

    action = re.sub(r",?\s+or click Skip\b", "", action, flags=re.I)
    action = re.sub(r"\s+", " ", action).strip().rstrip(".")
    return action


def lowercase_first(text):
    return text[:1].lower() + text[1:] if text else text


def sample_expected(module, action, expected):
    action_low = action.lower()
    text = expected.strip().rstrip(".")

    if "cancel" in action_low or "back or close" in action_low:
        if "confirmation" in text.lower() or "modal" in module.lower():
            return f"Removes the confirmation prompt and leaves the {module.lower()} information unchanged"
        return f"Removes the displayed window and displays the previous {module.lower()} page"
    if re.search(r"click (?:the )?ok\b", action_low):
        return "Removes the success message"
    if "close" in action_low and "confirm" not in action_low:
        return f"Removes the displayed window and displays the previous {module.lower()} page"

    text = re.sub(r"^The system\b", "System", text, flags=re.I)
    if re.match(r"^(System|Displays|Removes)\b", text):
        return text

    # Convert the passive constructions that occur repeatedly in the source cases.
    passive = [
        (r"^(.+?) (?:is|are) created\b(.*)$", "creates"),
        (r"^(.+?) (?:is|are) saved\b(.*)$", "saves"),
        (r"^(.+?) (?:is|are) updated\b(.*)$", "updates"),
        (r"^(.+?) (?:is|are) removed\b(.*)$", "removes"),
        (r"^(.+?) (?:is|are) recorded\b(.*)$", "records"),
        (r"^(.+?) (?:is|are) issued\b(.*)$", "issues"),
        (r"^(.+?) (?:is|are) downloaded\b(.*)$", "downloads"),
    ]
    for pattern, verb in passive:
        match = re.match(pattern, text, re.I)
        if match:
            subject = lowercase_first(match.group(1))
            return f"System {verb} {subject}{match.group(2)}"

    display_match = re.match(r"^(.+?) displays\b(.*)$", text, re.I)
    if display_match:
        subject = lowercase_first(display_match.group(1))
        rest = display_match.group(2)
        if rest.strip():
            return f"Displays {subject} and{rest}"
        return f"Displays {subject}"

    if re.match(r"^(Only|All|Each|Current|Authorized|Chosen|Available|Matching|Charts|Metrics|Trend|Additional)\b", text):
        return f"Displays {lowercase_first(text)}"

    replacements = [
        (r"^A (.+?) appears\b(.*)$", r"Displays a \1\2"),
        (r"^An (.+?) appears\b(.*)$", r"Displays an \1\2"),
        (r"^A (.+?) opens\b(.*)$", r"Displays a \1\2"),
        (r"^The (.+?) opens\b(.*)$", r"Displays the \1\2"),
        (r"^The (.+?) becomes\b(.*)$", r"System updates the \1 and makes it\2"),
        (r"^The (.+?) changes\b(.*)$", r"System changes the \1\2"),
        (r"^The (.+?) updates\b(.*)$", r"System updates the \1\2"),
        (r"^The (.+?) switches\b(.*)$", r"System switches the \1\2"),
        (r"^The (.+?) closes\b(.*)$", r"Removes the \1\2"),
        (r"^The (.+?) persists\b(.*)$", r"System saves the \1\2"),
        (r"^Settings persist\b(.*)$", r"System saves the settings\1"),
        (r"^Consent is recorded\b(.*)$", r"System records the consent\1"),
        (r"^Acknowledgement is recorded\b(.*)$", r"System records the acknowledgement\1"),
        (r"^One (.+?) becomes\b(.*)$", r"System updates one \1 and makes it\2"),
    ]
    for pattern, replacement in replacements:
        if re.match(pattern, text, re.I):
            return re.sub(pattern, replacement, text, flags=re.I)

    # Sample-style fallback for uncommon result sentences.
    if any(word in text.lower() for word in ["page", "screen", "modal", "list", "details", "information", "data", "records", "charts", "metrics"]):
        return f"Displays the updated {module.lower()} information"
    return f"System completes the {module.lower()} action and displays a success message"


def final_expected(module, action, original_expected):
    low = action.lower()
    module_text = module.lower()
    original_low = original_expected.lower()

    if "correct ngitify url" in low:
        return "Displays the Login page"
    if "mobile application icon" in low:
        return "Displays the mobile Login page"
    if "password visibility" in low:
        return "System changes the password between masked and visible format"
    if "login button" in low and "email and password" in low:
        return "System verifies the email and password and displays the authorized user dashboard"
    if "send code button" in low:
        return "System verifies the email and displays the OTP page"
    if "resend" in low:
        return "System sends a new OTP and displays a confirmation message"
    if "reset password button" in low:
        return "System updates the password in the database and displays a success message"
    if "back to login" in low or "go to login" in low:
        return "Displays the Login page"

    if "cancel" in low or "back or close" in low:
        return f"Removes the confirmation prompt and leaves the {module_text} information unchanged"
    if re.search(r"click (?:the )?ok\b", low):
        return "Removes the success message"
    if "close" in low and "confirm" not in low:
        return f"Removes the displayed window and displays the previous {module_text} page"

    if "confirm" in low or "yes, logout" in low or "update status" in low:
        if "logout" in low:
            return "System logs the user out and displays the Login page"
        if "delete" in low:
            return f"System removes the confirmation prompt, removes the selected {module_text} information, and displays a success message"
        if "cancellation" in low or "cancel appointment" in low:
            return "System cancels the appointment and displays a success message"
        if "reschedule" in low:
            return "System updates the appointment schedule and displays a success message"
        if "transfer" in low:
            return "System updates the patient branch in the database and displays a success message"
        if "status" in low or any(word in original_low for word in ["activate", "deactivate", "archive", "restore"]):
            return f"System updates the {module_text} status in the database and displays a success message"
        return f"System removes the confirmation prompt, updates the {module_text} information in the database, and displays a success message"

    if low.startswith("input"):
        if "search" in low:
            return f"System searches the database and displays the matching {module_text} records"
        if "otp" in low:
            return "System verifies the OTP and displays the Reset Password page"
        if "onboarding" in module_text:
            return "System saves the onboarding information and displays the next onboarding page"
        if "confirmation" in original_low or "prompt" in original_low:
            return "Displays a user confirmation prompt"
        if "upload" in low:
            return f"System saves the {module_text} information in the database and displays the uploaded file"
        return f"System saves the {module_text} information in the database and displays a success message"

    if any(word in low for word in ["export csv", "export pdf", "download"]):
        return f"System downloads the selected {module_text} file"
    if "search" in low or "filter" in low or "date range" in low:
        return f"Displays the matching {module_text} records"
    if "previous" in low or "next" in low or "pagination" in module_text:
        return f"Displays the selected {module_text} page"
    if "mark all as read" in low:
        return "System updates all notification statuses to read and removes the unread notification badge"
    if "mark as unread" in low or "mark as read" in low:
        return "System updates the selected notification status and displays the updated unread count"
    if "stay logged in" in low:
        return "System renews the session and removes the session warning prompt"
    if "toggle" in low or "preference controls" in low:
        return f"System updates the {module_text} setting and displays a success message"
    if "delete" in low:
        return "Displays a user confirmation prompt"
    if low.startswith("click add "):
        return f"Displays the add {module_text} form"
    if low.startswith("click edit "):
        return f"Displays the edit {module_text} form"
    if any(word in low for word in ["save button", "submit button"]):
        return "Displays a user confirmation prompt"

    if any(word in low for word in ["sidebar", "menu item", "tab", "dashboard", "profile shortcut"]):
        if "role navigation" in module_text:
            return "Displays the selected authorized page and indicates the active menu item"
        return f"Displays the {module_text} page"
    if any(word in low for word in ["details", "view", "open", "bell", "row", "card", "summary", "gallery", "tooth", "notification"]):
        return f"Displays the selected {module_text} information"
    if "clear" in low:
        return f"System removes the selected filters and displays all authorized {module_text} records"
    return f"Displays the updated {module_text} information"


def main():
    document = Document(SOURCE)
    total_before = total_after = 0
    removed_by_role = []

    for table in document.tables:
        rows = list(table.rows[6:])
        total_before += len(rows)
        removed = 0
        for row in rows:
            cells = visible_cells(row)
            module, action, expected = (cells[0].text.strip(), cells[1].text.strip(), cells[2].text.strip())
            if not is_valid_case(action):
                table._tbl.remove(row._tr)
                removed += 1
                continue
            new_action = finalize_action(module, action, sample_action(module, action))
            new_expected = final_expected(module, new_action, expected)
            replace_cell_text(cells[1], new_action)
            replace_cell_text(cells[2], new_expected)
        role = visible_cells(table.rows[1])[3].text.strip()
        count = len(table.rows) - 6
        total_after += count
        removed_by_role.append((role, count, removed))

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    print(f"Before: {total_before}; after: {total_after}; removed: {total_before - total_after}")
    for role, count, removed in removed_by_role:
        print(f"{role}: {count} valid cases ({removed} removed)")


if __name__ == "__main__":
    main()
