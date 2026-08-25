from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE = Path(r"C:\Users\Administrator\Documents\integration testing.docx")
OUTPUT = Path(r"C:\Users\Administrator\Desktop\NGITIFY DENTIME\output\Ngitify_Integration_Testing_All_Roles_Detailed.docx")


def c(module, action, expected):
    return (module, action, expected)


WEB_COMMON = [
    c("Login Page", "Open the Ngitify web URL.", "The Login page displays the Ngitify branding, email and password fields, password visibility control, Forgot Password link, and Login button."),
    c("Login Page", "Click the password visibility control.", "The password switches between masked and visible text without changing its value."),
    c("Login Page", "Click Login with both fields blank or with an invalid email format.", "Required-field or email-format validation displays and no login request is completed."),
    c("Login Page", "Enter incorrect credentials and click Login.", "An authentication error displays, protected pages remain inaccessible, and the entered password is not exposed."),
    c("Login Page", "Enter valid active-account credentials and click Login.", "The system creates an authenticated session and opens the dashboard authorized for that account's role."),
    c("Forgot Password", "Click Forgot Password, enter a registered email, and click Send Code.", "The reset-code screen displays and the system sends or queues one OTP for the registered account without revealing account details."),
    c("Forgot Password", "Enter an invalid, expired, or incomplete OTP and submit it.", "The OTP error displays and the password reset fields remain unavailable."),
    c("Forgot Password", "Click Resend Code after the resend control becomes available.", "A new OTP is issued, the resend timer restarts, and the previous OTP can no longer complete the reset."),
    c("Forgot Password", "Verify a valid OTP, enter mismatched or noncompliant passwords, and click Reset Password.", "Password validation displays and the account password is not changed."),
    c("Forgot Password", "Verify a valid OTP, enter matching valid passwords, and click Reset Password.", "A success message displays, Go to Login is available, and the new password works while the old password does not."),
    c("Forgot Password", "Click Back to Login or Go to Login.", "The Login page displays and temporary reset inputs are cleared."),
    c("Role Navigation", "Open the sidebar and review its menu entries.", "Only pages authorized for the signed-in role display; restricted-role menu entries are absent."),
    c("Role Navigation", "Click each visible sidebar menu item and the sidebar collapse or expand control.", "Each item opens its named page, the active item is identified, and collapsing or expanding does not lose the current page."),
    c("Protected Routes", "Paste a protected URL belonging to another role into the address bar.", "The system blocks the page and redirects to an authorized page or Access Denied view without exposing restricted data."),
    c("Session", "Remain inactive until the session warning appears, then click Stay Logged In.", "The warning closes, the authenticated session is renewed, and the current unsaved page remains displayed."),
    c("Session", "Let the session warning expire without choosing Stay Logged In.", "The session ends, protected data is cleared from view, and the Login page displays."),
    c("Logout Modal", "Click Logout from the account or sidebar menu.", "A Confirm Logout modal displays with Yes, Logout and Cancel controls; the session remains active until confirmed."),
    c("Logout Modal", "Click Cancel or close the Confirm Logout modal.", "The modal closes and the user remains signed in on the current page."),
    c("Logout Modal", "Click Yes, Logout in the Confirm Logout modal.", "The session ends, the Login page displays, and browser Back cannot reopen protected content."),
    c("Notifications", "Click the notification bell.", "The Notifications page or panel displays the current unread count and notifications permitted for the signed-in user."),
    c("Notifications", "Use notification search and type or read-status filters, then clear them.", "The list shows only matching notifications; clearing filters restores the complete authorized list."),
    c("Notifications", "Open an unread notification and close its details.", "The notification details display, the item becomes read, the unread count decreases, and Close returns to the list."),
    c("Notifications", "Click Mark All as Read.", "All currently accessible unread notifications become read and the unread badge updates to zero."),
    c("Notifications", "Click Mark as Unread or Mark as Read on an item.", "The selected item's read state and the unread badge update consistently without changing other items."),
    c("Notifications", "Click Previous or Next in a multi-page notification list.", "The requested page loads once, page information updates, and navigation is disabled at the first or last page."),
    c("Activity Logs", "Open Activity Logs, search, choose a date range, and clear the filters.", "Only matching activities display; clearing filters restores the authorized activity history."),
    c("Activity Logs", "Click View Details on an activity and then Close.", "The detail modal identifies the action, actor, target, time, and available metadata; Close returns to the unchanged list."),
    c("Activity Logs", "Click Export CSV.", "A CSV containing the currently authorized activity records downloads with readable column headings."),
    c("Profile", "Open My Profile and review the displayed account information.", "The signed-in user's current profile, role, and branch information display without another user's data."),
    c("Profile", "Click Edit Profile, change valid editable fields, and click Cancel.", "The edit view closes or resets and no profile change is saved."),
    c("Profile", "Click Edit Profile, enter invalid data, and click Save or Confirm.", "Field-specific validation displays and the invalid profile is not saved."),
    c("Profile", "Save valid profile changes and click Confirm or OK in the confirmation modal.", "A success state displays and the updated values persist in the profile and relevant account views."),
    c("Account Settings", "Enter an incorrect current password or mismatched new passwords and submit Change Password.", "An error displays, the password remains unchanged, and the user stays signed in."),
    c("Account Settings", "Submit a valid current password and matching compliant new password.", "The password changes successfully and the system applies its configured re-login behavior."),
    c("Account Settings", "Change notification preferences and click Save Notifications.", "A success message displays and the selected preferences persist after reopening Settings."),
    c("NgitiBot", "Open NgitiBot from the visible menu or button, enter a message, and click Send.", "The conversation opens, the patient-safe prompt is submitted once, and a response or clear recoverable error displays."),
    c("NgitiBot", "Use New Chat, conversation history, rename, archive, or delete controls.", "The selected conversation action is reflected in the visible history and does not alter unrelated chats."),
    c("NgitiBot Modal", "Open a rename, archive, or delete confirmation and click Cancel.", "The modal closes and the conversation name and state remain unchanged."),
    c("NgitiBot Modal", "Open a rename, archive, or delete confirmation and click Confirm or Delete.", "The requested action completes once, a status message displays when applicable, and the chat list refreshes."),
]


SCHEDULE = [
    c("Schedule Management", "Open Schedule Management and review the calendar or appointment table.", "Authorized appointments display with patient, dentist, branch, schedule, and status information appropriate to the signed-in role."),
    c("Schedule Management", "Search appointments and apply date, dentist, branch, or status filters, then click Clear.", "Only matching appointments display and Clear restores the default authorized schedule."),
    c("Schedule Management", "Click Export CSV and Export PDF.", "Each file downloads with the currently authorized schedule data and readable headings."),
    c("Add Schedule Entry", "Click Add Schedule Entry, choose an existing patient, and complete valid appointment fields.", "The form displays available patient, branch, dentist, service, date, time, and notes controls and accepts valid values."),
    c("Add Schedule Entry", "Choose Add New Patient and submit missing, invalid, or duplicate identity information.", "Required or duplicate warnings display and neither a patient nor appointment is created."),
    c("Add Schedule Entry Modal", "Enter appointment information and click Cancel or close the modal.", "The modal closes, no appointment is created, and the previous schedule remains displayed."),
    c("Add Schedule Entry Modal", "Submit a valid available schedule entry.", "One appointment is created, a success message displays, and the entry appears in the staff schedule and patient's appointments."),
    c("Appointment Details", "Click View on a schedule entry.", "The details modal displays the selected appointment and only actions permitted for its current status and the user's role."),
    c("Appointment Details Modal", "Click Close or the modal close icon.", "The details modal closes without changing the appointment."),
    c("Edit Appointment", "Click Edit or Full Update, change valid fields, and save.", "The existing appointment updates once, availability is revalidated, and the changed details appear across staff and patient views."),
    c("Edit Appointment", "Enter conflicting or invalid appointment information and save.", "A clear validation or availability error displays and the existing appointment remains unchanged."),
    c("Appointment Status Modal", "Select Confirm, In Clinic, Complete, or another allowed next status and click Cancel.", "The confirmation closes and the appointment retains its previous status."),
    c("Appointment Status Modal", "Select an allowed next status and click Confirm or Update Status.", "The status changes once, the schedule refreshes, and the patient-facing appointment and related notification reflect the change."),
    c("Cancel Appointment Modal", "Click Cancel Appointment, enter or select a reason, then close or click Cancel.", "The modal closes and the appointment remains active with its previous status."),
    c("Cancel Appointment Modal", "Click Cancel Appointment, provide required information, and confirm cancellation.", "The appointment becomes cancelled, the slot is released when applicable, and the patient sees the cancellation and notification."),
    c("Guest Appointment", "Open an eligible guest request and click Confirm Guest Request.", "The request becomes a confirmed appointment without creating duplicate schedule entries."),
    c("Guest Appointment", "Click Register or Link Patient and choose an existing match or create a valid patient.", "The appointment links to exactly one patient record and retains its original request details."),
    c("Guest Appointment", "Click Resend Link for an eligible guest request.", "The activation or registration link is reissued once and a success or recoverable delivery message displays."),
    c("Schedule Pagination", "Click Previous or Next in a multi-page schedule.", "The requested page displays with no duplicate rows, and the controls disable at the boundaries."),
]


STAFF_MANAGEMENT = [
    c("Staff Management", "Open the role's staff list and review visible staff cards or rows.", "Only staff within the signed-in user's role and branch scope display with accurate role, branch, activation, and lifecycle status."),
    c("Staff Management", "Search staff and apply Active, Needs Activation, Inactive, Archived, or All filters.", "The list displays only matching in-scope accounts and an appropriate empty state when no record matches."),
    c("Add Staff", "Click Add, submit missing, malformed, or duplicate staff information.", "Required, format, and duplicate validation display and no staff account is created."),
    c("Add Staff Modal", "Enter staff information and click Cancel or close.", "The modal closes and no new staff account appears."),
    c("Add Staff Modal", "Submit valid staff identity, contact, role, and branch information.", "One staff account is created with the selected assignments and a success or activation-delivery result displays."),
    c("View Staff", "Click View on a staff record and then Close or Back.", "The selected staff details display without sensitive credentials; Close or Back returns to the same list state."),
    c("Edit Staff", "Change valid editable staff information and click Save or Done.", "The existing staff account updates once and the new information persists in the list and detail view."),
    c("Edit Staff", "Enter invalid data or click Cancel.", "Invalid values are not saved; Cancel closes the form and preserves the original staff record."),
    c("Staff Lifecycle Modal", "Choose Activate, Deactivate, Archive, or Restore and review Impact Preview, then click Cancel.", "The modal closes and the staff status and access remain unchanged."),
    c("Staff Lifecycle Modal", "Provide the required reason and click Confirm.", "The allowed lifecycle action completes once, list status updates, account access follows the new state, and the action is logged."),
    c("Staff Management", "Attempt an action on a staff account outside the signed-in user's scope.", "The control is unavailable or the server rejects the request without exposing or changing the out-of-scope account."),
]


PATIENT_MANAGEMENT = [
    c("Patients Management", "Open Patients Management and review the patient list.", "Only patients authorized for the role and branch display with accurate identity, branch, contact, and lifecycle information."),
    c("Patients Management", "Search patients and apply status or branch filters, then clear them.", "Only matching authorized patients display; clearing filters restores the allowed list."),
    c("Patients Management", "Click Export CSV and Export PDF.", "Each download contains the currently authorized patient list and excludes out-of-scope records."),
    c("Add Patient", "Click Add Patient and submit blank, malformed, or duplicate identity and contact data.", "Required, format, and possible-duplicate messages display and no duplicate patient is created."),
    c("Add Patient", "Complete identity, contact, medical or dental, and consent steps using Continue and Back.", "Each step retains valid entered data, prevents invalid progression, and displays the appropriate consent information."),
    c("Add Patient Modal", "Click Cancel or close before completion.", "The form closes and no patient record is created."),
    c("Add Patient", "Submit a complete valid patient record and confirm required consent.", "One patient record is created, a success state displays, and the patient appears in authorized lists."),
    c("Duplicate Patient", "Choose Open Existing Record after a possible duplicate is found.", "The existing patient's authorized detail view opens and no additional patient is created."),
    c("Patient Details", "Click View on a patient and then Close or Back.", "The selected patient's authorized details display; Close or Back returns to the same filtered list."),
    c("Edit Patient", "Save valid non-clinical patient changes.", "The existing patient updates once and the new information persists in staff and patient profile views as permitted."),
    c("Edit Patient", "Enter invalid values or click Cancel.", "Validation prevents invalid saving; Cancel preserves the original patient record."),
    c("Patient Lifecycle Modal", "Choose Activate, Deactivate, Archive, or Restore and click Cancel.", "The modal closes and the patient's lifecycle status and access remain unchanged."),
    c("Patient Lifecycle Modal", "Provide the required reason and click Confirm.", "The permitted lifecycle change completes, list placement and account access update, and the action is logged."),
    c("Transfer Patient Branch", "Open Transfer Branch and review the target branch, upcoming appointments, and Impact Preview.", "The modal accurately identifies records affected by the proposed transfer before any change occurs."),
    c("Transfer Patient Branch Modal", "Select a target branch and click Cancel.", "The modal closes and the patient's branch and appointments remain unchanged."),
    c("Transfer Patient Branch Modal", "Select an allowed target branch, provide a reason, and click Confirm Transfer.", "The patient moves once to the target branch, related appointments follow the confirmed handling, and one audit entry is created."),
]


MOBILE_PATIENT = [
    c("Mobile Login", "Launch the Ngitify mobile application.", "The Login screen displays branding, email and password fields, password visibility, Forgot Password, and Login controls without clipped content."),
    c("Mobile Login", "Tap the password visibility icon.", "The password switches between masked and visible text without altering its value."),
    c("Mobile Login", "Tap Login with blank, malformed, or incorrect credentials.", "Field or authentication errors display and the patient dashboard does not open."),
    c("Mobile Login", "Enter valid active patient credentials and tap Login.", "A patient session is created and the privacy, onboarding, or Home screen opens according to the account's completion state."),
    c("Mobile Forgot Password", "Tap Forgot Password, enter a registered email, and tap Send Code.", "The OTP screen displays, one reset code is issued, and no sensitive account details are revealed."),
    c("Mobile Forgot Password", "Submit an invalid or expired OTP, then tap Resend when enabled.", "The invalid code is rejected; Resend issues a new code and restarts its timer."),
    c("Mobile Forgot Password", "Verify a valid OTP but submit mismatched or weak new passwords.", "Password validation displays and the current password remains valid."),
    c("Mobile Forgot Password", "Verify a valid OTP and submit matching compliant passwords.", "A success state displays, Go to Login or Close returns to Login, and only the new password authenticates."),
    c("Privacy Consent", "Review Privacy Information and tap Close.", "The privacy details close and the consent screen remains displayed without recording consent."),
    c("Privacy Consent", "Tap I Agree and Continue.", "Consent is recorded once and the next incomplete onboarding screen or dashboard opens."),
    c("Onboarding", "Tap Get Started on the Welcome screen.", "The preferred-name onboarding screen displays and progress begins for the signed-in patient."),
    c("Onboarding", "Enter a preferred name and tap Continue, or tap Skip.", "A valid name is saved when provided; Skip leaves it optional and opens the next onboarding step."),
    c("Onboarding", "Complete or skip Help Focus, Routine, and Experience screens.", "Each selection or skip state is saved once and the next onboarding screen displays in sequence."),
    c("Onboarding Notifications", "Choose reminder preferences and tap Save Preferences, or tap Skip.", "Chosen preferences persist; Skip retains defaults and both actions continue onboarding."),
    c("Onboarding Ready", "Tap Go to Dashboard or another displayed next action such as Log Today.", "Onboarding is marked complete and the selected Home or Oral Care destination opens."),
    c("Mobile Navigation", "Tap Home, Health, Visits, Records, and Profile in the bottom navigation.", "Each tab opens the matching patient screen and the selected tab remains visually active."),
    c("Mobile Home", "Open Home and review summary cards, today's information, upcoming appointment, and Recommended Visit Window.", "The dashboard displays the signed-in patient's current data and appropriate empty, loading, or retry states."),
    c("Mobile Home", "Tap a visible appointment, Oral Health, Recommended Visit Window, notification, or profile shortcut.", "The matching detail screen or modal opens with the selected patient's data."),
    c("Mobile Logout Modal", "Tap Logout and then Cancel or close the modal.", "The modal closes and the patient remains signed in on the current screen."),
    c("Mobile Logout Modal", "Tap Logout and confirm Logout.", "The patient session ends, Login displays, and protected mobile screens cannot be reopened."),
    c("Mobile Appointment Booking", "Tap Book Appointment and use Back and Next through branch, date, time, service, and notes steps.", "Valid choices persist between steps, unavailable selections are disabled, and missing required choices prevent progression."),
    c("Mobile Appointment Booking", "Open the privacy summary and tap I Understand or close it.", "The summary closes; I Understand records acknowledgement while Close alone does not submit the booking."),
    c("Mobile Appointment Booking", "Submit a complete booking with an available slot.", "One appointment request is created, a success state displays, and it appears under Visits and the staff schedule."),
    c("Mobile Appointment Booking", "Submit after the selected slot becomes unavailable or the request fails.", "A recoverable error and Retry or reselection option display, and no duplicate appointment is created."),
    c("Mobile Visits", "Open Upcoming and Past appointments and tap an appointment.", "The appropriate appointment list and selected appointment details display with actions valid for its status."),
    c("Mobile Reschedule Modal", "Tap Reschedule, choose a new valid slot, then tap Back or Cancel.", "The modal closes and the original appointment schedule remains unchanged."),
    c("Mobile Reschedule Modal", "Tap Confirm Reschedule for an available slot.", "The existing appointment is updated once and the new time appears in patient and staff schedule views."),
    c("Mobile Cancellation Modal", "Tap Cancel Appointment and then Back or close.", "The modal closes and the appointment remains active."),
    c("Mobile Cancellation Modal", "Provide the required reason and tap Confirm Cancellation.", "The appointment becomes cancelled, the released slot is handled correctly, and a notification or confirmation displays."),
    c("Mobile Oral Care - Today", "Select a date, mark oral-care items, enter symptom details or severity and notes, then tap Save.", "The daily log saves once and its selected values reappear when the same date is reopened."),
    c("Mobile Oral Care Modal", "Open symptom details and tap Done or Close.", "Done applies valid symptom details; Close without saving preserves the previously saved details."),
    c("Mobile Oral Care - Calendar", "Move to the previous or next period and select a logged date.", "The calendar changes period and displays the selected date's saved oral-care record."),
    c("Mobile Oral Care - Trends", "Switch between 7-day and 30-day views.", "Trend summaries recalculate for the selected range using only the patient's saved logs."),
    c("Mobile Oral Care - Education", "Search, filter by category, and open an education article.", "Matching patient-safe topics display and the selected article opens with its title and content."),
    c("Mobile Records", "Open each visible Medical Records tab and use Try Again after a simulated load failure.", "The selected authorized record section displays; Try Again reloads it without duplicating data."),
    c("Mobile Records - Treatments", "Expand and collapse treatment history or tap Show All.", "The selected treatment details display in chronological context and Close returns to the list."),
    c("Mobile Records - Odontogram", "Tap a tooth on the odontogram.", "The tooth's latest dentist-recorded status displays read-only for the patient."),
    c("Mobile Records - X-Rays", "Open an X-ray, tap Explain This Record if available, and tap Back.", "The authorized image and approved explanation display; Back returns to the same records position."),
    c("Mobile Notifications", "Open an unread notification and return to the list.", "The correct details display, the item becomes read, and the unread badge decreases consistently."),
    c("Mobile Notifications", "Trigger a loading error and tap Try Again.", "The list retries loading and displays either current notifications or a clear empty or error state."),
    c("Mobile Profile", "Open Profile and tap Edit Profile, Settings, Activity Logs, or Medical Records.", "Each visible shortcut opens its matching patient screen."),
    c("Mobile Edit Profile", "Change valid fields or profile image and tap Back or Cancel before saving.", "No profile update is committed and the previously saved profile remains displayed."),
    c("Mobile Edit Profile", "Enter invalid values and tap Save.", "Field validation displays and invalid profile data is not saved."),
    c("Mobile Edit Profile", "Save valid changes and tap Confirm or Close on the success modal.", "The profile updates once, the success modal closes, and the new values persist across mobile and authorized web views."),
    c("Mobile Settings", "Change password using incorrect, mismatched, then valid password values.", "Invalid submissions show errors; a valid submission changes the password and applies the configured session behavior."),
    c("Mobile Change Email Modal", "Open Change Email and tap Close or Cancel.", "The modal closes and the account email remains unchanged."),
    c("Mobile Change Email Modal", "Verify the password, submit a valid unused email, and confirm the requested change.", "A verification message displays and the email changes only after the required verification process."),
    c("Mobile Settings", "Change notification toggles or reminder time and leave and reopen Settings.", "The saved preferences persist and control only eligible patient notifications."),
    c("Mobile Activity Logs", "Open Activity Logs, tap Load More, then use Back.", "Additional authorized activities append without duplicates and Back returns to Profile."),
    c("Mobile NgitiBot", "Open NgitiBot, send a message or quick prompt, and close it.", "The conversation uses the signed-in patient's allowed context, displays a response or Retry state, and Close returns to the previous screen."),
    c("Mobile NgitiBot Modal", "Open rename, archive, or delete and tap Cancel.", "The modal closes and the selected conversation remains unchanged."),
    c("Mobile NgitiBot Modal", "Confirm rename, archive, or delete.", "The selected conversation action completes once and the active or archived chat list refreshes."),
    c("Mobile Pre-Operation", "Open the pre-operation instructions and try to continue without required acknowledgement.", "The required agreement warning displays and the workflow does not advance."),
    c("Mobile Pre-Operation", "Select the agreement checkbox and continue or sign as requested.", "Acknowledgement is recorded for the correct patient and the next permitted step displays."),
]


ADMIN_EXTRA = [
    c("Administrator Dashboard", "Open Dashboard and review summary cards, charts, calendar, alerts, and recent data.", "Current system-wide authorized metrics display with loading, empty, and error states handled without stale cross-role data."),
    c("Administrator Dashboard", "Click Add Patient, Manage Appointments, a summary card, or chart link.", "Each shortcut opens the matching management page and carries any intended filter context."),
    c("Administrator Dashboard", "Move the calendar to previous or next periods and close a visible alert.", "The calendar changes period; closing the alert removes it without changing underlying operational data."),
    *STAFF_MANAGEMENT,
    *PATIENT_MANAGEMENT,
    c("Inventory Management", "Open Inventory and search, filter, or review stock and batch history.", "Authorized inventory rows, quantities, expiry or low-stock states, and expanded batch history display accurately."),
    c("Inventory Management", "Click Export CSV and Export PDF.", "The current authorized inventory data downloads in the selected format."),
    c("Add Inventory Item", "Submit missing, invalid, or duplicate item information.", "Validation displays and no inventory item or batch is created."),
    c("Add Inventory Item Modal", "Enter item or batch details and click Cancel.", "The modal closes and inventory remains unchanged."),
    c("Add Inventory Item Modal", "Submit valid item and initial batch information.", "One item and valid batch are created, totals recalculate, and a success state displays."),
    c("Receive Stock", "Add a valid new batch to an existing inventory item.", "The batch appears in history and the item's available quantity and status recalculate."),
    c("Edit Inventory", "Save valid edits or click Cancel.", "Save persists allowed item changes; Cancel closes without altering the item."),
    c("Delete Inventory Batch Modal", "Click delete on an eligible batch and then Cancel.", "The confirmation closes and the batch and stock totals remain unchanged."),
    c("Delete Inventory Batch Modal", "Confirm deletion of an eligible batch.", "The batch is removed once, totals recalculate, and blocked dependent batches display an explanatory error instead."),
    c("Branch Management", "Open branches, search, and click Add Branch.", "The authorized branch list and an add form with required branch information display."),
    c("Add/Edit Branch Modal", "Submit invalid or duplicate branch information.", "Validation displays and no invalid or duplicate branch is saved."),
    c("Add/Edit Branch Modal", "Click Cancel or submit valid branch information and click Done.", "Cancel preserves data; Done creates or updates one branch and refreshes the list."),
    c("Branch Status Modal", "Choose Activate or Deactivate and click Cancel.", "The modal closes and branch status remains unchanged."),
    c("Branch Status Modal", "Confirm Activate or Deactivate.", "The branch status updates once and dependent assignment restrictions are enforced with a clear message when blocked."),
    c("Branch Analytics", "Click View Analytics, change filters, clear them, and export CSV or PDF.", "Charts and totals reflect the selected branch and period; Clear restores defaults and exports match the displayed scope."),
    c("Branch Analytics Modal", "Open branch analytics details and click Close.", "The selected branch metrics display and Close returns to Branch Management."),
    c("System Configuration", "Switch among configuration sections and website-content tabs.", "The selected settings form displays its current saved values without saving tab navigation alone."),
    c("System Configuration", "Edit clinic, appointment, email-template, feature, or website-content values and click Reset Default.", "A default or confirmation state displays and no reset is persisted until the required confirmation or save occurs."),
    c("System Configuration", "Add or remove a procedure or website service card, then click Cancel or leave without saving.", "Unsaved list changes are discarded and the last saved configuration remains active."),
    c("System Configuration", "Upload valid supported media, then use Reset Media.", "The preview updates for valid media; Reset restores the configured default after confirmation and unsupported files show validation."),
    c("System Configuration", "Save valid settings and click OK on feedback.", "A success message displays, OK closes it, settings persist, and affected public or booking pages use the new values."),
    c("System Configuration", "Save invalid configuration values.", "Validation or an error message displays and the last valid configuration remains active."),
    c("Database Backup", "Open Backup and click Refresh.", "Current backup status, schedule, retention, verification, and history reload without starting a backup."),
    c("Database Backup", "Click Create Backup Now and cancel its confirmation.", "The confirmation closes and no manual backup job is started."),
    c("Database Backup", "Confirm Create Backup Now.", "One backup job starts, progress or status displays, and a completed or failed history entry is recorded."),
    c("Database Backup", "Edit automatic schedule or retention using invalid values and save.", "Validation displays and existing backup settings remain unchanged."),
    c("Database Backup", "Save a valid automatic schedule and retention setting.", "A success state displays and the settings persist after Refresh or reopening the page."),
    c("Database Backup", "Click Verify on a completed backup.", "Verification runs for the selected file and records a valid or failed integrity result without changing the backup."),
    c("Database Backup", "Click Download on a completed backup.", "The selected backup downloads with the expected file identity; unavailable files show a clear error."),
    c("Integrity Tools", "Run one integrity check or Run All.", "Each requested check reports status, affected-record count, and expandable details without modifying data."),
    c("Integrity Tools", "Expand and collapse affected records.", "The selected issue details display and collapse without changing any record."),
    c("Safe Auto-Fix Modal", "Open Safe Auto-Fix and click Cancel.", "The modal closes and no affected record changes."),
    c("Safe Auto-Fix Modal", "Confirm Safe Auto-Fix for an eligible issue.", "Only listed safe changes are applied, a result summary displays, rechecking reflects the fix, and the action is audited."),
    c("Archive Review", "Refresh, search, and filter archived records by type.", "Only matching archived records display with restoration or deletion eligibility and explanatory blockers."),
    c("Restore Archive Modal", "Click Restore and then Cancel.", "The modal closes and the record remains archived."),
    c("Restore Archive Modal", "Confirm Restore for an eligible record.", "The record returns once to its active list with a valid restored state and the action is logged."),
    c("Permanent Delete Modal", "Click Delete Review and then Cancel.", "The confirmation closes and the archived record remains recoverable."),
    c("Permanent Delete Modal", "Confirm permanent deletion for an eligible record.", "The record is deleted once and removed from Archive Review; blocked dependencies prevent deletion and display the reason."),
    c("Roles and Permissions", "Switch role tabs and review the permission matrix.", "The selected role's current permissions display and locked administrator protections cannot be removed."),
    c("Roles and Permissions", "Change editable permissions and navigate away or cancel before saving.", "Unsaved permission changes do not affect authorization."),
    c("Roles and Permissions", "Save allowed permission changes.", "A success message displays and new sessions enforce the updated permissions while preserving protected system access."),
    c("System Audit Logs", "Search, filter by range, paginate, and click Export CSV.", "Only matching audit entries display and the exported file matches the authorized result set."),
    c("System Audit Details", "Click View Details and then Close.", "Immutable actor, role, action, target, time, result, and available metadata display; Close returns to the list."),
]


OWNER_EXTRA = [
    c("Owner Dashboard", "Open Dashboard and review organization summary cards, recent activity, alerts, and calendar.", "Authorized organization-wide totals and records display with clear loading, empty, and error states."),
    c("Owner Dashboard", "Click Manage Staff, Manage Schedule, Patient Records, or another visible shortcut.", "The matching page opens with organization-level scope permitted to the owner."),
    c("Owner Dashboard", "Use previous, next, or date controls on the calendar and close an alert.", "The calendar displays the selected period and the alert closes without changing operational records."),
    *STAFF_MANAGEMENT,
    *PATIENT_MANAGEMENT,
    c("Branch Management", "Open Branches, search, and review branch information and status.", "All owner-authorized branches display with accurate status and operational summary information."),
    c("Branch Management", "Click Add Branch and submit invalid or duplicate information.", "Required and duplicate validation display and no branch is created."),
    c("Branch Modal", "Enter branch information and click Cancel.", "The modal closes and no branch information changes."),
    c("Branch Modal", "Submit valid Add or Edit branch information and click Done.", "One branch is created or updated and the saved details persist in branch and assignment views."),
    c("Branch Status Modal", "Open Activate or Deactivate and click Cancel.", "The confirmation closes and the branch retains its current status."),
    c("Branch Status Modal", "Confirm the allowed status change.", "The branch status updates once; blocked dependencies show a reason and leave the branch unchanged."),
    c("Branch Analytics", "Open analytics, apply branch or date filters, clear them, and export CSV or PDF.", "Metrics reflect the selected scope, Clear restores defaults, and exports match the displayed data."),
    c("Inventory Management", "Search and review item stock, batch history, expiry, and low-stock states.", "Current authorized inventory and accurate calculated totals display."),
    c("Inventory Management", "Add or edit an item or receive a batch, then click Cancel.", "The modal closes and inventory quantities and item information remain unchanged."),
    c("Inventory Management", "Save a valid item edit or received batch.", "The item or batch persists once and totals and warning states recalculate."),
    c("Inventory Batch Modal", "Open batch deletion and click Cancel, then repeat and confirm an eligible deletion.", "Cancel preserves the batch; Confirm removes it once or displays a dependency blocker and updates totals."),
    c("Material Usage", "Open Material Usage, search, filter by date, and expand a log.", "Only matching authorized logs display and expansion shows the appointment, patient, item, quantity, and recorder details."),
    c("Material Usage Modal", "Click Log New Entry, add or remove item rows, and click Cancel.", "The modal closes and no material quantity or inventory stock changes."),
    c("Material Usage Modal", "Submit invalid quantities or an ineligible appointment.", "Validation displays and neither a usage log nor stock deduction occurs."),
    c("Material Usage Modal", "Submit a valid usage record for a completed appointment.", "One usage log is created and inventory totals decrease by the recorded batch quantities."),
    c("Roles and Permissions", "Review role tabs and permission matrix.", "Current organization roles and owner-authorized permissions display; protected permissions remain locked."),
    c("Roles and Permissions", "Change editable permissions and cancel or save.", "Cancel discards changes; Save persists allowed changes and future authorization follows the updated matrix."),
]


BRANCH_MANAGER_EXTRA = [
    c("Branch Manager Dashboard", "Open Dashboard and review branch patient, staff, appointment, alert, and calendar information.", "Only the assigned branch's current totals and records display with clear loading, empty, and error states."),
    c("Branch Manager Dashboard", "Click Add Patient, Manage Schedule, a summary card, or another visible shortcut.", "The matching page opens with the manager's branch scope applied."),
    c("Branch Manager Dashboard", "Use calendar previous, next, or date controls and close an alert.", "The selected calendar period displays and the alert closes without changing records."),
    *STAFF_MANAGEMENT,
    *PATIENT_MANAGEMENT,
    c("Branch Analytics", "Open Analytics and review summary cards and charts.", "Metrics use only the manager's assigned branch and reflect the selected period."),
    c("Branch Analytics", "Apply or clear available filters and export CSV or PDF.", "Charts recalculate for the filters, Clear restores defaults, and exports contain only assigned-branch data."),
    c("Inventory Management", "Open Inventory and search or review item and batch information.", "Only inventory available to the assigned branch displays with accurate stock and warning states."),
    c("Inventory Modal", "Add, edit, receive, or delete an inventory item or batch and click Cancel.", "The modal closes and inventory remains unchanged."),
    c("Inventory Modal", "Submit valid allowed item or batch changes.", "The branch inventory updates once, totals recalculate, and an activity record is created."),
    c("Branch Scope", "Attempt to view or modify another branch through a copied URL or altered request.", "The request is rejected or redirected and no other-branch data is exposed or changed."),
]


DENTIST_EXTRA = [
    c("Dentist Dashboard", "Open Dashboard and review today's schedule, assigned patients, alerts, summary data, and calendar.", "Only dentist-authorized appointments and patients display with accurate current status."),
    c("Dentist Dashboard", "Click View EMR, Manage Patients, Manage Appointments, or a visible shortcut.", "The selected patient or management page opens within the dentist's authorized assignment scope."),
    c("Dentist Dashboard", "Use calendar previous, next, or date controls and close an alert.", "The selected schedule period displays and closing the alert does not change clinical or appointment data."),
    c("Dentist Patients", "Search assigned patients and click View EMR.", "Only assigned or otherwise authorized patients display and the selected patient's EMR opens."),
    c("Patient EMR", "Switch among Overview, Medical and Dental History, Treatment History, Odontogram, and Radiograph Images.", "Each tab displays the selected patient's latest matching clinical data without mixing patient records."),
    c("Medical and Dental History", "Click Edit Medical History, change fields, and click Cancel.", "The editor closes and the previously saved clinical history remains unchanged."),
    c("Medical and Dental History", "Save valid history changes.", "The history updates once with recorder and time information and becomes visible in authorized patient records."),
    c("Treatment Log Modal", "Click Add Treatment Log, enter incomplete or invalid information, and save.", "Validation displays and no treatment log is created."),
    c("Treatment Log Modal", "Enter treatment information and click Cancel.", "The modal closes and treatment history remains unchanged."),
    c("Treatment Log Modal", "Submit a valid treatment log.", "One chronological treatment entry is created for the selected patient and appears in authorized patient records."),
    c("Delete Treatment Modal", "Open delete confirmation and click Cancel.", "The confirmation closes and the treatment entry remains."),
    c("Delete Treatment Modal", "Confirm deletion of an eligible treatment entry.", "The selected entry is removed or archived once and the action is audited; blocked entries display a reason."),
    c("Treatment History", "Search or filter by date and click Export PDF.", "Only matching treatment entries display and the generated PDF represents the selected patient's authorized history."),
    c("EMR Print Preview", "Open print preview, then click Close or Print.", "Close returns without printing; Print opens the system print workflow for the selected patient's formatted record."),
    c("Odontogram", "Select a tooth and open its update modal.", "The modal displays the selected tooth number and its latest saved condition and treatment fields."),
    c("Odontogram Modal", "Change tooth information and click Cancel.", "The modal closes and the odontogram remains unchanged."),
    c("Odontogram Modal", "Save valid tooth information.", "The selected tooth updates once, the odontogram redraws, and the change appears in authorized patient records."),
    c("Radiograph Upload", "Choose an unsupported, oversized, or missing image and click Upload or Save.", "A file validation error displays and no radiograph is stored."),
    c("Radiograph Upload", "Choose a valid image, enter metadata, and click Cancel.", "The upload closes and no new radiograph appears."),
    c("Radiograph Upload", "Upload a valid radiograph for the selected patient.", "One image is stored under the correct patient and appears in the gallery with its metadata."),
    c("Radiograph Gallery", "Open an image and use Original, Enhanced, zoom, Fit, Grab, or Compare controls.", "The viewer changes only the presentation or comparison state and preserves the stored original."),
    c("Radiograph Gallery", "Click Back or Close in the viewer.", "The viewer closes and returns to the selected patient's radiograph gallery."),
    c("Delete Radiograph Modal", "Open delete confirmation and click Cancel.", "The modal closes and the radiograph remains available."),
    c("Delete Radiograph Modal", "Confirm deletion of an eligible radiograph.", "The selected image is removed once or a dependency blocker displays, and the action is audited."),
    c("AI Radiograph Review", "Request AI enhancement or analysis and wait for completion.", "A separate review result or recoverable error displays without replacing the original image or publishing unapproved findings."),
    c("Radiograph Annotation", "Add, edit, reset, or remove a visual annotation.", "The active annotation layer updates as selected while the original radiograph remains unchanged."),
    c("Dentist Finding Modal", "Add or edit a finding and click Cancel.", "The modal closes and no finding changes."),
    c("Dentist Finding Modal", "Save a valid dentist finding or archive and restore it.", "The finding state updates once with dentist attribution and remains linked to the correct image and patient."),
    c("Delete Finding Modal", "Click Cancel, then repeat and confirm Delete.", "Cancel preserves the finding; Delete removes the selected eligible finding once and refreshes the review."),
    c("Radiograph Review Approval", "Approve the completed review summary.", "Only dentist-approved findings and explanation become available to the patient; unapproved AI content remains internal."),
    c("Material Usage", "Search, filter, and expand material usage logs.", "Matching dentist-authorized usage entries display with appointment, item, quantity, and date details."),
    c("Material Usage Modal", "Add or remove item rows and click Cancel.", "The modal closes and no usage record or stock deduction occurs."),
    c("Material Usage Modal", "Submit invalid quantities, insufficient stock, or an ineligible appointment.", "Validation displays and inventory remains unchanged."),
    c("Material Usage Modal", "Submit valid materials for a completed appointment.", "One usage log is created and the correct inventory batches and totals are reduced."),
    c("Dentist Scope", "Attempt to open an unassigned patient's EMR through a copied URL.", "Access is denied or redirected and no unassigned patient information is exposed."),
]


SECRETARY_EXTRA = [
    c("Secretary Dashboard", "Open Dashboard and review branch patients, appointments, alerts, summary cards, and calendar.", "Only assigned-branch operational data displays; clinical edit controls are absent."),
    c("Secretary Dashboard", "Click Add Patient, Manage Appointments, a summary card, or another shortcut.", "The corresponding operational page opens with assigned-branch scope."),
    c("Secretary Dashboard", "Use calendar previous, next, or date controls and close an alert.", "The selected calendar period displays and the alert closes without changing records."),
    c("Secretary Check-In Modal", "Click Check In for an eligible appointment and then Cancel.", "The confirmation closes and the appointment retains its previous status."),
    c("Secretary Check-In Modal", "Confirm Check In for an eligible appointment.", "The appointment changes to the permitted in-clinic status once and the staff schedule refreshes."),
    *PATIENT_MANAGEMENT,
    c("Patient View", "Open a patient and click View EMR.", "The patient's authorized record displays in read-only form for the secretary; dentist-only clinical edit controls are absent."),
    c("Patient View", "Click Edit Profile, save allowed demographic changes, and return with Back.", "Allowed non-clinical changes persist and Back returns to the selected patient or list."),
    c("Secretary Scope", "Attempt to open another branch's patient or appointment using a copied URL.", "Access is denied or redirected and no out-of-branch data is exposed or changed."),
]


WEB_PATIENT_EXTRA = [
    c("Patient Dashboard", "Open Dashboard and review date controls, appointment summary, Oral Care, and Recommended Visit Window.", "Only the signed-in patient's current information displays with appropriate loading, empty, and retry states."),
    c("Patient Dashboard", "Click My Appointments, Details, Open My Appointments, or Oral Care.", "The matching patient page or selected appointment detail opens."),
    c("Recommended Visit Window Modal", "Click Why am I seeing this, then Close.", "The modal explains the recommendation's patient-safe source and timing; Close returns to Dashboard without changing data."),
    c("Patient Appointments", "Switch between Upcoming, History, and Book views.", "The selected tab displays the correct signed-in patient's appointment state and valid actions."),
    c("Patient Appointment Details", "Click Details and then Close.", "The selected appointment's branch, dentist, service, schedule, status, and notes permitted to the patient display; Close returns to the list."),
    c("Patient Booking", "Click Book New Appointment and proceed through required branch, service, date, time, and notes controls.", "Available options display, missing required choices block progression, and selected values persist through Review and Back."),
    c("Patient Booking Privacy Modal", "Open the privacy summary and click Close.", "The summary closes without submitting the request or changing entered booking details."),
    c("Patient Booking", "Submit a valid available appointment request.", "One request is created, the success modal displays, and the appointment appears in the patient's list and staff schedule."),
    c("Patient Booking Success Modal", "Click Close or Go to Appointments.", "Close dismisses the modal; Go to Appointments opens the created request in the appointments page."),
    c("Patient Booking", "Submit after a slot conflict or simulated request failure, then click Try Again.", "A clear error displays, no duplicate request is created, and Try Again resubmits or reloads safely."),
    c("Patient Reschedule Modal", "Choose Reschedule and click Back or Close.", "The modal closes and the original appointment schedule remains unchanged."),
    c("Patient Reschedule Modal", "Choose an available replacement slot and confirm.", "The existing appointment updates once and both patient and staff schedule views show the new time."),
    c("Patient Cancellation Modal", "Choose Cancel Appointment and click Back or Close.", "The modal closes and the appointment remains active."),
    c("Patient Cancellation Modal", "Provide the required reason and confirm cancellation.", "The appointment becomes cancelled once and the patient receives confirmation or a related notification."),
    c("Medical Records", "Switch among all visible record tabs and click Try Again after a load failure.", "Each authorized section shows the latest patient record; Try Again reloads without duplicating entries."),
    c("Medical Records - Treatments", "Expand and collapse a treatment entry.", "The selected treatment details display and collapse without altering the clinical record."),
    c("Medical Records - Odontogram", "Select a tooth.", "The tooth's latest dentist-recorded status displays read-only."),
    c("Medical Records - Radiographs", "Open an authorized radiograph and then Close.", "The correct image and approved patient explanation display; Close returns to Records."),
    c("Oral Care - Today", "Select a date, mark care items, enter symptom details, and save the daily log.", "The log saves once and the same values display when that date is reopened."),
    c("Oral Care Symptom Modal", "Change symptom severity or notes and click Close or Done.", "Done applies valid details; closing without applying preserves the previous saved values."),
    c("Oral Care - History", "Open History or Calendar and select a logged date.", "The selected date's saved oral-care information displays accurately."),
    c("Oral Care - Trends", "Switch between 7-day and 30-day views.", "Trend values and risk factors recalculate for the selected range using only the patient's logs."),
    c("Oral Care", "Click Book Preventive Visit or Open NgitiBot.", "The booking page or NgitiBot opens with relevant patient-safe oral-care context."),
    c("Dental Health Education", "Search topics, select a category, click Show All, and open a topic.", "The library filters correctly and the selected patient-safe article opens."),
    c("Dental Health Education", "Click Open NgitiBot from a topic.", "NgitiBot opens with the selected education context and does not expose restricted clinical notes."),
    c("Patient Settings", "Click View Profile, Edit Profile, or Activity Logs.", "Each shortcut opens the signed-in patient's corresponding page."),
    c("Change Email Modal", "Open Change Email and click Cancel.", "The modal closes and the current email remains unchanged."),
    c("Change Email Modal", "Submit a valid unused email, verify credentials if requested, and click OK.", "A verification or success message displays; OK closes it and the email changes only after required verification."),
    c("Patient Privacy Settings", "Change visible privacy or data-sharing controls and save.", "Allowed choices persist and affect only the configured patient-facing sharing behavior."),
    c("Patient Scope", "Alter a record or appointment URL to reference another patient.", "Access is denied and no other patient's appointments, profile, or clinical records display."),
]


CASES = {
    "Patient / Mobile Application": MOBILE_PATIENT,
    "Admin / Web Application": WEB_COMMON + SCHEDULE + ADMIN_EXTRA,
    "Owner / Web Application": WEB_COMMON + SCHEDULE + OWNER_EXTRA,
    "Branch Manager / Web Application": WEB_COMMON + SCHEDULE + BRANCH_MANAGER_EXTRA,
    "Dentist / Web Application": WEB_COMMON + SCHEDULE + DENTIST_EXTRA,
    "Secretary / Web Application": WEB_COMMON + SCHEDULE + SECRETARY_EXTRA,
    "Patient / Web Application": WEB_COMMON + WEB_PATIENT_EXTRA,
}


def unique_cells(row):
    result = []
    seen = set()
    for cell in row.cells:
        key = id(cell._tc)
        if key not in seen:
            seen.add(key)
            result.append(cell)
    return result


def replace_cell_text(cell, value):
    paragraphs = cell.paragraphs
    paragraph = paragraphs[0]
    for extra in paragraphs[1:]:
        cell._tc.remove(extra._p)
    runs = paragraph.runs
    if runs:
        runs[0].text = value
        for extra in runs[1:]:
            paragraph._p.remove(extra._r)
    else:
        paragraph.add_run(value)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def fill_table(table, prototype_tr, cases):
    for row in list(table.rows[6:]):
        table._tbl.remove(row._tr)
    for module, action, expected in cases:
        table._tbl.append(deepcopy(prototype_tr))
        row = table.rows[-1]
        cells = unique_cells(row)
        if len(cells) != 5:
            raise RuntimeError(f"Expected 5 visible cells, found {len(cells)}")
        replace_cell_text(cells[0], module)
        replace_cell_text(cells[1], action)
        replace_cell_text(cells[2], expected)
        replace_cell_text(cells[3], "")
        replace_cell_text(cells[4], "")
        prevent_row_split(row)


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    document = Document(SOURCE)
    if len(document.tables) != 7:
        raise RuntimeError(f"Expected 7 role tables, found {len(document.tables)}")

    prototype_tr = deepcopy(document.tables[1].rows[6]._tr)
    table_keys = [
        "Patient / Mobile Application",
        "Admin / Web Application",
        "Owner / Web Application",
        "Branch Manager / Web Application",
        "Dentist / Web Application",
        "Secretary / Web Application",
        "Patient / Web Application",
    ]

    for table, key in zip(document.tables, table_keys):
        cases = CASES[key]
        fill_table(table, prototype_tr, cases)
        # Clarify the execution setup while retaining the supplied document structure.
        cells3 = unique_cells(table.rows[3])
        cells4 = unique_cells(table.rows[4])
        if key.endswith("Mobile Application"):
            replace_cell_text(cells3[1], "Install and open the mobile app; use prepared patient, appointment, notification, and clinical test data")
            replace_cell_text(cells4[1], "Execute every listed action on supported mobile devices and record Actual Results and Remarks")
        else:
            replace_cell_text(cells3[1], "Open a supported browser; use prepared accounts and test data for the stated role and branch scope")
            replace_cell_text(cells4[1], "Enter the correct URL, execute every listed action, and record Actual Results and Remarks")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    for key in table_keys:
        print(f"{key}: {len(CASES[key])} cases")
    print(f"Total: {sum(len(CASES[key]) for key in table_keys)} cases")


if __name__ == "__main__":
    main()
