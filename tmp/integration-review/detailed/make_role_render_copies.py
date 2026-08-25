from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


SOURCE = Path(r"C:\Users\Administrator\Desktop\NGITIFY DENTIME\output\Ngitify_Integration_Testing_All_Roles_Detailed.docx")
OUTDIR = Path(r"C:\Users\Administrator\Desktop\NGITIFY DENTIME\tmp\integration-review\detailed\role-render-docs")
NAMES = ["mobile-patient", "admin", "owner", "branch-manager", "dentist", "secretary", "web-patient"]


def main():
    OUTDIR.mkdir(parents=True, exist_ok=True)
    master = Document(SOURCE)
    table_xml = [deepcopy(table._tbl) for table in master.tables]
    sect_pr = deepcopy(master.sections[-1]._sectPr)
    for name, selected in zip(NAMES, table_xml):
        doc = Document()
        body = doc._element.body
        for child in list(body):
            body.remove(child)
        body.append(deepcopy(selected))
        body.append(deepcopy(sect_pr))
        output = OUTDIR / f"{name}.docx"
        doc.save(output)
        print(output)


if __name__ == "__main__":
    main()
