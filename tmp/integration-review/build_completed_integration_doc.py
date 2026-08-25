from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(r"C:\Users\Administrator\Documents\integration testing.docx")
OUTPUT = Path(r"C:\Users\Administrator\Desktop\NGITIFY DENTIME\output\Ngitify_Integration_Testing_Completed.docx")


CASES = {
    "mobile_patient": [
        (
            "REGISTRATION + LOGIN",
            "Complete patient pre-registration or staff registration, open the activation link, set a valid password, and sign in through the mobile app.",
            "The patient account is activated once, the mobile login accepts the new credentials, and the patient dashboard displays the same registered profile.",
        ),
        (
            "LOGIN + ACCOUNT SETTINGS",
            "Request a password reset in the mobile app, verify the OTP, save a valid new password, then sign in and open Account Settings.",
            "The reset password is accepted, the old password no longer works, and the authenticated patient can open Account Settings without session or profile errors.",
        ),
        (
            "ACCOUNT SETTINGS + PATIENTS MANAGEMENT",
            "Edit the patient profile in the mobile app, save the changes, then open the same patient record from the web Patients Management module.",
            "Validated profile changes persist and the web patient record shows the same updated non-clinical information without creating a duplicate patient.",
        ),
        (
            "PATIENT APPOINTMENTS + SCHEDULE MANAGEMENT",
            "Book an appointment in the mobile app, then locate the request in the staff web schedule and open My Appointments again.",
            "One appointment is created with the selected branch, service, date, and time; it appears in both the staff schedule and the patient's upcoming appointments.",
        ),
        (
            "PATIENT APPOINTMENTS + PATIENT NOTIFICATIONS",
            "Reschedule an eligible upcoming appointment in the mobile app and review the updated appointment and notification inbox.",
            "The existing appointment is updated without a duplicate, the new schedule appears on staff and patient views, and a related patient notification is available.",
        ),
        (
            "SCHEDULE MANAGEMENT + PATIENT APPOINTMENTS",
            "Have staff change the appointment status, then refresh My Appointments and Notifications in the mobile app.",
            "The mobile appointment displays the staff-selected status and details, and the unread notification count reflects the new appointment update.",
        ),
        (
            "ORAL HEALTH MANAGEMENT",
            "Save a Daily Oral Health Log, open the same date in Calendar, then review the 7-day and 30-day Trends and Recommended Visit Window.",
            "The saved answers appear on the selected date, trend values are recalculated from patient logs, and the recommendation shows its source, reason, and visit timing.",
        ),
        (
            "ORAL HEALTH MANAGEMENT + DENTAL HEALTH EDUCATION",
            "Record oral-health information that has related education content, then open the contextual topic and search the education library.",
            "The contextual education matches the recorded care information, and library category, search, and article views remain available.",
        ),
        (
            "ORAL HEALTH MANAGEMENT + NGITIBOT CARE CONTEXT",
            "After saving oral-health information, open NgitiBot and request an explanation of the patient's care context and recommendation.",
            "NgitiBot displays the patient's current Oral Health Management information, Recommended Visit Window, and related Dental Health Education context in the conversation.",
        ),
        (
            "PATIENT EMR + PATIENT MEDICAL RECORDS",
            "Have the dentist update medical and dental history, treatment logs, or the odontogram, then refresh Records in the mobile app.",
            "The patient sees the permitted, latest EMR information under the correct record section; staff-only controls and restricted data are not exposed.",
        ),
        (
            "AI-ASSISTED RADIOGRAPH + PATIENT RADIOGRAPH EXPLANATION",
            "Have the dentist upload and approve a radiograph review, then open X-Rays from the patient's mobile Records.",
            "The approved radiograph and dentist-approved explanation are available to the correct patient; unapproved AI review content is not published.",
        ),
        (
            "ACCOUNT SETTINGS + PATIENT NOTIFICATIONS",
            "Change notification preferences, trigger an eligible appointment update, open Notifications, and mark the new item as read.",
            "Preferences persist, allowed notifications are delivered, and opening the item updates its read state and unread count consistently.",
        ),
    ],
    "administrator": [
        (
            "WEBSITE ONLINE BOOKING + REGISTRATION",
            "Submit a website appointment request, complete patient pre-registration, and verify the account through the activation or password setup flow.",
            "The booking data carries into one patient onboarding record, the verified account can sign in, and no duplicate patient or appointment is created.",
        ),
        (
            "REGISTRATION + PATIENTS MANAGEMENT + SCHEDULE MANAGEMENT",
            "Open an online-booking patient from Patients Management, complete any required registration details, and place the request on the schedule.",
            "The patient record and schedule entry reference the same patient, branch, service, and appointment details and remain searchable in both modules.",
        ),
        (
            "STAFFS MANAGEMENT + LOGIN",
            "Add a dentist, secretary, or branch manager, use the delivered activation link to set a password, and sign in as the new staff member.",
            "One staff account is created with the selected role and branch access, activation succeeds, and login opens the correct role dashboard.",
        ),
        (
            "STAFF ACCOUNT LIFECYCLE + LOGIN",
            "Deactivate a staff account and verify access is blocked; reactivate or restore it and sign in again.",
            "The lifecycle status persists, inactive or archived staff cannot access protected pages, and restored active staff regain only their assigned role permissions.",
        ),
        (
            "BRANCH MANAGEMENT + STAFFS/PATIENTS MANAGEMENT",
            "Add or edit a branch, assign staff and patients to it, then search the related records and open branch analytics.",
            "The branch appears once, assigned users are linked to it, branch-scoped lists show the correct records, and analytics use the updated branch data.",
        ),
        (
            "PATIENTS MANAGEMENT + PATIENT ACCOUNT LIFECYCLE",
            "Add a patient, edit validated non-clinical information, deactivate the account, then reactivate or restore it.",
            "A single patient record is maintained, edits persist, access follows the current lifecycle status, and lifecycle actions are reflected in patient and archive views.",
        ),
        (
            "PATIENT BRANCH TRANSFER + SCHEDULE MANAGEMENT",
            "Preview and confirm transfer of an active patient to another branch, including review of any upcoming appointments.",
            "The patient branch assignment changes through the dedicated transfer flow, affected upcoming appointments are handled as confirmed, and the patient remains a single record.",
        ),
        (
            "SCHEDULE MANAGEMENT + PATIENT APPOINTMENTS + NOTIFICATIONS",
            "Add or reschedule an appointment, update its status, then verify the patient's appointment list and notification inbox.",
            "The same appointment is updated across staff and patient views, the selected status persists, and the patient receives the corresponding notification.",
        ),
        (
            "INVENTORY MANAGEMENT + MATERIAL USAGE LOG",
            "Add an inventory item and stock batch, then log materials used for a completed appointment.",
            "The batch is available for use, the material usage record links to the completed appointment, and inventory quantity reflects the recorded consumption.",
        ),
        (
            "SYSTEM CONFIGURATION + WEBSITE CONTENT AND MEDIA",
            "Update validated website content or media in System Configuration, save, and open the affected public website page.",
            "The saved configuration persists and the public website displays the updated content or media without changing unrelated sections.",
        ),
        (
            "ARCHIVE REVIEW + ACCOUNT LIFECYCLE",
            "Archive a staff or patient record, locate it in Archive Review, restore it, then repeat the flow for an eligible permanent deletion.",
            "Archived records leave active lists, restoration returns them with valid state, and only eligible records can be permanently deleted after confirmation.",
        ),
        (
            "DATABASE BACKUP + BACKUP HISTORY",
            "Create a manual database backup, monitor completion, then verify and download the completed backup from history.",
            "Progress and final status are recorded, the backup file is available, verification reports a valid result, and the download corresponds to the selected history entry.",
        ),
        (
            "DATABASE BACKUP + AUTOMATIC BACKUP SETTINGS",
            "Enable automatic backup, save a valid schedule, reopen the page, and confirm that a scheduled run is represented in backup history when due.",
            "The schedule settings persist, invalid values are rejected, and automatic backup records use the configured timing and status fields.",
        ),
        (
            "INTEGRITY TOOLS + SAFE AUTO-FIX + AUDIT LOGS",
            "Run all or selected integrity checks, apply Safe Auto-Fix to an eligible issue, then open the related audit or activity details.",
            "Checks report affected records, only eligible issues are changed, the follow-up check reflects the fix, and the action is recorded with actor and result details.",
        ),
        (
            "ACTIVITY LOGS + SYSTEM AUDIT LOGS",
            "Perform representative account, schedule, branch, and configuration changes, then search and open their activity and audit entries.",
            "Matching logs are searchable and their detail views identify the correct actor, role, action, target, time, and recorded outcome.",
        ),
    ],
    "owner": [
        (
            "STAFFS MANAGEMENT + LOGIN",
            "Add a branch manager, dentist, or secretary, complete account activation, and sign in as the new staff member.",
            "The staff account is created once with the selected role and branch assignment, and activation opens the correct role dashboard.",
        ),
        (
            "STAFF ACCOUNT LIFECYCLE + ACTIVITY LOGS",
            "Deactivate, archive, restore, or reissue access for an eligible staff account, then review the resulting activity log.",
            "The staff status and access change as selected, protected access follows the new status, and the owner action is recorded with the target account.",
        ),
        (
            "BRANCH MANAGEMENT + BRANCH ANALYTICS",
            "Add a branch, update its information or status, assign operational records, and open branch analytics.",
            "The branch list shows the saved state, assignments use the correct branch, inactive status is enforced where applicable, and analytics reflect the branch records.",
        ),
        (
            "PATIENTS MANAGEMENT + PATIENT ACCOUNT LIFECYCLE",
            "Add or edit a patient, then deactivate and restore the account while checking active and archived patient lists.",
            "One validated patient record is maintained and its availability, login access, and archive placement follow the current lifecycle status.",
        ),
        (
            "PATIENT BRANCH TRANSFER + SCHEDULE MANAGEMENT",
            "Transfer an active patient to a target branch after reviewing the transfer impact and upcoming appointments.",
            "The dedicated transfer action updates branch ownership once, handles affected appointments as confirmed, and preserves the patient's EMR identity.",
        ),
        (
            "SCHEDULE MANAGEMENT + PATIENT APPOINTMENTS + NOTIFICATIONS",
            "Create or edit a schedule entry and update its status, then verify the patient-facing appointment and notification.",
            "Staff and patient views show the same appointment details and status, with one related patient notification and no duplicate schedule entry.",
        ),
        (
            "INVENTORY MANAGEMENT + MATERIAL USAGE LOG",
            "Receive a stock batch and record materials used for a completed dental appointment.",
            "The usage log references the correct appointment and item, and the available stock and low-stock state reflect the recorded quantity.",
        ),
        (
            "PATIENT EMR + PATIENT MEDICAL RECORDS",
            "As an owner with dentist access, update clinical history, treatment logs, or odontogram data, then verify the patient record view.",
            "Authorized clinical updates persist in the same EMR and the patient can view the permitted latest information without staff editing controls.",
        ),
        (
            "AI-ASSISTED RADIOGRAPH + PATIENT RADIOGRAPH EXPLANATION",
            "As an owner with dentist access, approve a radiograph review and verify the result in the patient's Medical Records.",
            "Only the dentist-approved radiograph explanation is published to the correct patient, while unapproved AI content remains unavailable.",
        ),
        (
            "ACCOUNT SETTINGS + NOTIFICATIONS + ACTIVITY LOGS",
            "Update owner account settings and notification preferences, trigger an operational event, then review notifications and activity details.",
            "Settings persist, eligible notifications appear with correct read state, and the related owner action is recorded in Activity Logs.",
        ),
    ],
    "branch_manager": [
        (
            "STAFFS MANAGEMENT + LOGIN",
            "Add a dentist or secretary for the manager's assigned branch, complete activation, and sign in as that staff member.",
            "The account is created once under the manager's branch and opens the correct role dashboard; accounts for other branches remain outside scope.",
        ),
        (
            "STAFF ACCOUNT LIFECYCLE + ACTIVITY LOGS",
            "Change the lifecycle status of an eligible dentist or secretary in the assigned branch and review the activity entry.",
            "The allowed status change persists only for in-scope staff, access follows the new status, and the action is logged with branch and actor details.",
        ),
        (
            "PATIENTS MANAGEMENT + PATIENT ACCOUNT LIFECYCLE",
            "Add or edit a branch patient, then deactivate and reactivate or restore the account.",
            "A single validated patient record is maintained in the assigned branch, and visibility and access follow the current lifecycle status.",
        ),
        (
            "PATIENT BRANCH TRANSFER + SCHEDULE MANAGEMENT",
            "Preview and confirm transfer of an active in-scope patient, including the listed upcoming appointment impact.",
            "The patient and confirmed related appointments are handled through one transfer transaction; out-of-branch patients cannot be changed.",
        ),
        (
            "SCHEDULE MANAGEMENT + PATIENT APPOINTMENTS",
            "Add or reschedule an appointment for an in-branch patient and update its status.",
            "The branch schedule and patient appointment view show the same details and status, with no duplicate appointment created.",
        ),
        (
            "SCHEDULE MANAGEMENT + PATIENT NOTIFICATIONS",
            "Confirm, reschedule, or cancel an appointment and verify the affected patient's notification inbox.",
            "The schedule change persists and produces the appropriate patient notification with a consistent unread count and appointment reference.",
        ),
        (
            "INVENTORY MANAGEMENT + SCHEDULE MANAGEMENT",
            "Receive stock for the assigned branch and review inventory availability before a scheduled procedure.",
            "The new batch appears under the correct inventory item and branch, and branch users see the updated quantity and low-stock state.",
        ),
        (
            "BRANCH ANALYTICS + SCHEDULE/PATIENT DATA",
            "Create representative patient and schedule activity, then open branch analytics and filter or refresh the results.",
            "Analytics reflect the assigned branch's current patients and appointments and do not include records owned by another branch.",
        ),
        (
            "NOTIFICATIONS + ACTIVITY LOGS",
            "Perform staff, patient, schedule, and inventory actions, then review notifications and search Activity Logs.",
            "Eligible operational notifications appear, and matching activity details identify the correct branch manager, action, target, and time.",
        ),
    ],
    "dentist": [
        (
            "SCHEDULE MANAGEMENT + PATIENTS MANAGEMENT",
            "Open an assigned appointment from the dentist schedule and navigate to the same patient's record.",
            "The selected appointment resolves to the correct patient, while patients outside the dentist's assignment are not accessible.",
        ),
        (
            "PATIENT EMR + MEDICAL AND DENTAL HISTORY",
            "Update the assigned patient's medical and dental history, save, and reopen the EMR.",
            "Validated history changes persist in the same patient EMR and become available in the permitted patient Medical Records view.",
        ),
        (
            "PATIENT EMR + TREATMENT LOGS",
            "Add a treatment log for the scheduled visit, then reopen Treatment History from staff and patient views.",
            "One treatment record is stored for the correct patient and visit, and the permitted treatment information is visible to the patient.",
        ),
        (
            "PATIENT EMR + INTERACTIVE DIGITAL ODONTOGRAM",
            "Record tooth conditions or treatment updates in the odontogram, save, and open the patient's odontogram view.",
            "The tooth-level changes persist under the correct patient and the patient sees a read-only representation of the latest approved data.",
        ),
        (
            "AI-ASSISTED RADIOGRAPH + PATIENT RADIOGRAPH EXPLANATION",
            "Upload a patient radiograph, review AI-assisted findings, approve or revise the dentist summary, then open the patient X-Ray view.",
            "The image remains linked to the correct patient, dentist verification is recorded, and only the approved explanation is published to the patient.",
        ),
        (
            "SCHEDULE STATUS + PATIENT NOTIFICATIONS",
            "Update the assigned appointment from confirmed to the appropriate visit status, then verify the patient's appointment and notification.",
            "The status is consistent across dentist schedule and patient view, and the patient receives the corresponding notification once.",
        ),
        (
            "COMPLETED APPOINTMENT + MATERIAL USAGE LOG",
            "For a completed appointment, log the materials used and reopen the Material Usage list.",
            "The material usage entry references the correct dentist, patient, appointment, item, and quantity, and the inventory balance is adjusted.",
        ),
        (
            "NGITIBOT CARE CONTEXT + PATIENT EMR",
            "After saving relevant care information, open the AI assistant or patient care companion context for the same patient.",
            "The assistant uses available current care context and recommendations without exposing another patient's information or replacing dentist judgment.",
        ),
        (
            "ACCOUNT SETTINGS + NOTIFICATIONS + ACTIVITY LOGS",
            "Update dentist settings, perform an EMR or schedule action, then review Notifications and Activity Logs.",
            "Settings persist, relevant notifications display with correct read state, and the clinical or schedule action is recorded with the dentist as actor.",
        ),
    ],
    "secretary": [
        (
            "WEBSITE ONLINE BOOKING + SCHEDULE MANAGEMENT",
            "Locate a submitted website appointment request, validate the patient details, and add or confirm it in the branch schedule.",
            "The request becomes one schedule entry for the correct patient, branch, service, date, and time without duplicate booking data.",
        ),
        (
            "REGISTRATION + PATIENTS MANAGEMENT",
            "Complete an online-booking patient's pre-registration or add a new patient, then send or use the patient activation link.",
            "One branch patient record is created, required fields are retained, and the patient can set a password and sign in after verification.",
        ),
        (
            "PATIENTS MANAGEMENT + SCHEDULE MANAGEMENT",
            "Edit allowed patient information and open the patient's existing appointment from the schedule.",
            "Validated non-clinical edits persist on the same patient record and the schedule continues to reference that patient without losing appointment details.",
        ),
        (
            "PATIENT BRANCH TRANSFER + SCHEDULE MANAGEMENT",
            "Preview and confirm transfer of an active in-branch patient after reviewing affected upcoming appointments.",
            "The patient branch changes through the dedicated transfer action, affected appointments are handled as confirmed, and restricted fields remain protected.",
        ),
        (
            "SCHEDULE MANAGEMENT + PATIENT APPOINTMENTS",
            "Add, reschedule, or cancel a patient appointment and update its allowed status.",
            "The same appointment is updated in the staff schedule and patient appointment list, and no extra appointment is created.",
        ),
        (
            "SCHEDULE MANAGEMENT + PATIENT NOTIFICATIONS",
            "Change appointment details or status, then verify the affected patient's notification inbox.",
            "A related patient notification is created with the correct appointment reference, and its unread state is reflected in the patient interface.",
        ),
        (
            "PATIENTS MANAGEMENT + PATIENT MEDICAL RECORDS",
            "Open an in-branch patient's record and inspect the permitted overview or imaging information without attempting clinical edits.",
            "The secretary can view only authorized patient information; restricted clinical fields and dentist-only editing actions remain unavailable.",
        ),
        (
            "NOTIFICATIONS + ACTIVITY LOGS",
            "Perform patient and schedule actions, open Notifications, then search and view the related Activity Log details.",
            "Eligible notifications display and matching logs identify the secretary, action, patient or appointment target, branch, and time.",
        ),
    ],
    "web_patient": [
        (
            "WEBSITE ONLINE BOOKING + REGISTRATION + LOGIN",
            "Submit an online booking request, complete pre-registration and account activation, then sign in through the patient web portal.",
            "One verified patient account and related appointment request are created, and login opens the dashboard with the same patient identity.",
        ),
        (
            "LOGIN + ACCOUNT SETTINGS",
            "Complete Forgot Password, OTP verification, and Reset Password, then sign in and open Account Settings.",
            "The new password is accepted, the old password is rejected, and the authenticated patient can access the correct profile and settings.",
        ),
        (
            "ACCOUNT SETTINGS + PATIENTS MANAGEMENT",
            "Edit the patient profile in the web portal, save, and verify the same non-clinical information from staff Patients Management.",
            "Validated changes persist on the existing patient record and are consistent across patient and authorized staff views.",
        ),
        (
            "PATIENT APPOINTMENTS + SCHEDULE MANAGEMENT",
            "Book an appointment in the web portal and locate it in the staff schedule and My Appointments.",
            "One appointment is created with matching patient, branch, service, date, and time across staff and patient views.",
        ),
        (
            "PATIENT APPOINTMENTS + PATIENT NOTIFICATIONS",
            "Reschedule or cancel an eligible appointment, then reopen My Appointments and Notifications.",
            "The existing appointment is updated without duplication, both interfaces show the new state, and the related notification is available.",
        ),
        (
            "ORAL HEALTH MANAGEMENT + DENTAL HEALTH EDUCATION",
            "Save a Daily Oral Health Log, review Calendar and Trends, then open a contextual education topic and the education library.",
            "The log persists by date, trends and Recommended Visit Window use the saved information, and relevant education content is accessible.",
        ),
        (
            "ORAL HEALTH MANAGEMENT + NGITIBOT CARE CONTEXT",
            "Open NgitiBot after oral-health logs and a recommendation are available, then ask about the current care context.",
            "NgitiBot displays patient-specific Oral Health Management information, Recommended Visit Window, and related education context in the conversation.",
        ),
        (
            "PATIENT EMR + PATIENT MEDICAL RECORDS",
            "After a dentist updates history, treatment, or odontogram data, refresh the patient Medical Records sections.",
            "The permitted latest EMR information appears under the correct patient and section, while staff-only controls remain unavailable.",
        ),
        (
            "AI-ASSISTED RADIOGRAPH + PATIENT RADIOGRAPH EXPLANATION",
            "After dentist approval of a radiograph review, open X-Rays and its explanation in the patient portal.",
            "The correct image and approved dentist explanation display; unapproved AI findings and another patient's radiographs remain inaccessible.",
        ),
        (
            "ACCOUNT SETTINGS + PATIENT NOTIFICATIONS",
            "Save notification preferences, trigger an eligible appointment event, open the notification, and verify its read state.",
            "Preferences persist, allowed events create notifications, and the selected item's read status and unread count update consistently.",
        ),
    ],
}


TABLE_ORDER = [
    ("mobile_patient", "Patient", "Mobile Application", "Launch the mobile app; use an active or activatable patient account and prepared staff test accounts."),
    ("administrator", "Administrator", "Web Application", "Open the correct web URL; use an active administrator account and prepared branch, staff, patient, appointment, inventory, and backup test data."),
    ("owner", "Owner", "Web Application", "Open the correct web URL; use an active owner account, including owner-with-dentist-access data where specified."),
    ("branch_manager", "Branch Manager", "Web Application", "Open the correct web URL; use an active branch manager assigned to a test branch with in-scope and out-of-scope records."),
    ("dentist", "Dentist", "Web Application", "Open the correct web URL; use an active dentist with assigned patients and appointments plus prepared radiograph and inventory data."),
    ("secretary", "Secretary", "Web Application", "Open the correct web URL; use an active secretary assigned to a test branch with booking, patient, and schedule data."),
    ("web_patient", "Patient", "Web Application", "Open the correct web URL; use an active or activatable patient account and prepared staff test accounts."),
]


def set_cell_text(cell, text, *, size=8.5, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=120, bottom=70, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_table_geometry(table):
    widths = [1940, 1840, 1840, 2240, 760, 740]
    table.autofit = False
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = table_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        table_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_layout = table_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        table_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid_cols = table._tbl.tblGrid.gridCol_lst
    for grid_col, width in zip(grid_cols, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        grid_index = 0
        for tc in row._tr.tc_lst:
            span_node = tc.tcPr.gridSpan
            span = int(span_node.val) if span_node is not None else 1
            cell_width = sum(widths[grid_index:grid_index + span])
            tc_w = tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(cell_width))
            tc_w.set(qn("w:type"), "dxa")
            grid_index += span


def clear_data_rows(table):
    while len(table.rows) > 6:
        table._tbl.remove(table.rows[-1]._tr)


def configure_metadata(table, access_type, system_type, precondition):
    set_cell_text(table.cell(0, 0), "INTEGRATION\nTEST DOCUMENT", size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(0, 2), "Application/System\nName", size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(
        table.cell(0, 3),
        "NGITIFY: A DENTAL INFORMATION MANAGEMENT SYSTEM WITH AI-DRIVEN PATIENT ENGAGEMENT FOR DENTIME DENTAL CLINIC",
        size=8.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_cell_text(table.cell(1, 0), "Test Cycle No.", size=8.5, bold=True)
    set_cell_text(table.cell(1, 1), "1", size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(1, 2), "User/Access Type", size=8.5, bold=True)
    set_cell_text(table.cell(1, 3), access_type, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(2, 0), "Date Tested", size=8.5, bold=True)
    set_cell_text(table.cell(2, 1), "08/25/2026", size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(2, 2), "Type of System", size=8.5, bold=True)
    set_cell_text(table.cell(2, 3), system_type, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(3, 0), "Pre-conditions", size=8.5, bold=True)
    set_cell_text(table.cell(3, 1), precondition, size=8.5)
    set_cell_text(table.cell(4, 0), "Verification Steps", size=8.5, bold=True)
    set_cell_text(
        table.cell(4, 1),
        "Execute each cross-module action in sequence; refresh the receiving view; verify persisted data, role/branch scope, notifications, and absence of duplicate records or errors.",
        size=8.5,
    )
    headers = {
        0: "Modules",
        1: "Action Description",
        3: "Expected Results",
        4: "Actual Results",
        5: "Remarks",
    }
    for idx, value in headers.items():
        set_cell_text(table.cell(5, idx), value, size=8.2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row_index in range(6):
        prevent_row_split(table.rows[row_index])
        for cell in table.rows[row_index].cells:
            set_cell_margins(cell, top=45, bottom=45)
    for cell in table.rows[5].cells:
        shade_cell(cell, "D9EAF7")


def add_case(table, module, action, expected):
    row = table.add_row()
    action_cell = row.cells[1].merge(row.cells[2])
    set_cell_text(row.cells[0], module, size=8.2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(1, 83, 139))
    set_cell_text(action_cell, action, size=8.2)
    set_cell_text(row.cells[3], expected, size=8.2)
    set_cell_text(row.cells[4], "", size=8.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[5], "", size=8.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    prevent_row_split(row)
    for cell in row.cells:
        set_cell_margins(cell)


def main():
    document = Document(SOURCE)
    if len(document.tables) != len(TABLE_ORDER):
        raise ValueError(f"Expected {len(TABLE_ORDER)} role tables, found {len(document.tables)}")

    for table, (key, access_type, system_type, precondition) in zip(document.tables, TABLE_ORDER):
        clear_data_rows(table)
        configure_metadata(table, access_type, system_type, precondition)
        for module, action, expected in CASES[key]:
            add_case(table, module, action, expected)
        set_table_geometry(table)

    for paragraph in document.paragraphs[:-1]:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.page_break_before = False

    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.25)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.core_properties.title = "Ngitify Integration Testing"
    document.core_properties.subject = "Role-based end-to-end integration test matrix aligned with mobile and web unit testing"
    document.core_properties.author = "Ngitify Project Team"
    document.save(OUTPUT)
    print(f"Saved {OUTPUT}")
    print("Case counts:", {key: len(value) for key, value in CASES.items()})
    print("Total cases:", sum(len(value) for value in CASES.values()))


if __name__ == "__main__":
    main()
