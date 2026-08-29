from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE = Path(r"C:\Users\Administrator\Desktop\NGITIFY DENTIME\output\Ngitify_Integration_Testing_Valid_Inputs_Sample_Wording.docx")
OUTPUT = Path(r"C:\Users\Administrator\Desktop\NGITIFY DENTIME\output\Ngitify_Integration_Testing_Final_FDD.docx")


def c(module, action, expected):
    return module, action, expected


def auth(role, mobile=False):
    app = "mobile" if mobile else "web"
    return [
        c("Login", f"Click the Ngitify {app} application", "Displays the Login page"),
        c("Login / Forgot Password", "Click Forgot Password button", "Displays the Forgot Password page"),
        c("Login / Forgot Password", "Input the correct registered email, then click Send Code button", "System verifies the email and displays the OTP page"),
        c("Login / Forgot Password", "Input the correct OTP, then click Verify button", "System verifies the OTP and displays the Reset Password page"),
        c("Login / Forgot Password", "Input the correct new password and confirm new password, then click Reset Password button", "System updates the password in the database and displays a success message"),
        c("Login / Forgot Password", "Click Back to Login button", "Displays the Login page"),
        c("Login", "Input the correct email and password, then click Login button", f"System verifies the account and displays the {role} dashboard"),
    ]


def dashboard(module="Dashboard"):
    return [c(module, f"Click the {module} on the navigation menu", "Displays the dashboard summary")]


def simple_view(module, label, result=None):
    result = result or f"Displays the {label.lower()}"
    return c(module, f"Click {label}", result)


def edit_flow(module, entity, button_label=None):
    button_label = button_label or f"Edit {entity}"
    target = entity.lower()
    return [
        c(module, f"Click {button_label} button", f"Displays the edit {target} form"),
        c(module, "Input all fields completely and in the correct format, then click Save button", "Displays a user confirmation prompt"),
        c(module, "Click Confirm button", f"System removes the confirmation prompt, updates the {target} information in the database, and displays a success message"),
        c(module, "Click OK button", "Removes the success message"),
        c(module, "Click Cancel button", f"Removes the confirmation prompt and leaves the {target} information unchanged"),
    ]


def create_flow(module, entity, button_label=None):
    button_label = button_label or f"Create {entity}"
    target = entity.lower()
    return [
        c(module, f"Click {button_label} button", f"Displays the create {target} form"),
        c(module, "Input all fields completely and in the correct format, then click Save button", "Displays a user confirmation prompt"),
        c(module, "Click Confirm button", f"System removes the confirmation prompt, saves the {target} information in the database, and displays a success message"),
        c(module, "Click OK button", "Removes the success message"),
        c(module, "Click Cancel button", f"Removes the confirmation prompt and leaves the {target} form unchanged"),
    ]


def status_flow(module, action, entity, result_verb):
    target = entity.lower()
    return [
        c(module, f"Click {action} button", "Displays a user confirmation prompt"),
        c(module, "Click Confirm button", f"System removes the confirmation prompt, {result_verb} the {target}, and displays a success message"),
        c(module, "Click OK button", "Removes the success message"),
        c(module, "Click Cancel button", f"Removes the confirmation prompt and leaves the {target} unchanged"),
    ]


def notifications_activity_settings_logout(notification_module="Notification"):
    rows = [
        c(notification_module, f"Click {notification_module} on the navigation menu", "Displays the Notifications page"),
        c("Activity Logs", "Click Activity Logs on the navigation menu", "Displays the Activity Logs page"),
        c("Account Settings / View Profile", "Click View Profile button", "Displays the user profile information"),
    ]
    rows += edit_flow("Account Settings / Edit Profile", "profile")
    rows += [
        c("Account Settings / Change Password", "Click Change Password button", "Displays the Change Password form"),
        c("Account Settings / Change Password", "Input the correct current password, new password, and confirm new password, then click Save button", "Displays a user confirmation prompt"),
        c("Account Settings / Change Password", "Click Confirm button", "System removes the confirmation prompt, updates the password in the database, and displays a success message"),
        c("Account Settings / Change Password", "Click OK button", "Removes the success message"),
        c("Account Settings / Change Password", "Click Cancel button", "Removes the confirmation prompt and leaves the password unchanged"),
        c("Logout", "Click Logout button", "Displays a user confirmation prompt"),
        c("Logout", "Click Confirm button", "System logs the user out and displays the Login page"),
        c("Logout", "Click Cancel button", "Removes the confirmation prompt and leaves the user logged in"),
    ]
    return rows


def schedule_management():
    rows = [
        simple_view("Schedule Management", "View Schedule List", "Displays the schedule list"),
        simple_view("Schedule Management", "View Schedule Details", "Displays the selected schedule details"),
    ]
    rows += edit_flow("Schedule Management / Edit Schedule Details", "schedule details", "Edit Schedule Details")
    rows += status_flow("Schedule Management / Approve Schedule Entry", "Approve Schedule Entry", "schedule entry", "approves")
    rows += status_flow("Schedule Management / Reject Schedule Entry", "Reject Schedule Entry", "schedule entry", "rejects")
    return rows


def user_accounts_management(final_action="Reset User Password"):
    rows = [
        simple_view("User Accounts Management", "View User Accounts List", "Displays the user accounts list"),
        c("User Accounts Management / Search User Accounts", "Input a valid user name or email, then click Search button", "System searches the database and displays matching user accounts"),
    ]
    rows += create_flow("User Accounts Management / Create User Account", "user account", "Create User Account")
    rows += edit_flow("User Accounts Management / Edit User Account Details", "user account", "Edit User Account Details")
    rows += status_flow("User Accounts Management / Deactivate User Account", "Deactivate User Account", "user account", "deactivates")
    if final_action == "Archive User Account":
        rows += status_flow("User Accounts Management / Archive User Account", "Archive User Account", "user account", "archives")
    else:
        rows += status_flow("User Accounts Management / Reset User Password", "Reset User Password", "user password", "resets")
    return rows


def patient_profile(edit=True, export=False):
    rows = [simple_view("Patient EMR / Patient Profile", "View Patient Profile", "Displays the selected patient profile")]
    if edit:
        rows += edit_flow("Patient EMR / Patient Profile", "patient profile", "Edit Patient Profile")
    if export:
        rows.append(c("Patient EMR / Patient Profile", "Click Export PDF button", "System downloads the patient profile PDF"))
    return rows


def staff_emr(history_edit=False, history_print=False, radiograph_upload=False, radiograph_delete=False):
    rows = patient_profile(edit=False)
    rows += [
        simple_view("Patient EMR / Medical and Dental History", "View Medical History", "Displays the selected patient medical history"),
        simple_view("Patient EMR / Medical and Dental History", "View Dental History", "Displays the selected patient dental history"),
    ]
    if history_edit:
        rows += edit_flow("Patient EMR / Medical and Dental History", "medical history", "Edit Medical History")
        rows += edit_flow("Patient EMR / Medical and Dental History", "dental history", "Edit Dental History")
    if history_print:
        rows += [
            c("Patient EMR / Medical and Dental History", "Click Print Medical History button", "System displays the medical history print window"),
            c("Patient EMR / Medical and Dental History", "Click Print Dental History button", "System displays the dental history print window"),
        ]
    rows += [
        simple_view("Patient EMR / Treatment History", "View Treatment Records", "Displays the selected patient treatment records"),
        simple_view("Patient EMR / Treatment History", "View Treatment Notes", "Displays the selected treatment notes"),
        simple_view("Patient EMR / Digital Odontogram", "View Digital Odontogram", "Displays the selected patient digital odontogram"),
        simple_view("Patient EMR / Digital Odontogram", "View Tooth History", "Displays the selected tooth history"),
        simple_view("Patient EMR / Radiograph", "View X-ray Images", "Displays the selected patient X-ray images"),
    ]
    if radiograph_upload:
        rows += create_flow("Patient EMR / Radiograph", "radiograph", "Upload Radiograph")
    if radiograph_delete:
        rows += status_flow("Patient EMR / Radiograph", "Delete Radiograph", "radiograph", "removes")
    return rows


def material_usage():
    return [
        c("Material Usage Logging", "Click Log New Entry button", "Displays the material usage entry form"),
        c("Material Usage Logging", "Input all fields completely and in the correct format, then click Save button", "Displays a user confirmation prompt"),
        c("Material Usage Logging", "Click Confirm button", "System removes the confirmation prompt, saves the material usage entry in the database, and displays a success message"),
        c("Material Usage Logging", "Click OK button", "Removes the success message"),
        c("Material Usage Logging", "Click Cancel button", "Removes the confirmation prompt and leaves the material usage information unchanged"),
        simple_view("Material Usage Logging", "View Total Logs Count", "Displays the total material usage logs count"),
        simple_view("Material Usage Logging", "View Monthly Logs Count", "Displays the monthly material usage logs count"),
        simple_view("Material Usage Logging", "View Most Used Item", "Displays the most used material item"),
        c("Material Usage Logging", "Input a valid procedure or patient name, then click Search button", "System searches the database and displays matching material usage logs"),
        c("Material Usage Logging", "Click Start Date filter", "Displays material usage logs from the selected start date"),
        c("Material Usage Logging", "Click End Date filter", "Displays material usage logs up to the selected end date"),
        simple_view("Material Usage Logging", "View Material Usage Logs List", "Displays the material usage logs list"),
    ]


def ai_image_enhancer(module="AI-Assisted Image Enhancer"):
    return [
        c(module, "Click Select Radiograph Image button", "Displays the radiograph image selection window"),
        c(module, "Click Run Adaptive Enhancement button", "System enhances the selected radiograph image and displays a success message"),
        c(module, "Click View Enhanced Image button", "Displays the enhanced radiograph image"),
        c(module, "Click Review AI Summary button", "Displays the AI radiograph summary"),
        c(module, "Input the correct radiograph findings, then click Save Findings button", "Displays a user confirmation prompt"),
        c(module, "Click Confirm button", "System removes the confirmation prompt, saves the radiograph findings in the database, and displays a success message"),
        c(module, "Click OK button", "Removes the success message"),
        c(module, "Click Cancel button", "Removes the confirmation prompt and leaves the radiograph findings unchanged"),
    ]


def supply_stock():
    rows = [
        simple_view("Supply and Stock Monitoring", "View Total Stock Count", "Displays the total stock count"),
        simple_view("Supply and Stock Monitoring", "View Item Categories List", "Displays the item categories list"),
        simple_view("Supply and Stock Monitoring", "View Low Stock Warning Banners", "Displays the low stock warning banners"),
    ]
    rows += edit_flow("Supply and Stock Monitoring / Update Item Stock Quantity", "item stock quantity", "Update Item Stock Quantity")
    rows += create_flow("Supply and Stock Monitoring / Add New Supply Entry", "supply entry", "Add New Supply Entry")
    rows += edit_flow("Supply and Stock Monitoring / Edit Item Details", "item details", "Edit Item Details")
    rows += [
        c("Supply and Stock Monitoring", "Input a valid item name or code, then click Search button", "System searches the database and displays matching supply items"),
        c("Supply and Stock Monitoring", "Click the In Stock, Low Stock, or Out of Stock filter", "Displays supply items with the selected stock status"),
    ]
    return rows


def branch_analytics():
    rows = [
        simple_view("Branch Management and Expansion Analytics", "View Clinic Branches List", "Displays the clinic branches list"),
        simple_view("Branch Management and Expansion Analytics", "View Individual Branch Analytics", "Displays the selected branch analytics"),
    ]
    rows += edit_flow("Branch Management and Expansion Analytics", "branch details", "Edit Branch Details")
    return rows


def emr_synchronization():
    rows = [
        simple_view("EMR Synchronization", "View Real-time Synchronization Status", "Displays the real-time synchronization status"),
        simple_view("EMR Synchronization", "View Last Synchronized Timestamp", "Displays the last synchronized timestamp"),
        simple_view("EMR Synchronization", "View Synchronization Conflict Logs", "Displays the synchronization conflict logs"),
        c("EMR Synchronization", "Click Manual Sync button", "System refreshes the synchronized EMR data and displays the synchronization status"),
        c("EMR Synchronization", "Click Retry Failed Database Sync Entries button", "System retries the failed synchronization entries and displays the updated status"),
        simple_view("EMR Synchronization", "View Data Conflict Details", "Displays the selected data conflict details"),
        c("EMR Synchronization", "Click the overriding record", "Displays the selected overriding record"),
    ]
    rows += status_flow("EMR Synchronization / Resolve Data Conflict Entry", "Resolve Data Conflict Entry", "data conflict entry", "resolves")
    return rows


def branch_transfer():
    rows = [
        simple_view("Branch Transfer Management", "View Pending Branch Transfer Requests", "Displays the pending branch transfer requests"),
        simple_view("Branch Transfer Management", "View Completed Transfer Logs History", "Displays the completed transfer logs history"),
        c("Branch Transfer Management", "Input a valid item name or branch, then click Search button", "System searches the database and displays matching transfer records"),
        c("Branch Transfer Management", "Click Create Transfer Request button", "Displays the create transfer request form"),
        c("Branch Transfer Management", "Click the source branch", "Displays the selected source branch"),
    ]
    return rows


def database_backup():
    return [
        c("Database Backup", "Click Refresh button", "System refreshes the database backup information"),
        simple_view("Database Backup", "View Backup Storage Status", "Displays the backup storage status"),
        simple_view("Database Backup", "View Local Backup Availability", "Displays the local backup availability"),
        simple_view("Database Backup", "View Cloud Backup Availability", "Displays the cloud backup availability"),
        simple_view("Database Backup", "View Scheduled Backup Interval Time", "Displays the scheduled backup interval time"),
        simple_view("Database Backup", "View Retention Period", "Displays the database backup retention period"),
        c("Database Backup", "Click Create Backup button", "Displays a user confirmation prompt"),
        c("Database Backup", "Click Confirm button", "System removes the confirmation prompt, creates the database backup, and displays a success message"),
        c("Database Backup", "Click OK button", "Removes the success message"),
        c("Database Backup", "Click Cancel button", "Removes the confirmation prompt and does not create a database backup"),
        simple_view("Database Backup", "View Backup History", "Displays the database backup history"),
        simple_view("Database Backup", "View Successful and Failed Backup Count", "Displays the successful and failed database backup count"),
        c("Database Backup", "Click Download Available Backup button", "System downloads the selected available database backup"),
    ]


def integrity_tools():
    return [
        c("Integrity Tools", "Click Run Data Validation Scan button", "System runs the data validation scan and displays the validation results"),
        simple_view("Integrity Tools", "View Database Error Reports", "Displays the database error reports"),
    ]


def system_configuration():
    return [
        c("System Configuration", "Input the correct clinic name, contact number, email address, and address, then click Save Changes button", "Displays a user confirmation prompt"),
        c("System Configuration", "Click Confirm button", "System removes the confirmation prompt, updates the system configuration in the database, and displays a success message"),
        c("System Configuration", "Click OK button", "Removes the success message"),
        c("System Configuration", "Click Cancel button", "Removes the confirmation prompt and leaves the system configuration unchanged"),
    ]


def audit_trail():
    return [
        c("Audit Trail", "Input a valid audit log keyword, then click Search button", "System searches the database and displays matching system audit logs"),
        c("Audit Trail", "Click the user or date filter", "Displays audit logs for the selected user or date"),
        c("Audit Trail", "Click Export Audit Trail Report button", "System downloads the audit trail report"),
    ]


def patient_readonly_emr(mobile=False):
    prefix = "Electronic Medical Record (EMR)" if mobile else "Patient EMR"
    rows = []
    if not mobile:
        rows += [
            simple_view(f"{prefix} / Patient Profile", "View Patient Details Header", "Displays the patient details header"),
        ]
        rows += edit_flow(f"{prefix} / Patient Profile", "patient profile", "Edit Patient Profile")
        rows.append(c(f"{prefix} / Patient Profile", "Click Export PDF button", "System downloads the patient profile PDF"))
    rows += [
        c(f"{prefix} / Medical and Dental History", "Click Medical and Dental History tab", "Displays the Medical and Dental History section"),
        simple_view(f"{prefix} / Medical and Dental History", "View Medical History Details", "Displays the patient medical history details"),
        simple_view(f"{prefix} / Medical and Dental History", "View Dental History Details", "Displays the patient dental history details"),
        c(f"{prefix} / Treatment History", "Click Treatment History tab", "Displays the Treatment History section"),
        simple_view(f"{prefix} / Treatment History", "View Past Treatment Logs", "Displays the past treatment logs"),
        c(f"{prefix} / Odontogram", "Click Odontogram tab", "Displays the patient odontogram"),
        simple_view(f"{prefix} / Odontogram", "View Personal Tooth Chart Diagram", "Displays the personal tooth chart diagram"),
        c(f"{prefix} / Radiograph Images", "Click Radiographs tab", "Displays the Radiograph Images section"),
        simple_view(f"{prefix} / Radiograph Images", "View Uploaded X-ray Images", "Displays the uploaded X-ray images"),
    ]
    if not mobile:
        rows.append(simple_view(f"{prefix} / Radiograph Images", "View Radiograph Summary and Findings", "Displays the radiograph summary and findings"))
    return rows


def patient_ai_engagement(mobile=False):
    module = "AI Patient Care Companion" if mobile else "AI Patient Engagement"
    return [
        c(f"{module} / NgitiBot", "Click NgitiBot button", "Displays the NgitiBot conversation window"),
        c(f"{module} / NgitiBot", "Input a patient care question, then click Send button", "System processes the question and displays the NgitiBot response"),
        c(f"{module} / Dental Health Education", "Click Dental Health Education button", "Displays the dental health educational materials"),
        c(f"{module} / Oral Health Management", "Click Oral Health Management button", "Displays the Oral Health Management page"),
        c(f"{module} / Predictive Visit", "Click Predictive Visit button", "Displays the dentist-recommended visit or AI-generated recommendations for the next visit"),
    ]


def patient_appointments(module="Appointment Scheduling"):
    return [
        c(module, "Click Book Appointment button", "Displays the appointment booking form"),
        c(module, "Input all appointment fields completely and in the correct format, then click Book Appointment button", "Displays a user confirmation prompt"),
        c(module, "Click Confirm button", "System removes the confirmation prompt, saves the appointment in the database, and displays a success message"),
        c(module, "Click OK button", "Removes the success message"),
        c(module, "Click Cancel button", "Removes the confirmation prompt and leaves the appointment booking form unchanged"),
        c(module, "Click Reschedule Appointment button", "Displays the appointment rescheduling form"),
        c(module, "Input the new appointment schedule, then click Reschedule button", "Displays a user confirmation prompt"),
        c(module, "Click Confirm Reschedule button", "System updates the appointment schedule in the database and displays a success message"),
        c(module, "Click Cancel Appointment button", "Displays a user confirmation prompt"),
        c(module, "Click Confirm Cancellation button", "System cancels the appointment in the database and displays a success message"),
        c(module, "Click Cancel button", "Removes the confirmation prompt and leaves the appointment unchanged"),
    ]


def website_booking():
    return [
        c("Website Booking", "Click Book Appointment button", "Displays the website appointment booking form"),
        c("Website Booking", "Input all appointment fields completely and in the correct format, then click Book Appointment button", "Displays a user confirmation prompt"),
        c("Website Booking", "Click Confirm button", "System removes the confirmation prompt, saves the website booking in the database, and displays a success message"),
        c("Website Booking", "Click OK button", "Removes the success message"),
        c("Website Booking", "Click Cancel button", "Removes the confirmation prompt and leaves the website booking form unchanged"),
    ]


MOBILE_PATIENT = (
    auth("patient", mobile=True)
    + dashboard("Dashboard (Smile Hub)")
    + patient_appointments()
    + patient_ai_engagement(mobile=True)
    + patient_readonly_emr(mobile=True)
    + notifications_activity_settings_logout("Notifications")
)


ADMIN = (
    auth("administrator")
    + dashboard()
    + schedule_management()
    + user_accounts_management(final_action="Archive User Account")
    + staff_emr(history_edit=True, history_print=False)
    + branch_analytics()
    + supply_stock()
    + branch_transfer()
    + database_backup()
    + integrity_tools()
    + system_configuration()
    + audit_trail()
    + notifications_activity_settings_logout()
)


OWNER = (
    auth("clinic owner")
    + dashboard()
    + schedule_management()
    + user_accounts_management()
    + staff_emr(history_edit=True, history_print=False, radiograph_upload=True, radiograph_delete=True)
    + material_usage()
    + ai_image_enhancer("AI-Assisted Radiograph Review")
    + branch_analytics()
    + supply_stock()
    + branch_transfer()
    + notifications_activity_settings_logout()
)


BRANCH_MANAGER = (
    auth("branch manager")
    + dashboard()
    + schedule_management()
    + user_accounts_management()
    + staff_emr(history_edit=True, history_print=False)
    + supply_stock()
    + branch_transfer()
    + notifications_activity_settings_logout()
)


DENTIST = (
    auth("dentist")
    + dashboard()
    + [
        simple_view("Schedule Management", "View Schedule List", "Displays the schedule list"),
        simple_view("Schedule Management", "View Schedule Details", "Displays the selected schedule details"),
    ]
    + staff_emr(history_edit=True, history_print=False)
    + material_usage()
    + ai_image_enhancer()
    + notifications_activity_settings_logout()
)


SECRETARY = (
    auth("clinic secretary")
    + dashboard()
    + schedule_management()
    + create_flow("Patient Registration", "patient account", "Register New Patient Account")
    + staff_emr(history_edit=True, history_print=False)
    + notifications_activity_settings_logout()
)


WEB_PATIENT = (
    auth("patient")
    + dashboard()
    + patient_appointments()
    + patient_readonly_emr(mobile=False)
    + patient_ai_engagement(mobile=False)
    + website_booking()
    + notifications_activity_settings_logout()
)


CASES = [MOBILE_PATIENT, ADMIN, OWNER, BRANCH_MANAGER, DENTIST, SECRETARY, WEB_PATIENT]


def unique_cells(row):
    result, seen = [], set()
    for cell in row.cells:
        key = id(cell._tc)
        if key not in seen:
            seen.add(key)
            result.append(cell)
    return result


def replace_cell_text(cell, value):
    paragraph = cell.paragraphs[0]
    for extra in cell.paragraphs[1:]:
        cell._tc.remove(extra._p)
    if paragraph.runs:
        paragraph.runs[0].text = value
        for extra in paragraph.runs[1:]:
            paragraph._p.remove(extra._r)
    else:
        paragraph.add_run(value)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def fill_table(table, prototype, cases):
    for row in list(table.rows[6:]):
        table._tbl.remove(row._tr)
    for module, action, expected in cases:
        table._tbl.append(deepcopy(prototype))
        row = table.rows[-1]
        cells = unique_cells(row)
        if len(cells) != 5:
            raise RuntimeError(f"Expected 5 visible cells, found {len(cells)}")
        replace_cell_text(cells[0], module)
        replace_cell_text(cells[1], action)
        replace_cell_text(cells[2], expected)
        replace_cell_text(cells[3], "")
        replace_cell_text(cells[4], "")
        prevent_row_split(row)


def main():
    document = Document(SOURCE)
    if len(document.tables) != 7:
        raise RuntimeError(f"Expected 7 role tables, found {len(document.tables)}")
    prototype = deepcopy(document.tables[1].rows[6]._tr)
    labels = [
        "Patient / Mobile Application",
        "Admin / Web Application",
        "Owner / Web Application",
        "Branch Manager / Web Application",
        "Dentist / Web Application",
        "Secretary / Web Application",
        "Patient / Web Application",
    ]
    for table, cases, label in zip(document.tables, CASES, labels):
        fill_table(table, prototype, cases)
        print(f"{label}: {len(cases)} cases")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(f"Total: {sum(map(len, CASES))} cases")
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
