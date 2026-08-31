from __future__ import annotations

import json
import re
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
WORKSPACE = ROOT.parents[1]
SOURCE_DATA = ROOT / "normalized_unit_cases.json"
OUTPUT_DIR = WORKSPACE / "output" / "documents"
OUTPUT_DOCX = OUTPUT_DIR / "NGITIFY_Complete_Functional_Testing.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
WHITE = "FFFFFF"
BLACK = "000000"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}

ALL_WEB_USERS = (
    "Administrator, Owner, Branch Manager, Dentist, Secretary, and Patient"
)
ALL_STAFF = "Administrator, Owner, Branch Manager, Dentist, and Secretary"


def clean_text(value: str) -> str:
    text = (value or "").replace("\u00a0", " ")
    text = re.sub(r"(?<=\w)�(?=\w)", "'", text)
    text = text.replace("�", '"')
    text = text.replace("\u2018", "'").replace("\u2019", "'")
    text = text.replace("\u201c", '"').replace("\u201d", '"')
    text = text.replace("â€™", "'").replace("â€œ", '"').replace("â€", '"')
    text = text.replace("â€“", "-").replace("â€”", "-")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", " ", text)
    return text.strip()


def ensure_period(text: str) -> str:
    text = clean_text(text)
    if not text:
        return text
    if text[-1] not in ".?!\"":
        text += "."
    return text


def lower_initial(text: str) -> str:
    if not text:
        return text
    if len(text) >= 2 and text[:2].isupper():
        return text
    return text[0].lower() + text[1:]


def correct_expected_source(text: str, system: str, function_name: str) -> str:
    text = clean_text(text)
    replacements = {
        "Invalid email and password": "Invalid email or password.",
        "Please compete the Captcha before submitting": "Please complete the captcha before submitting.",
        "No matching topic yet": "No Matching Topics Yet",
        "Select another date or save a log for this day": "Select another date or save a Patient Log for this day.",
        "Try another category or clear your search to browse the full Dental Health Education library": "Try another category or clear your search to browse the full library.",
        "Image exceeds 20 MB": "Image exceeds 20 MB limit",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)

    if function_name == "VIEW EMR" and "No radiographs yet" in text:
        if system == "Mobile Application":
            text = text.replace("No radiographs yet", "No radiograph images on file")
        else:
            text = text.replace("No radiographs yet", "No radiograph images yet")
            text = text.replace(
                "Radiographs uploaded by the clinic will appear here for your reference",
                "Dental X-rays shared by your clinic will appear here.",
            )
    if function_name == "VIEW EMR" and system == "Web Application":
        text = text.replace("No odontogram recorded yet", "No odontogram updates have been recorded yet.")
    return clean_text(text)


def normalize_expected(text: str, system: str, function_name: str) -> str:
    text = correct_expected_source(text, system, function_name)
    text = ensure_period(text)

    if text.startswith("The message must display "):
        return ensure_period("The system must display the message " + text[len("The message must display "):])
    if text.startswith("The message must show "):
        return ensure_period("The system must display the message " + text[len("The message must show "):])
    if text.startswith("Success modal must display with title "):
        return ensure_period("The system must display a success modal titled " + text[len("Success modal must display with title "):])
    if text.startswith("A confirmation message must appear"):
        remainder = text[len("A confirmation message must appear"):].lstrip(" ,")
        return ensure_period("The system must display a confirmation prompt" + (", " + remainder if remainder else ""))
    if text.startswith("Only "):
        body = text[5:]
        body = re.sub(r"\s+must display\.?$", "", body, flags=re.I)
        return ensure_period("The system must display only " + lower_initial(body))
    if text.startswith("Table updates to "):
        return ensure_period("The system must update the table to " + text[len("Table updates to "):])
    if text.startswith("The table updates to "):
        return ensure_period("The system must update the table to " + text[len("The table updates to "):])
    if text.startswith("Table shows "):
        return ensure_period("The system must display a table that shows " + text[len("Table shows "):])
    if text.startswith("The table shows "):
        return ensure_period("The system must display a table that shows " + text[len("The table shows "):])
    if text.startswith("NgitiBot must "):
        predicate = text[len("NgitiBot must "):].rstrip(".")
        if predicate.startswith("explain "):
            return ensure_period("The system must display NgitiBot's explanation of " + predicate[len("explain "):])
        if predicate.startswith("display "):
            return ensure_period("The system must display NgitiBot's response stating " + predicate[len("display "):])
        if predicate.startswith("not provide "):
            remainder = predicate[len("not provide "):].replace(" and must recommend ", " and must recommend ")
            return ensure_period("The system must prevent NgitiBot from providing " + remainder)
        return ensure_period("The system must have NgitiBot " + predicate)

    subject_match = re.match(r"^The (.+?) must (.+)$", text, flags=re.I)
    if subject_match:
        subject, predicate = subject_match.groups()
        predicate = predicate.rstrip(".")
        subject_lower = lower_initial(subject)
        if predicate.startswith("display"):
            rest = predicate[len("display"):].strip()
            return ensure_period(f"The system must display {subject_lower}" + (f" {rest}" if rest else ""))
        if predicate.startswith("be disabled"):
            return ensure_period(f"The system must keep {subject_lower} disabled")
        if predicate.startswith("not be available"):
            return ensure_period(f"The system must hide or disable {subject_lower}")
        if predicate.startswith("be available"):
            return ensure_period(f"The system must make {subject_lower} available")
        if predicate.startswith("be updated"):
            return ensure_period(f"The system must update {subject_lower}" + predicate[len("be updated"):])
        if predicate.startswith("be saved"):
            return ensure_period(f"The system must save {subject_lower}" + predicate[len("be saved"):])
        if predicate.startswith("be removed"):
            return ensure_period(f"The system must remove {subject_lower}" + predicate[len("be removed"):])
        if predicate.startswith("be restored"):
            return ensure_period(f"The system must restore {subject_lower}" + predicate[len("be restored"):])
        if predicate.startswith("be downloaded"):
            return ensure_period(f"The system must download {subject_lower}" + predicate[len("be downloaded"):])
        if predicate.startswith("be "):
            rest = predicate[len("be "):].replace(" must ", " ")
            return ensure_period(f"The system must ensure that {subject_lower} is {rest}")
        if predicate.startswith("save"):
            rest = predicate[len("save"):].replace(" must ", " ")
            return ensure_period(f"The system must save {subject_lower}{rest}")
        if predicate.startswith("not provide"):
            rest = predicate.replace("not provide", "does not provide", 1).replace(" must recommend ", " recommends ")
            return ensure_period(f"The system must ensure that {subject_lower} {rest}")
        if predicate.startswith("remain"):
            return ensure_period(f"The system must keep {subject_lower} " + predicate[len("remain"):].strip())
        if predicate.startswith("load"):
            return ensure_period(f"The system must load {subject_lower} " + predicate[len("load"):].strip())
        return ensure_period(f"The system must ensure that {subject_lower} {predicate}")

    if text.startswith("no matching article"):
        remainder = text[len("no matching article"):].strip()
        return ensure_period("The system must display no matching articles and " + lower_initial(remainder))
    quoted_display = re.match(r'^(".+?") must display\.?$', text)
    if quoted_display:
        return ensure_period("The system must display " + quoted_display.group(1))
    if text.startswith('"No Symptoms"'):
        return 'The system must prevent "No Symptoms" and another symptom from remaining selected at the same time.'
    if text.startswith('"'):
        return ensure_period("The system must display " + text)
    if text.startswith("Related "):
        return ensure_period("The system must display " + lower_initial(text))
    if text.startswith("All "):
        if "must be marked as read" in text:
            remainder = text.replace("All unread notifications must be marked as read", "all unread notifications as read")
            remainder = remainder.replace("the inbox must display", "make the inbox display")
            return ensure_period("The system must mark " + lower_initial(remainder))
        return ensure_period("The system must ensure that " + lower_initial(text).replace(" must ", " "))
    return ensure_period("The system must produce the following observable result: " + lower_initial(text))


def polish_expected(text: str) -> str:
    replacements = {
        "and the selected batch must be removed": "and remove the selected batch",
        "The system must display in-progress and duplicate backup must be prevented.": "The system must display the in-progress backup and prevent a duplicate backup.",
        "The system must prevent NgitiBot from providing a diagnosis and must recommend consulting the dentist.": "The system must prevent NgitiBot from providing a diagnosis and display a recommendation to consult the dentist.",
        'The system must display recommendation "Insufficient Data" and the message must display': 'The system must display the "Insufficient Data" recommendation and show the message',
        'The system must display recommendation "Insufficient Data" and an unsupported visit window must not be created.': 'The system must display the "Insufficient Data" recommendation and must not create an unsupported visit window.',
        "The system must ensure that patient not be able to book another appointment.": "The system must prevent the patient from booking another appointment.",
        "The system must ensure that appointment not be cancelled.": "The system must leave the appointment unchanged and must not cancel it.",
        'The system must display selected Daily Oral Health Log and the selected date must display "Patient Log"': 'The system must display the selected Daily Oral Health Log and label the selected date "Patient Log"',
        "The system must display related Gum Care and Flossing Dental Health Education must display.": "The system must display the related Gum Care and Flossing Dental Health Education topics.",
        'The system must display no matching articles and the message must display': 'The system must display no matching articles and show the message',
        "The system must display patient's odontogram and the patient must not be able to edit it.": "The system must display the patient's odontogram in read-only form.",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return ensure_period(text)


def to_success_phrase(text: str, function_name: str) -> str:
    text = ensure_period(text)
    prefix = "The system must "
    rest = text[len(prefix):].rstrip(".") if text.startswith(prefix) else text.rstrip(".")
    lower_rest = rest.lower()
    if lower_rest.startswith("display "):
        return ensure_period("Must be able to " + rest)
    if lower_rest.startswith(("update ", "save ", "download ", "restore ", "end ")):
        phrase = "Must " + rest
        if "successfully" not in phrase.lower():
            phrase += " successfully"
        return ensure_period(phrase)
    if lower_rest.startswith("ensure that "):
        clause = rest[len("ensure that "):]
        return ensure_period(f"Must successfully complete the {function_name.title()} flow so that {clause}")
    return ensure_period("Must successfully " + rest)


def role_for_web(index: int) -> str:
    if index == 1:
        return "Public User / Patient"
    if index == 2:
        return "Patient"
    if 3 <= index <= 11:
        return ALL_WEB_USERS
    if 12 <= index <= 15:
        return "Administrator, Owner, and Branch Manager"
    if index == 16:
        return ALL_STAFF
    if index in (17, 18):
        return "Administrator and Secretary"
    if index in (19, 20):
        return "Administrator, Owner, and Branch Manager"
    if 21 <= index <= 23:
        return "Administrator and Owner"
    if index == 24:
        return "Administrator, Owner, and Branch Manager"
    if 25 <= index <= 28:
        return ALL_STAFF
    if index == 29:
        return "Administrator, Owner, and Dentist"
    if index in (30, 31, 32):
        return "Owner and Dentist"
    if index == 33:
        return "Patient"
    if index == 34:
        return "Administrator, Owner, Branch Manager, and Dentist (read-only)"
    if 35 <= index <= 38:
        return "Administrator, Owner, and Branch Manager"
    if index == 39:
        return "Owner and Dentist"
    if 40 <= index <= 49:
        return "Administrator"
    if index in (50, 51):
        return ALL_WEB_USERS
    if index in (52, 53):
        return "Administrator"
    return "Patient"


def case_record(source: dict, system: str, access: str, source_tag: str) -> dict:
    function_name = clean_text(source["function"]).upper()
    scenarios = []
    for i, scenario in enumerate(source.get("scenarios", []), 1):
        number = clean_text(scenario.get("number")) or str(i)
        scenarios.append({
            "number": number,
            "condition": ensure_period(clean_text(scenario.get("input"))),
            "expected": normalize_expected(scenario.get("expected", ""), system, function_name),
        })
    return {
        "module": clean_text(source["module"]).upper(),
        "function": function_name,
        "access": access,
        "system": system,
        "preconditions": ensure_period(
            "The listed user must satisfy this condition: " + lower_initial(clean_text(source.get("preconditions")))
        ),
        "action": ensure_period(
            f"The listed user performs the {function_name.title()} function using the condition stated in each test scenario"
        ),
        "verification": (
            "Perform each test scenario and verify that the system produces the corresponding observable output."
        ),
        "scenarios": scenarios,
        "source_tag": source_tag,
    }


def manual_case(module: str, function: str, access: str, system: str, precondition: str, scenarios: list[tuple[str, str]]) -> dict:
    return {
        "module": module,
        "function": function,
        "access": access,
        "system": system,
        "preconditions": ensure_period(precondition),
        "action": ensure_period(
            f"The listed user performs the {function.title()} function using the condition stated in each test scenario"
        ),
        "verification": "Perform each test scenario and verify that the system produces the corresponding observable output.",
        "scenarios": [
            {"number": str(i), "condition": ensure_period(condition), "expected": ensure_period(expected)}
            for i, (condition, expected) in enumerate(scenarios, 1)
        ],
        "source_tag": "integration",
    }


def build_cases() -> list[dict]:
    source = json.loads(SOURCE_DATA.read_text(encoding="utf-8"))
    web = [
        case_record(case, "Web Application", role_for_web(i), f"unit-web-{i}")
        for i, case in enumerate(source["web"], 1)
    ]
    mobile = [
        case_record(case, "Mobile Application", "Patient", f"unit-mobile-{i}")
        for i, case in enumerate(source["mobile"], 1)
    ]

    # The role-specific dashboard result replaces the generic successful-login output.
    login_case = next(c for c in web if c["function"] == "LOGIN")
    remaining_login_scenarios = login_case["scenarios"][1:]
    login_case["scenarios"] = [
        {"number": "1", "condition": "Correct Administrator email and password.", "expected": "The system must display the Administrator dashboard."},
        {"number": "2", "condition": "Correct Owner email and password.", "expected": "The system must display the Owner dashboard."},
        {"number": "3", "condition": "Correct Branch Manager email and password.", "expected": "The system must display the Branch Manager dashboard."},
        {"number": "4", "condition": "Correct Dentist email and password.", "expected": "The system must display the Dentist dashboard."},
        {"number": "5", "condition": "Correct Secretary email and password.", "expected": "The system must display the Secretary dashboard."},
        {"number": "6", "condition": "Correct Patient email and password.", "expected": "The system must display the Patient dashboard."},
    ] + [
        {**s, "number": str(i)} for i, s in enumerate(remaining_login_scenarios, 7)
    ]

    # Align role-sensitive outputs and visible modal composition to the implementation.
    for case in web + mobile:
        if case["function"] == "SET-UP PASSWORD":
            for scenario in case["scenarios"]:
                if "Activation Unavailable" in scenario["expected"]:
                    scenario["expected"] = 'The system must display the "Activation Unavailable" page with the message "Invalid or expired activation link."'
        if case["function"] == "EDIT PROFILE":
            for scenario in case["scenarios"]:
                if scenario["number"] == "1":
                    if case["system"] == "Mobile Application":
                        scenario["expected"] = 'The system must display the "Profile Updated" modal with the message "Your profile information has been successfully saved."'
                    else:
                        scenario["expected"] = 'The system must display the role-appropriate profile update confirmation and show the saved profile information.'
                if "recommended 2mb size" in scenario["expected"].lower():
                    scenario["expected"] = 'The system must display the "Upload Failed" modal with the message "File exceeds the recommended 2MB size."'
        if case["function"] == "CHANGE EMAIL":
            for scenario in case["scenarios"]:
                if "Verification email sent" in scenario["expected"]:
                    scenario["expected"] = 'The system must display the "Request Link Sent" confirmation, state that a verification email was sent, and log the user out for reactivation.'
        if case["function"] == "CHANGE PASSWORD":
            for scenario in case["scenarios"]:
                if scenario["number"] == "1":
                    scenario["expected"] = 'The system must display the password-change success confirmation and log the user out for security.'
        if case["function"] == "ADD NEW STAFF":
            for scenario in case["scenarios"]:
                if scenario["number"] == "1":
                    scenario["expected"] = 'The system must display the selected role-specific success confirmation and state that the activation email was sent.'
        if case["function"] == "EDIT STAFF INFO":
            for scenario in case["scenarios"]:
                if scenario["number"] == "1":
                    scenario["expected"] = 'The system must display the selected role-specific profile update success confirmation.'
        if case["function"] == "WEBSITE CONTENT AND MEDIA":
            for scenario in case["scenarios"]:
                if scenario["number"] == "2":
                    scenario["expected"] = 'The system must preview the uploaded Website Logo with Text and indicate that Save Changes publishes it to the website.'
                if scenario["number"] == "3":
                    scenario["expected"] = 'The system must display the "Image exceeds 20 MB limit" error and identify the 20 MB upload limit.'
        if case["function"] == "SAFE AUTO-FIX":
            for scenario in case["scenarios"]:
                if scenario["number"] == "4":
                    scenario["expected"] = 'The system must display the successful auto-fix result and rerun the affected integrity check.'

    for case in web + mobile:
        for scenario in case["scenarios"]:
            scenario["expected"] = polish_expected(scenario["expected"])

    # Add integration-covered outputs that do not have their own unit-test component.
    web_extras = [
        manual_case(
            "DASHBOARD MODULE", "DASHBOARD", ALL_WEB_USERS, "Web Application",
            "The listed user must be authenticated and authorized for the web application",
            [
                ("The Administrator opens Dashboard", "The system must display the Administrator dashboard summary and administrative navigation."),
                ("The Owner opens Dashboard", "The system must display the Owner dashboard summary and owner navigation."),
                ("The Branch Manager opens Dashboard", "The system must display the assigned-branch dashboard summary and branch management navigation."),
                ("The Dentist opens Dashboard", "The system must display the Dentist dashboard with the assigned schedule and patient summary."),
                ("The Secretary opens Dashboard", "The system must display the Secretary dashboard with appointment and patient-workflow summaries."),
                ("The Patient opens Dashboard", "The system must display the Patient dashboard with appointment and oral-health information."),
            ],
        ),
        manual_case(
            "ACCOUNT SETTINGS MODULE", "VIEW PROFILE", ALL_WEB_USERS, "Web Application",
            "The listed user must be authenticated and the My Profile page must be accessible",
            [("The user opens My Profile", "The system must display the authenticated user's available profile information and role details.")],
        ),
        manual_case(
            "NOTIFICATIONS MODULE", "STAFF NOTIFICATION INBOX", ALL_STAFF, "Web Application",
            "The listed staff user must be authenticated and the Notifications page must be accessible",
            [
                ("The staff user has notifications", "The system must display the staff notification list and unread indicators."),
                ("The staff user opens an unread notification", "The system must display the notification details and mark the notification as read."),
            ],
        ),
        manual_case(
            "LOGIN MODULE", "LOGOUT", ALL_WEB_USERS, "Web Application",
            "The listed user must be authenticated",
            [("The user confirms Logout", "The system must end the authenticated session and display the Login page.")],
        ),
    ]
    mobile_extras = [
        manual_case(
            "DASHBOARD MODULE", "PATIENT DASHBOARD", "Patient", "Mobile Application",
            "The patient must be authenticated in the mobile application",
            [("The patient opens Home", "The system must display the Patient dashboard with the next appointment, recommended visit information, and navigation shortcuts.")],
        ),
        manual_case(
            "ACCOUNT SETTINGS MODULE", "VIEW PROFILE", "Patient", "Mobile Application",
            "The patient must be authenticated and the Profile tab must be accessible",
            [("The patient opens Profile", "The system must display the patient's available profile information and account actions.")],
        ),
        manual_case(
            "ACTIVITY LOGS MODULE", "VIEW ACTIVITY LOGS", "Patient", "Mobile Application",
            "The patient must be authenticated and Activity Logs must be accessible from Profile",
            [
                ("The patient has recorded account activity", "The system must display the patient's activity log entries."),
                ("The patient has no recorded activity", "The system must display the activity-log empty state."),
            ],
        ),
        manual_case(
            "LOGIN MODULE", "LOGOUT", "Patient", "Mobile Application",
            "The patient must be authenticated",
            [("The patient confirms Logout", "The system must end the authenticated session and display the mobile Login screen.")],
        ),
    ]

    # Insert web integration extras after shared authentication/account cases and mobile extras after mobile login/account cases.
    cases = web[:11] + web_extras + web[11:] + mobile[:8] + mobile_extras + mobile[8:]

    manual_success = {
        "PATIENT BRANCH TRANSFER": "Must update the patient's branch assignment to the selected target branch successfully.",
        "BRANCH STATUS": "Must deactivate the selected branch successfully and display its updated status.",
        "TREATMENT LOGS": "Must add the treatment log successfully and display it in the patient's EMR.",
        "INTERACTIVE DIGITAL ODONTOGRAM": "Must be able to display the patient's current tooth data in the digital odontogram.",
        "AI-ASSISTED RADIOGRAPH REVIEW": "Must upload the radiograph successfully and display it in the patient record.",
        "ADD NEW INVENTORY ITEM": "Must add the new inventory item successfully and display it in the inventory tracker.",
        "ADD INVENTORY STOCK": "Must add the stock batch successfully and display the updated available quantity.",
        "MATERIAL USAGE LOG": "Must save the material usage log successfully and update the corresponding inventory quantities.",
        "BACKUP STATUS AND HISTORY": "Must be able to display existing backup records with their status, file availability, and verification result.",
        "CREATE DATABASE BACKUP": "Must create the database backup successfully and display it in the backup history.",
        "SAFE AUTO-FIX": "Must update the eligible records successfully and rerun the affected integrity check.",
        "RECOMMENDED VISIT WINDOW": "Must be able to display the dentist-recommended visit window as the primary planned-care recommendation.",
    }
    for case in cases:
        successful = case["scenarios"][0]
        successful["number"] = "1"
        successful["expected"] = manual_success.get(
            case["function"],
            to_success_phrase(successful["expected"], case["function"]),
        )

        if case["function"] == "LOGIN":
            if case["system"] == "Web Application":
                successful["condition"] = "Enter the correct email and password for the listed role."
                successful["expected"] = "Must successfully log in and display the role-appropriate dashboard."
            else:
                successful["condition"] = "Enter the correct patient email and password."
                successful["expected"] = "Must successfully log in and display the Patient dashboard."
        elif case["function"] == "DASHBOARD":
            successful["condition"] = "The listed user opens Dashboard."
            successful["expected"] = "Must be able to display the role-appropriate dashboard summary and navigation."
        elif case["function"] == "VIEW EMR" and case["system"] == "Mobile Application":
            successful["expected"] = "Must be able to display the patient's Medical and Dental History."

        case["scenarios"] = [successful]
        case["verification"] = "Complete the successful flow and verify that the system displays the stated output."
    return cases


def set_run_font(run, size=9, bold=False, color=BLACK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGINS.items():
        tag = "left" if side == "start" else "right" if side == "end" else side
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_text(cell, text, *, bold=False, color=BLACK, size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8, color=MID_GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    run = header_p.add_run("NGITIFY  |  FUNCTIONAL TEST SPECIFICATION")
    set_run_font(run, size=8, bold=True, color=MID_GRAY)
    footer_p = section.footer.paragraphs[0]
    add_page_number(footer_p)


def add_cover(doc: Document, cases: list[dict]):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("FUNCTIONAL TESTING DOCUMENT")
    set_run_font(r, size=24, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("NGITIFY: A Dental Information Management System with AI-Driven Patient Engagement for Dentime Dental Clinic")
    set_run_font(r, size=13, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(
        "Scope: observable functional outputs for the public website, role-based web application, and patient mobile application."
    )
    set_run_font(r, size=10.5, color=BLACK)

    heading = doc.add_paragraph(style="Heading 1")
    heading.add_run("Coverage at a Glance")

    roles = ["Administrator", "Owner", "Branch Manager", "Dentist", "Secretary", "Patient", "Public User"]
    coverage_rows = []
    for role in roles:
        role_cases = [c for c in cases if role in c["access"]]
        systems = sorted({c["system"] for c in role_cases})
        coverage_rows.append((role, ", ".join(systems) if systems else "-", str(len(role_cases))))

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [3000, 4500, 1860])
    for i, label in enumerate(("User / Access Type", "System Coverage", "Functional Cases")):
        set_cell_text(table.rows[0].cells[i], label, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[i], BLUE)
    set_repeat_table_header(table.rows[0])
    for role, systems, count in coverage_rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], role, size=9)
        set_cell_text(cells[1], systems, size=9)
        set_cell_text(cells[2], count, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_geometry(table, [3000, 4500, 1860])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run(f"Total functional cases: {len(cases)}")
    set_run_font(r, size=10, bold=True, color=DARK_BLUE)
    doc.add_page_break()


def add_case(doc: Document, case: dict, case_id: str, case_index: int, last: bool):
    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label.paragraph_format.space_after = Pt(6)
    run = label.add_run(f"FUNCTIONALITY TEST: {case['access'].upper()} - {case['system'].upper()}")
    set_run_font(run, size=9.5, bold=True, color=DARK_BLUE)

    meta = doc.add_table(rows=8, cols=2)
    meta.style = "Table Grid"
    widths = [2400, 6960]
    set_table_geometry(meta, widths)

    set_cell_text(meta.rows[0].cells[0], "FUNCTIONAL TEST DOCUMENT", bold=True, color=WHITE, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(meta.rows[0].cells[0], BLUE)
    set_cell_text(meta.rows[0].cells[1], f"Module Name: {case['module']}", bold=True, color=DARK_BLUE, size=8.5)
    shade_cell(meta.rows[0].cells[1], LIGHT_BLUE)

    detail_rows = (
        ("Function Name", case["function"]),
        ("User / Access Type", case["access"]),
        ("Type of System", case["system"]),
        ("Case ID", case_id),
        ("Pre-Conditions", case["preconditions"]),
        ("Action Description", case["action"]),
        ("Verification Steps", case["verification"]),
    )
    for row_index, (label_text, value) in enumerate(detail_rows, 1):
        set_cell_text(meta.rows[row_index].cells[0], label_text, bold=True, color=DARK_BLUE, size=8.5)
        shade_cell(meta.rows[row_index].cells[0], LIGHT_GRAY)
        set_cell_text(meta.rows[row_index].cells[1], value, size=8.5)

    scenario_header = meta.add_row().cells
    set_cell_text(scenario_header[0], "Test Scenario", bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(scenario_header[1], "Expected Results", bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(scenario_header[0], BLUE)
    shade_cell(scenario_header[1], BLUE)
    scenario = case["scenarios"][0]
    cells = meta.add_row().cells
    set_cell_text(cells[0], f"{scenario['number']}. {scenario['condition']}", size=8.5)
    set_cell_text(cells[1], scenario["expected"], size=8.5)
    set_table_geometry(meta, widths)

    if not last:
        if case_index % 2 == 0:
            doc.add_page_break()
        else:
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(8)


def audit_cases(cases: list[dict]):
    assert cases, "No functional cases generated"
    required_roles = ["Administrator", "Owner", "Branch Manager", "Dentist", "Secretary", "Patient"]
    for role in required_roles:
        assert any(role in c["access"] for c in cases), f"Missing role coverage: {role}"
    for case in cases:
        for key in ("module", "function", "access", "system", "preconditions", "action", "verification"):
            assert case[key], f"Missing {key}: {case}"
        assert case["scenarios"], f"No scenarios: {case['function']}"
        for scenario in case["scenarios"]:
            assert scenario["condition"] and scenario["expected"]
            assert scenario["expected"].startswith("Must "), scenario["expected"]
            assert (
                scenario["expected"].startswith("Must be able to ")
                or scenario["expected"].startswith("Must successfully ")
                or " successfully" in scenario["expected"]
            ), scenario["expected"]
            forbidden = ("Actual Results", "Remarks", "Date Tested", "Test Cycle No.")
            assert not any(word in scenario["expected"] for word in forbidden)


def main():
    cases = build_cases()
    audit_cases(cases)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    (ROOT / "final_cases.json").write_text(json.dumps(cases, indent=2, ensure_ascii=False), encoding="utf-8")

    doc = Document()
    configure_document(doc)
    add_cover(doc, cases)
    for index, case in enumerate(cases, 1):
        prefix = "WEB" if case["system"] == "Web Application" else "MOB"
        add_case(doc, case, f"FT-{prefix}-{index:03d}", index, index == len(cases))

    props = doc.core_properties
    props.title = "NGITIFY Complete Functional Testing Document"
    props.subject = "Role-based functional output verification for web and mobile applications"
    props.author = "NGITIFY Project Team"
    props.keywords = "functional testing, NGITIFY, Dentime, web, mobile"
    doc.save(OUTPUT_DOCX)

    manifest = {
        "output": str(OUTPUT_DOCX),
        "case_count": len(cases),
        "scenario_count": sum(len(c["scenarios"]) for c in cases),
        "systems": dict(Counter(c["system"] for c in cases)),
        "modules": sorted({c["module"] for c in cases}),
        "excluded_fields": ["Test Cycle No.", "Date Tested", "Actual Results", "Remarks"],
    }
    (ROOT / "final_manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    print(json.dumps(manifest, indent=2))


if __name__ == "__main__":
    main()
