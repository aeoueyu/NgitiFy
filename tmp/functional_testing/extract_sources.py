from __future__ import annotations

import json
from pathlib import Path
from typing import Any

import fitz
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent
DOWNLOADS = Path(r"C:\Users\Administrator\Downloads")

DOCX_FILES = [
    DOWNLOADS / "integration testing.docx",
    DOWNLOADS / "unit testing web.docx",
    DOWNLOADS / "unit testing mobile.docx",
    DOWNLOADS / "functionality testing.docx",
]
PDF_FILE = DOWNLOADS / "functional testing sample (1).pdf"


def iter_block_items(parent):
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def cell_text(cell) -> str:
    return "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def extract_docx(path: Path) -> dict[str, Any]:
    doc = Document(path)
    blocks = []
    for idx, item in enumerate(iter_block_items(doc)):
        if isinstance(item, Paragraph):
            if item.text.strip() or item._p.xpath('.//w:drawing | .//w:pict'):
                blocks.append({
                    "index": idx,
                    "type": "paragraph",
                    "style": item.style.name if item.style else None,
                    "text": item.text,
                    "has_image": bool(item._p.xpath('.//w:drawing | .//w:pict')),
                })
        else:
            blocks.append({
                "index": idx,
                "type": "table",
                "rows": [[cell_text(c) for c in row.cells] for row in item.rows],
            })

    sections = []
    for s in doc.sections:
        sections.append({
            "page_width_in": s.page_width.inches,
            "page_height_in": s.page_height.inches,
            "top_margin_in": s.top_margin.inches,
            "bottom_margin_in": s.bottom_margin.inches,
            "left_margin_in": s.left_margin.inches,
            "right_margin_in": s.right_margin.inches,
            "header_distance_in": s.header_distance.inches,
            "footer_distance_in": s.footer_distance.inches,
            "header": [p.text for p in s.header.paragraphs],
            "footer": [p.text for p in s.footer.paragraphs],
        })

    return {
        "path": str(path),
        "sections": sections,
        "inline_shapes": len(doc.inline_shapes),
        "blocks": blocks,
    }


def extract_pdf(path: Path) -> dict[str, Any]:
    pdf = fitz.open(path)
    pages = []
    render_dir = ROOT / "renders" / "functional_sample"
    render_dir.mkdir(parents=True, exist_ok=True)
    for page_index, page in enumerate(pdf):
        pages.append({
            "page": page_index + 1,
            "width": page.rect.width,
            "height": page.rect.height,
            "text": page.get_text("text"),
        })
        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
        pix.save(render_dir / f"page-{page_index + 1}.png")
    return {"path": str(path), "pages": pages}


def main() -> None:
    ROOT.mkdir(parents=True, exist_ok=True)
    for path in DOCX_FILES:
        result = extract_docx(path)
        (ROOT / f"{path.stem}.json").write_text(
            json.dumps(result, indent=2, ensure_ascii=False), encoding="utf-8"
        )
    (ROOT / "functional testing sample (1).json").write_text(
        json.dumps(extract_pdf(PDF_FILE), indent=2, ensure_ascii=False), encoding="utf-8"
    )


if __name__ == "__main__":
    main()
