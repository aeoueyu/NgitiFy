from pathlib import Path
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SOURCE = Path(r"C:\Users\Administrator\Documents\4.4 DESCRIPTION OF THE SYSTEM.docx")
OUTPUT = Path(r"C:\Users\Administrator\Documents\4.4 DESCRIPTION OF THE SYSTEM - Captioned.docx")

captions = [
"Login (Web)",
"Forgot Password (Web)",
"Email Verification - Forgot Password (Web)",
"Set New Password - Forgot Password (Web)",
"Administrator Dashboard 1 (Web)",
"Administrator Dashboard 2 (Web)",
"Administrator Dashboard 3 (Web)",
"Administrator Dashboard 4 (Web)",
"NgitifyBot (Web)",
"Conversation History (Web)",
"Schedule Management (Web)",
"Create Schedule Entry (Web)",
"My Appointments (Web)",
"Book Appointment - Branch Selection (Web)",
"Book Appointment - Date and Procedure Selection (Web)",
"Manage Patients (Web)",
"Add New Patient - Identity (Web)",
"Electronic Medical Record - Patient Profile Overview (Web)",
"Electronic Medical Record - Patient Profile Details (Web)",
"Electronic Medical Record - Medical and Dental History 1 (Web)",
"Electronic Medical Record - Medical and Dental History 2 (Web)",
"Electronic Medical Record - Medical and Dental History 3 (Web)",
"Electronic Medical Record - Treatment History 1 (Web)",
"Electronic Medical Record - Treatment History 2 (Web)",
"Add Treatment Log (Web)",
"Add Treatment Notes (Web)",
"Clinical 2D Odontogram (Web)",
"Update Odontogram Entry 1 (Web)",
"Update Odontogram Entry 2 (Web)",
"Odontogram Legend (Web)",
"Odontogram History (Web)",
"Radiograph History (Web)",
"AI-Assisted Radiograph Review (Web)",
"Upload Radiograph (Web)",
"Radiograph Viewer (Web)",
"AI Radiograph Analysis - Quality (Web)",
"AI Radiograph Analysis - Findings (Web)",
"AI Radiograph Analysis - Summary (Web)",
"Dentist-Approved Radiograph Summary (Web)",
"Inventory Tracker (Web)",
"Inventory Item Details (Web)",
"Add New Inventory Item (Web)",
"Add Supply or Stock (Web)",
"Edit Inventory Item (Web)",
"Material Usage Log 1 (Web)",
"Material Usage Log 2 (Web)",
"Log New Material Usage (Web)",
"Manage Secretaries (Web)",
"Add New Secretary (Web)",
"Edit Secretary Profile (Web)",
"Manage Dentists (Web)",
"Add New Dentist (Web)",
"Edit Dentist Profile (Web)",
"Manage Branch Managers (Web)",
"Add New Branch Manager (Web)",
"Edit Branch Manager (Web)",
"Manage Owners (Web)",
"Add New Owner (Web)",
"Edit Owner (Web)",
"Notifications (Web)",
"Notification Details (Web)",
"Branch Management (Web)",
"Add New Branch (Web)",
"Branch Details (Web)",
"Edit Branch (Web)",
"My Activity Logs (Web)",
"Activity Log Details (Web)",
"System Audit Logs (Web)",
"Audit Log Details (Web)",
"System Configuration - Clinic Information (Web)",
"System Configuration - Appointment Settings (Web)",
"System Configuration - Email Templates (Web)",
"System Configuration - Feature Toggles (Web)",
"Website Content - Branding (Web)",
"Website Content - Logo Settings (Web)",
"Website Content - Home Page 1 (Web)",
"Website Content - Home Page 2 (Web)",
"Website Content - Home Page 3 (Web)",
"Website Content - Home Page 4 (Web)",
"Website Content - About and Locations 1 (Web)",
"Website Content - About and Locations 2 (Web)",
"Website Content - About and Locations 3 (Web)",
"Website Content - Services 1 (Web)",
"Website Content - Services 2 (Web)",
"Website Content - Services 3 (Web)",
"Website Content - Contact 1 (Web)",
"Website Content - Contact 2 (Web)",
"Website Content - Contact 3 (Web)",
"Website Content - Appointment 1 (Web)",
"Website Content - Appointment 2 (Web)",
"Website Content - Appointment 3 (Web)",
"Archive Review (Web)",
"Database Backup (Web)",
"Database Backup History (Web)",
"Integrity Tools 1 (Web)",
"Integrity Tools 2 (Web)",
"Integrity Tools 3 (Web)",
"Integrity Tools - Stale Unverified Accounts (Web)",
"Integrity Tools - Assigned Dentist Mismatches (Web)",
"My Profile 1 (Web)",
"My Profile 2 (Web)",
"Change Email Address (Web)",
"Account Security Settings (Web)",
"Notification Settings (Web)",
"My Appointments - Empty State (Web)",
"My Appointments - Appointment Entry (Web)",
"Appointment Details (Web)",
"Reschedule Appointment (Web)",
"Cancel Appointment (Web)",
"Appointment History (Web)",
"Cancelled Appointment Details (Web)",
"Book Appointment - Branch and Procedure (Web)",
"Book Appointment - Select Date (Web)",
"Book Appointment - Select Time (Web)",
"Book Appointment - Procedure and Notes (Web)",
"Book Appointment - Review and Confirm (Web)",
"Privacy Policy Summary (Web)",
"Booking Submitted (Web)",
"Active Appointment Conflict (Web)",
"Oral Health Management Dashboard (Web)",
"Oral Health Log Details (Web)",
"Edit Oral Health Log (Web)",
"Symptom Context (Web)",
"Oral Health Log - Empty State (Web)",
"Oral Health Calendar (Web)",
"Oral Health Trends (Web)",
"Oral Health Risk Factors (Web)",
"Dental Health Education Library (Web)",
"Dental Health Education Article (Web)",
"Login (Mobile)",
"Set New Password (Mobile)",
"Forgot Password (Mobile)",
"Email Verification - Forgot Password (Mobile)",
"Home Dashboard (Mobile)",
"Care Tools (Mobile)",
"NgitifyBot (Mobile)",
"Conversation History (Mobile)",
"Conversation Options (Mobile)",
"Archived Conversations (Mobile)",
"Oral Health Management (Mobile)",
"Oral Health Calendar (Mobile)",
"Daily Oral Health Log (Mobile)",
"Log Oral Health - Symptoms (Mobile)",
"Symptom Details (Mobile)",
"Log Oral Health - Care Habits (Mobile)",
"Log Oral Health - Risk Factors (Mobile)",
"Review Oral Health Log (Mobile)",
"Oral Health Trends - 7 Days (Mobile)",
"Oral Health Trends - 30 Days (Mobile)",
"Oral Health Trends - Care Habits (Mobile)",
"Oral Health Trends - Symptoms (Mobile)",
"Oral Health Trends - Risk Factors (Mobile)",
"Recommended Visit Window (Mobile)",
"Oral Health Calendar - Selected Date (Mobile)",
"Suggested Next Action (Mobile)",
"Current Oral Health Factors (Mobile)",
"Edit Oral Health Factors (Mobile)",
"Dental Health Education (Mobile)",
"Dental Health Education Library (Mobile)",
"Dental Health Education Article (Mobile)",
"Appointments Dashboard (Mobile)",
"Past Visits (Mobile)",
"Reschedule Appointment (Mobile)",
"Cancel Appointment (Mobile)",
"Electronic Medical Record - Medical and Dental History (Mobile)",
"Electronic Medical Record - Treatment History (Mobile)",
"Electronic Medical Record - Odontogram (Mobile)",
"Odontogram Tooth Details 1 (Mobile)",
"Odontogram Tooth Details 2 (Mobile)",
"Odontogram History (Mobile)",
"Radiograph History (Mobile)",
"Radiograph Viewer (Mobile)",
"Profile (Mobile)",
"Edit Profile (Mobile)",
"My Details (Mobile)",
"Activity Logs (Mobile)",
"Settings (Mobile)",
"Change Password - Verify Identity (Mobile)",
"Change Password - Set New Password (Mobile)",
"Change Email Address (Mobile)",
"Notification Settings (Mobile)",
]

descriptions = {
1: "Figure 1 shows the web login page, where an authorized user enters an email address and password to access the system. It also provides a Forgot Password link for account recovery.",
2: "Figure 2 shows the web Forgot Password page, where the user enters a registered email address to request a password-reset code.",
3: "Figure 3 shows the email verification step of account recovery. The user enters the verification code sent to the registered email address before proceeding.",
4: "Figure 4 shows the final password-reset step, where the user creates and confirms a new password before returning to the login page.",
8: "Figures 5, 6, 7, and 8 show the administrator dashboard and its primary views. The dashboard summarizes patient, staff, and dental-record indicators; displays visit trends, treatment distribution, clinic events, recent activity, and appointments; and provides access to the sidebar and NgitifyBot assistant.",
10: "Figures 9 and 10 show NgitifyBot and its conversation-history panel. Users can ask role-permitted questions, start a new conversation, reopen recent conversations, and manage saved chats.",
12: "Figures 11 and 12 show schedule management and the Create Schedule Entry form. Authorized users can review clinic schedules and create entries by selecting a source, branch, patient, dentist, date, procedure, time, status, and optional notes.",
15: "Figures 13, 14, and 15 show the web appointment area and booking workflow. A patient can review appointment counts and statuses, select a branch, view its booking rules, and choose an available date and procedure before continuing.",
17: "Figures 16 and 17 show patient management and the first step of adding a patient. Authorized staff can search and filter patient records, perform record actions, export data, and enter a new patient's identity information.",
22: "Figures 18 through 22 show the patient profile and the medical and dental history sections of the electronic medical record. These views consolidate demographic, contact, consent, allergy, medical-condition, medication, and related clinical information for authorized review.",
24: "Figures 23 and 24 show treatment-history views in the electronic medical record. Users can review treatment dates, procedures, dentists, clinical details, and other recorded information for each visit.",
26: "Figures 25 and 26 show the forms for adding a treatment log and its clinical notes. Staff can record the procedure, category, branch, charges, next appointment, and follow-up instructions.",
31: "Figures 27 through 31 show the clinical odontogram, tooth-entry workflow, legend, and odontogram history. Dental staff can select a tooth and surface, record its stage and condition, add clinical notes, and review previously recorded odontogram changes.",
39: "Figures 32 through 39 show radiograph history, uploading, viewing, AI-assisted review, and the dentist-approved summary. The system stores radiograph details, presents quality and finding suggestions for professional verification, and preserves the approved interpretation in the patient's record.",
44: "Figures 40 through 44 show the inventory tracker and item-maintenance screens. Authorized users can monitor stock levels and expiry status, inspect item history, create an inventory item, add stock, and update item details.",
47: "Figures 45, 46, and 47 show the material usage log and the form for recording materials consumed during a completed appointment. Recorded quantities are deducted from inventory to keep stock balances current.",
50: "Figures 48, 49, and 50 show secretary account management. Administrators can search, filter, add, and edit secretary profiles and assign each secretary to a clinic branch.",
53: "Figures 51, 52, and 53 show dentist account management. Administrators can create and update dentist profiles, including professional credentials, specialization, contact information, and branch assignment.",
56: "Figures 54, 55, and 56 show branch-manager account management. Administrators can add or edit a manager's personal information and assigned branch.",
59: "Figures 57, 58, and 59 show owner account management. Administrators can add and edit owners, configure their clinic access, and maintain professional and contact information.",
61: "Figures 60 and 61 show the notification list and an expanded notification. Users can search and filter notices, mark them as read or unread, and review the type, status, timestamp, message, and reference details.",
65: "Figures 62 through 65 show branch management. Administrators can review all branches, add a branch and its address, inspect branch statistics and activity, and update branch information or manager assignment.",
69: "Figures 66 through 69 show personal activity logs and system-wide audit logs with their detail views. These records identify the date, time, user, category, action, and recorded message for accountability and security review.",
73: "Figures 70 through 73 show system configuration for clinic information, appointment rules, email templates, and feature toggles. Administrators can maintain clinic details, booking limits and procedures, system messages, and module availability.",
75: "Figures 74 and 75 show website branding settings. Administrators can update public-facing names and social links as well as the website logo and standalone logo icon.",
79: "Figures 76 through 79 show the configurable content of the public home page. Administrators can manage hero and comfort images, headings, call-to-action labels, service copy, journey content, highlights, and related home-page text.",
82: "Figures 80, 81, and 82 show the configurable About and Locations content. These settings control the page imagery, location highlights, titles, descriptions, calls to action, and supporting clinic information.",
85: "Figures 83, 84, and 85 show the configurable Services page. Administrators can manage the page image and descriptive content and add or remove individual service cards with their category, image, procedures, title, and description.",
88: "Figures 86, 87, and 88 show the configurable Contact page. The settings manage contact imagery, social and map images, call-to-action labels, contact titles, form labels, and the page description.",
91: "Figures 89, 90, and 91 show the configurable Appointment page. Administrators can set appointment imagery, labels, guidance text, form wording, branch choices, and guest-booking instructions.",
92: "Figure 92 shows the archive review page, where authorized users can search archived records, filter them by role or review status, inspect the archive reason, and restore eligible accounts.",
94: "Figures 93 and 94 show database backup status, automated scheduling, retention settings, verification information, and backup history. Administrators can create, verify, restore, and monitor protected system backups.",
99: "Figures 95 through 99 show system integrity checks and detailed issue lists. The tools scan appointments, branches, inventory, security records, and patient assignments, then allow authorized users to review warnings and apply supported fixes.",
101: "Figures 100 and 101 show the web profile page, which presents the user's personal, demographic, home-address, and account information and provides access to profile editing.",
102: "Figure 102 shows the Change Email Address form. The user enters a new email address and current password, verifies the request, and receives an activation link for the new address.",
104: "Figures 103 and 104 show account security and notification settings. Users can update their password or security questions and choose which account, appointment, and dental-care notifications they receive.",
111: "Figures 105 through 111 show appointment-list states, appointment details, rescheduling, cancellation, appointment history, and cancelled-appointment details. Patients can review current and past requests and, when permitted, change or cancel a pending appointment with a reason.",
119: "Figures 112 through 119 show the complete web booking workflow. The patient selects a branch, date, time, and procedure; adds notes; reviews the privacy summary; confirms the request; and receives either a submission confirmation or a notice about an existing active appointment.",
127: "Figures 120 through 127 show web-based Oral Health Management. Patients can view a calendar and trends, open or edit daily logs, record symptom context and risk factors, and review recent patterns while clinical entries remain read-only where appropriate.",
129: "Figures 128 and 129 show the dental health education library and an article view. Patients can browse approved topics and open concise educational guidance related to daily oral care.",
133: "Figures 130 through 133 show mobile login and password recovery. A patient can sign in, request a reset using a registered email address, enter the verification code, and create a new password.",
135: "Figures 134 and 135 show the mobile home dashboard and care tools. The dashboard summarizes the next visit and oral-health activity, while the tools provide direct access to appointments, the electronic medical record, Oral Health Management, and settings.",
139: "Figures 136 through 139 show NgitifyBot on mobile and its conversation-management views. Patients can exchange supported messages, open conversation history, rename, pin, archive, or delete a chat, and revisit archived conversations.",
147: "Figures 140 through 147 show the mobile Oral Health Management dashboard, calendar, and four-step daily logging workflow. Patients record symptoms, optional symptom context, daily care habits, risk factors, and notes, then review the information before saving the log.",
152: "Figures 148 through 152 show oral-health trends for seven- and thirty-day periods. The views summarize check-ins, care habits, symptoms, and risk factors so patients can recognize patterns in their recorded data.",
157: "Figures 153 through 157 show guidance generated from recent oral-health logs, including a recommended visit window, selected-date context, suggested next actions, and editable current factors. These items support patient awareness without replacing professional diagnosis.",
160: "Figures 158 through 160 show mobile dental health education. Patients can open the education area, browse approved topics, and read an article containing practical, non-diagnostic oral-care guidance.",
164: "Figures 161 through 164 show mobile appointment management. Patients can view upcoming and past visits and submit a permitted reschedule or cancellation request with the required details.",
170: "Figures 165 through 170 show the mobile electronic medical record for medical and dental history, treatment history, odontogram data, tooth details, and odontogram history. The screens provide authorized, read-only access to clinical information recorded by the dental team.",
172: "Figures 171 and 172 show mobile radiograph history and the radiograph viewer. Patients can open a stored image and review the dentist-approved summary and related treatment information.",
177: "Figures 173 through 177 show the mobile profile, editable profile information, personal details, activity logs, and settings. Patients can maintain account information, review recorded account actions, and access security and notification controls.",
181: "Figures 178 through 181 show mobile security and notification settings. Users can verify their identity and set a new password, request an email-address change, and choose which appointment, visit-window, oral-health, symptom, and education reminders are enabled.",
}

assert len(captions) == 181

def set_font(run, name="Times New Roman", size=Pt(12), bold=False):
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)

def insert_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = paragraph._parent.add_paragraph()
    paragraph._parent._element.remove(p._p)
    new_p.addnext(p._p)
    p._p.getparent().remove(new_p)
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run)
    return p

doc = Document(SOURCE)
figure_paragraphs = []
for p in doc.paragraphs:
    if re.fullmatch(r"\s*FIGURE\s*\d+\s*", p.text, flags=re.IGNORECASE):
        figure_paragraphs.append(p)

if len(figure_paragraphs) != len(captions):
    raise RuntimeError(f"Expected {len(captions)} figure labels, found {len(figure_paragraphs)}")

for i, (p, title) in enumerate(zip(figure_paragraphs, captions), 1):
    # Preserve the existing paragraph properties/page-flow behavior, replacing
    # only its label and local formatting.
    for r in p.runs:
        r._element.getparent().remove(r._element)
    run = p.add_run(f"Figure {i}. {title}")
    set_font(run)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = False
    # Keep the immediately preceding figure paragraph on the same page as its
    # caption. This prevents a screenshot at the bottom of one page from leaving
    # its caption alone at the top of the next page.
    prev = p._p.getprevious()
    while prev is not None and not prev.xpath('.//w:drawing | .//w:pict'):
        prev = prev.getprevious()
    if prev is not None:
        from docx.text.paragraph import Paragraph
        image_p = Paragraph(prev, p._parent)
        image_p.paragraph_format.keep_with_next = True

# Insert shared narrative after the last caption in each logical group. Work in
# reverse order so insertion never invalidates later paragraph references.
for end_figure in sorted(descriptions, reverse=True):
    insert_after(figure_paragraphs[end_figure - 1], descriptions[end_figure])

# Apply consistent typography to the two platform labels while preserving the
# document's existing hierarchy and all image sizing.
for p in doc.paragraphs:
    if p.text.strip() in {"Web Application", "Mobile Application"}:
        for r in p.runs:
            set_font(r, bold=True)

doc.core_properties.title = "4.4 Description of the System - NGITIFY"
doc.save(OUTPUT)
print(f"Saved {OUTPUT}")
print(f"Updated {len(figure_paragraphs)} captions and inserted {len(descriptions)} descriptions")
