from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

SOURCE = Path(r"C:\Users\Administrator\Documents\4.4 DESCRIPTION OF THE SYSTEM - Captioned.docx")
OUTPUT = Path(r"C:\Users\Administrator\Documents\4.4 DESCRIPTION OF THE SYSTEM - Captioned with Public Website.docx")
IMAGE_DIR = Path(r"C:\Users\Administrator\AppData\Local\Temp\msohtmlclip1\01")

new_figures = [
    (IMAGE_DIR / "clip_image001.png", "Home Page (Web)", 6.25),
    (IMAGE_DIR / "clip_image003.png", "About Page (Web)", 6.25),
    (IMAGE_DIR / "clip_image005.png", "Services Page (Web)", 6.25),
    (IMAGE_DIR / "clip_image007.png", "Contact Page (Web)", 6.25),
    (IMAGE_DIR / "clip_image009.png", "Appointment Page 1 (Web)", 6.25),
    (IMAGE_DIR / "clip_image011.png", "Appointment Page 2 (Web)", 3.20),
]

for path, _, _ in new_figures:
    if not path.exists():
        raise FileNotFoundError(path)


def set_font(run, name="Times New Roman", size=Pt(12), bold=False):
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)


def renumber_description(text, offset=6):
    """Shift figure references in a narrative paragraph by the given offset."""
    return re.sub(r"(?<![A-Za-z])\d+(?![A-Za-z])", lambda m: str(int(m.group()) + offset), text)


doc = Document(SOURCE)

# Renumber the existing 181 captions and their narrative references before
# inserting the new front matter. Captions become Figures 7 through 187.
existing_caption_elements = set()
for p in doc.paragraphs:
    match = re.fullmatch(r"Figure (\d+)\. (.+)", p.text.strip())
    if not match:
        continue
    old_num = int(match.group(1))
    title = match.group(2)
    for run in p.runs:
        run._element.getparent().remove(run._element)
    run = p.add_run(f"Figure {old_num + 6}. {title}")
    set_font(run)
    existing_caption_elements.add(p._p)

for p in doc.paragraphs:
    if p._p in existing_caption_elements:
        continue
    text = p.text.strip()
    if not re.match(r"^Figures? \d", text):
        continue
    shifted = renumber_description(text)
    for run in p.runs:
        run._element.getparent().remove(run._element)
    run = p.add_run(shifted)
    set_font(run)

# Locate the current login screenshot, which is the first drawing after the Web
# Application label. New public website figures are inserted before it.
web_label = next(p for p in doc.paragraphs if p.text.strip() == "Web Application")
anchor = None
node = web_label._p.getnext()
while node is not None:
    if node.xpath('.//w:drawing | .//w:pict'):
        from docx.text.paragraph import Paragraph
        anchor = Paragraph(node, web_label._parent)
        break
    node = node.getnext()
if anchor is None:
    raise RuntimeError("Could not locate the web login figure")

for num, (image_path, title, width_inches) in enumerate(new_figures, 1):
    image_p = anchor.insert_paragraph_before()
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.paragraph_format.keep_with_next = True
    # Start clean two-figure spreads for Figures 3 and 5, and give the long
    # appointment form its own page.
    if num in {3, 5, 6}:
        image_p.paragraph_format.page_break_before = True
    image_run = image_p.add_run()
    image_run.add_picture(str(image_path), width=Inches(width_inches))

    caption_p = anchor.insert_paragraph_before()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_before = Pt(3)
    caption_p.paragraph_format.space_after = Pt(3)
    caption_run = caption_p.add_run(f"Figure {num}. {title}")
    set_font(caption_run)

# A shared description follows the second appointment screenshot, matching the
# grouping convention used elsewhere in the document.
description_p = anchor.insert_paragraph_before()
description_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
description_p.paragraph_format.first_line_indent = Inches(0.5)
description_p.paragraph_format.space_before = Pt(6)
description_p.paragraph_format.space_after = Pt(6)
description_p.paragraph_format.line_spacing = 1.0
description = (
    "Figures 1 through 6 show the public-facing pages of the Dentime Dental Clinic website. "
    "Visitors can learn about the clinic, browse available services, review contact information, "
    "and open the appointment page. The appointment request form collects the guest's identifying "
    "and contact details, selected branch, preferred schedule, procedure, concern or message, data "
    "privacy consent, and CAPTCHA verification before the request is submitted."
)
set_font(description_p.add_run(description))

# Begin the existing authenticated-system sequence on a fresh page.
anchor.paragraph_format.page_break_before = True

# Remove the source document's trailing empty paragraph, which otherwise
# exports as an unnecessary blank final page in Microsoft Word.
while doc.paragraphs and not doc.paragraphs[-1].text.strip() and not doc.paragraphs[-1]._p.xpath('.//w:drawing | .//w:pict'):
    trailing = doc.paragraphs[-1]._p
    trailing.getparent().remove(trailing)

doc.core_properties.title = "4.4 Description of the System - NGITIFY"
doc.save(OUTPUT)
print(f"Saved {OUTPUT}")
print("Inserted 6 public website figures; renumbered existing figures to 7-187")
