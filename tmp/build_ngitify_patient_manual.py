from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path


OUT = Path('output/documents/NGITIFY_Patient_User_Manual_Web_and_Mobile_Revised_v2.docx')
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = '01538B'
CYAN = '2DCCF6'
NAVY = '17364A'
PALE = 'EAF7FB'
LIGHT = 'F3F7F9'
MUTED = '607D8B'
WHITE = 'FFFFFF'
RED = 'B91C1C'
GREEN = '15803D'


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_border(cell, color='B8DCE8', size='8', val='single'):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in('w:tcBorders')
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right'):
        el = borders.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            borders.append(el)
        el.set(qn('w:val'), val)
        el.set(qn('w:sz'), size)
        el.set(qn('w:color'), color)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for tag, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{tag}'))
        if node is None:
            node = OxmlElement(f'w:{tag}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def set_table_widths(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    total = sum(int(w * 1440) for w in widths)
    tbl_w.set(qn('w:w'), str(total))
    tbl_w.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(int(width * 1440)))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(int(widths[idx] * 1440)))
            tc_w.set(qn('w:type'), 'dxa')


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run('Page ')
    r.font.name = 'Calibri'
    r.font.size = Pt(9)
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    r._r.extend([fld_char1, instr, fld_char2])


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.8)
sec.bottom_margin = Inches(0.75)
sec.left_margin = Inches(0.85)
sec.right_margin = Inches(0.85)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(NAVY)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.18

for name, size, color, before, after in [
    ('Title', 30, NAVY, 0, 8),
    ('Subtitle', 14, MUTED, 0, 12),
    ('Heading 1', 17, BLUE, 18, 8),
    ('Heading 2', 14, BLUE, 14, 6),
    ('Heading 3', 12, NAVY, 10, 4),
]:
    st = styles[name]
    st.font.name = 'Calibri'
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = name != 'Subtitle'
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for name in ('List Number', 'List Bullet'):
    st = styles[name]
    st.font.name = 'Calibri'
    st.font.size = Pt(10.5)
    st.paragraph_format.left_indent = Inches(0.38)
    st.paragraph_format.first_line_indent = Inches(-0.19)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.18

if 'Screenshot Caption' not in styles:
    sc = styles.add_style('Screenshot Caption', WD_STYLE_TYPE.PARAGRAPH)
else:
    sc = styles['Screenshot Caption']
sc.font.name = 'Calibri'
sc.font.size = Pt(9)
sc.font.italic = True
sc.font.color.rgb = RGBColor.from_string(MUTED)
sc.paragraph_format.space_before = Pt(3)
sc.paragraph_format.space_after = Pt(10)
sc.paragraph_format.keep_with_next = False

header = sec.header
p = header.paragraphs[0]
p.text = 'NGITIFY  |  PATIENT USER MANUAL'
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for run in p.runs:
    run.font.name = 'Calibri'
    run.font.size = Pt(8.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(MUTED)
add_page_number(sec.footer.paragraphs[0])


def add_title(text, subtitle=None):
    p = doc.add_paragraph(style='Title')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)
    if subtitle:
        s = doc.add_paragraph(style='Subtitle')
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.add_run(subtitle)


def add_note(label, text, tone='blue'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [6.55])
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    fill = {'blue': PALE, 'warn': 'FFF7E6', 'danger': 'FDECEC', 'success': 'ECFDF3'}.get(tone, PALE)
    border = {'blue': CYAN, 'warn': 'D97706', 'danger': RED, 'success': GREEN}.get(tone, CYAN)
    set_cell_shading(cell, fill)
    set_cell_border(cell, border, '10')
    set_cell_margins(cell, 140, 170, 140, 170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f'{label}: ')
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE if tone == 'blue' else NAVY)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


figure_counter = 0
def add_screenshot(caption, guidance='Capture the full screen with the relevant menu, form, or result visible.'):
    global figure_counter
    figure_counter += 1
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [6.25])
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'F8FBFC')
    set_cell_border(cell, '7AB8CB', '10', 'dashed')
    set_cell_margins(cell, 260, 220, 260, 220)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run('INSERT SCREENSHOT HERE')
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    q = cell.add_paragraph(guidance)
    q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    q.paragraph_format.space_after = Pt(10)
    q.runs[0].font.size = Pt(9)
    q.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    cap = doc.add_paragraph(f'Figure {figure_counter}. {caption}', style='Screenshot Caption')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_steps(steps):
    for step in steps:
        p = doc.add_paragraph(style='List Number')
        if isinstance(step, tuple):
            label, detail = step
            r = p.add_run(label)
            r.bold = True
            p.add_run(detail)
        else:
            p.add_run(step)


def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        if isinstance(item, tuple):
            label, detail = item
            r = p.add_run(label)
            r.bold = True
            p.add_run(detail)
        else:
            p.add_run(item)


def add_procedure(title, purpose, steps, screenshot=None, note=None, level=2):
    doc.add_heading(title, level=level)
    if purpose:
        p = doc.add_paragraph()
        r = p.add_run('Purpose. ')
        r.bold = True
        p.add_run(purpose)
    add_steps(steps)
    if note:
        add_note(note[0], note[1], note[2] if len(note) > 2 else 'blue')
    if screenshot:
        add_screenshot(screenshot)


def add_matrix(headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header_text)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(str(text))
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


# Cover
doc.add_paragraph().paragraph_format.space_after = Pt(85)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PATIENT GUIDE')
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor.from_string(CYAN)
add_title('NgitiFy Patient User Manual', 'Web and Mobile Applications')
doc.add_paragraph().paragraph_format.space_after = Pt(30)
add_note('Document status', 'Screenshot-ready draft based on the current patient web and mobile implementation. Replace every marked figure box with the corresponding screenshot before publication.')
doc.add_paragraph().paragraph_format.space_after = Pt(110)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Dentime Dental Clinic')
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor.from_string(BLUE)
p = doc.add_paragraph('Manual version: 1.2  |  Revised: September 2026')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
doc.add_page_break()

# Front matter
doc.add_heading('Copyright and Use Notice', level=1)
doc.add_paragraph('This manual is intended for authorized use, operation, demonstration, training, and evaluation of the NgitiFy patient web and mobile applications. The contents may be revised as the system is improved or updated.')
doc.add_paragraph('Patient and account information used for training, testing, or screenshots must be handled responsibly. Do not place real or unauthorized patient information in demonstration images. Blur or replace personal identifiers before publishing this manual.')
add_note('Clinical notice', 'NgitiFy supports appointment management, self-tracking, patient education, and explanation of authorized records. It does not diagnose dental disease or replace advice, diagnosis, or treatment from a dentist.', 'warn')

doc.add_heading('How to Use This Draft', level=1)
add_bullets([
    'Replace each “INSERT SCREENSHOT HERE” box with the requested web or mobile screenshot.',
    'Keep the figure caption below each screenshot. Renumber figures only if you add or remove images.',
    'Replace bracketed deployment details such as the production web address, Android APK download link, and iOS TestFlight public link.',
    'After inserting screenshots, update the table of contents in Microsoft Word and check page breaks.',
    'Use a training account with fictional information when capturing screenshots.',
])

doc.add_heading('Table of Contents', level=1)
toc_lines = [
    '1. General Information', '2. System Summary', '3. Getting Started',
    '4. Patient Web Application', '5. Patient Mobile Application',
    '6. Shared Rules and Important Notes', '7. Troubleshooting',
    'Appendix A. Screenshot Checklist', 'Appendix B. Quick Reference',
]
for line in toc_lines:
    doc.add_paragraph(line, style='List Bullet')
add_note('Word tip', 'After all screenshots are inserted, use References > Table of Contents to generate page numbers from the heading styles.')
doc.add_page_break()

# 1
doc.add_heading('1. General Information', level=1)
doc.add_heading('1.1 Purpose', level=2)
doc.add_paragraph('This manual explains how patients access and use NgitiFy through the public website, authenticated patient web portal, and mobile application. It covers account activation, sign-in and recovery, appointments, Oral Health Management, Dental Health Education, electronic medical records, notifications, NgitiBot, profile maintenance, settings, and activity logs.')
doc.add_heading('1.2 Intended User', level=2)
doc.add_paragraph('The intended user is a registered Dentime Dental Clinic patient with an active NgitiFy account. Some public website functions, including appointment requests, may be used before an account is activated.')
doc.add_heading('1.3 Terms Used in This Manual', level=2)
add_matrix(['Term', 'Meaning'], [
    ('Patient Log', 'Information entered by the patient in the Daily Oral Health Log.'),
    ('Clinic Record', 'Information recorded by authorized clinic personnel.'),
    ('Recommended Visit Window', 'Timing guidance based on supported clinic information and approved system rules.'),
    ('Dentist Suggested Next Visit', 'Follow-up timing entered by the dentist and shown as the primary planned-care recommendation when available.'),
    ('NgitiBot', 'The patient AI care companion for dental education and explanations. It does not provide a diagnosis.'),
    ('EMR', 'Electronic Medical Record containing authorized medical, dental, treatment, odontogram, and radiograph information.'),
], [1.75, 4.75])

# 2
doc.add_heading('2. System Summary', level=1)
doc.add_heading('2.1 System Overview', level=2)
doc.add_paragraph('NgitiFy is the patient-facing digital service of Dentime Dental Clinic. The web application supports public clinic information, online appointment requests, pre-registration, account access, and a full patient portal. The mobile application provides convenient access to the same patient account and adds a guided first-login onboarding flow.')
doc.add_heading('2.2 Requirements', level=2)
add_matrix(['Website', 'Android', 'iOS'], [
    ('Desktop or laptop computer', 'Compatible Android phone', 'Compatible iPhone'),
    ('Stable internet connection', 'Wi-Fi or mobile-data connection', 'Wi-Fi or mobile-data connection'),
    ('Current Chrome, Edge, Firefox, or Safari', 'Sufficient storage and permission to install an APK', 'TestFlight installed from the App Store'),
    ('Production website address', 'Clinic-provided NgitiFy APK download link', 'Clinic-provided TestFlight public link'),
    ('Active patient account for the portal', 'Active patient account', 'Active patient account'),
], [2.2, 2.15, 2.15])
doc.add_heading('2.3 Patient Access Level', level=2)
doc.add_paragraph('Patients may view and manage only the information and functions associated with their authenticated account. Patient-entered logs may be edited by the patient. Clinic records, dentist recommendations, odontograms, and approved radiograph findings are read-only in the patient applications.')
doc.add_heading('2.4 Web and Mobile Navigation', level=2)
add_matrix(['Area', 'Web', 'Mobile'], [
    ('Home', 'Dashboard', 'Home tab'), ('Appointments', 'My Appointments', 'Visits tab'),
    ('Oral health', 'Oral Health Management', 'Health tab'), ('Records', 'Electronic Medical Record', 'Records tab'),
    ('Account', 'My Profile and Settings', 'Profile tab'), ('AI help', 'Floating NgitiBot button', 'Floating NgitiBot button'),
], [1.35, 2.55, 2.6])

# 3
doc.add_heading('3. Getting Started', level=1)
add_procedure('3.1 Open the Public Website', 'Access clinic information or start an appointment request.', [
    'Use a desktop or laptop computer. The patient website instructions in this manual are intended for these devices only.',
    'Open a supported web browser.', 'Enter [INSERT PRODUCTION WEB ADDRESS] in the address bar and press Enter.',
    'Use Home, About, Services, Contact Us, or Appointment in the top navigation.',
    'Select LOGIN when you are ready to access the patient portal.'
], 'NgitiFy public website on a desktop or laptop')
add_procedure('3.2.1 Install NgitiFy on Android Using the APK Link', 'Download and install the Android application using the secure APK link provided by Dentime Dental Clinic.', [
    'On the Android phone, open [INSERT NGITIFY ANDROID APK DOWNLOAD LINK].',
    'Tap Download or Download APK. If the browser displays a warning about downloading an APK file, continue only when the link was provided by Dentime Dental Clinic.',
    'When the download finishes, open the download notification or the Downloads folder, then tap the NgitiFy APK file.',
    'If installation is blocked, tap Settings and allow Install unknown apps for the browser or file manager used to open the APK. Return to the installer afterward.',
    'Tap Install and wait for the installation to finish.',
    'Tap Open, or find and tap the NgitiFy icon on the device.',
    'When the NgitiFy login screen appears, enter the active patient account credentials.'
], 'Android APK download, security permission, installation, and NgitiFy sign-in', note=('Install securely', 'Use only the current APK link supplied by Dentime Dental Clinic. After installation, you may turn off the Install unknown apps permission again. When an updated version is released, download the new APK from the clinic link and install it over the existing application.', 'warn'))
add_procedure('3.2.2 Install NgitiFy on iOS Through TestFlight', 'Install the iOS beta using the public TestFlight invitation.', [
    'On the iPhone, open the App Store.',
    'Search for TestFlight, then tap Get or the download icon.',
    'Open [INSERT NGITIFY TESTFLIGHT PUBLIC LINK] or scan [INSERT NGITIFY TESTFLIGHT QR CODE].',
    'On the invitation page, tap View in TestFlight or Start Testing.',
    'In TestFlight, tap Accept if the invitation has not yet been accepted.',
    'Tap Install beside NgitiFy and wait for installation to finish.',
    'Tap Open, then sign in using the active NgitiFy patient account.',
    'When a newer beta build is available, open TestFlight and tap Update beside NgitiFy.'
], 'iOS TestFlight public invitation and NgitiFy installation', note=('TestFlight availability', 'The public link may stop accepting testers when its limit is reached or when the clinic disables it. TestFlight beta builds also expire, so install the latest build provided by the clinic.', 'warn'))
add_procedure('3.3 Request an Appointment from the Public Website', 'Send a booking request before or without signing in.', [
    'Open the Appointment page.', 'Enter First Name, Last Name, Phone, Email, Date of Birth, and Gender.',
    'Select the Branch. Available dates and times are based on that branch.',
    'Select a Preferred Date. Past dates, Sundays, and fully booked dates cannot be selected.',
    'Select an available Preferred Time.', 'Select the Procedure.',
    'Optionally enter a Concern / Message. If entered, provide enough detail for the clinic.',
    'Open and review the Data Privacy Notice, then select the consent acknowledgement.',
    'Complete the captcha verification.', 'Select the request-submission button and wait for Request received.'
], 'Public website appointment-request form and success message', note=('Phone format', 'Enter the Philippine mobile number as 9xxxxxxxxx; the +63 prefix is displayed separately.'))
add_procedure('3.4 Complete Patient Pre-Registration', 'Complete the clinic registration form sent after an eligible appointment request.', [
    'Open the secure pre-registration link sent by the clinic. Do not share this link.',
    'Confirm the appointment-linked First Name, Last Name, Phone, Email, Branch, Birthdate, and Gender.',
    'Complete Identity: occupation, civil status, nationality, religion, optional phone details, and referral information.',
    'Complete Contacts: home address and emergency-contact information. Complete guardian details when required.',
    'Complete Medical & Dental: reason for consultation, dental history, physician information where applicable, allergies, medical history, blood type, blood pressure, and relevant notes.',
    'Review the Data Privacy Act information and digital consent.',
    'Enter the signer name, signer role, and date signed, then acknowledge each required consent.',
    'Review all sections, correct any validation messages, and submit the form.',
    'Wait for the completion message, then check the registered email for the activation link.'
], 'Patient pre-registration steps and final review', note=('Minor patients', 'Guardian information and the appropriate signer role are required when the registration rules identify the patient as a minor.', 'warn'))
add_procedure('3.5 Activate the Account and Set a Password', 'Create the password for a clinic-created patient account.', [
    'Open the activation link sent to the registered email address.', 'Confirm that the page identifies the NgitiFy account activation process.',
    'Enter a new password containing at least 8 characters, one uppercase letter, one lowercase letter, one number, and one special character.',
    'Re-enter the same password in the confirmation field.', 'Select the activation button.',
    'After the success message appears, continue to Login.'
], 'Account activation and password setup')
add_procedure('3.6 Sign In on the Web', 'Open the authenticated patient portal.', [
    'Select LOGIN on the public website.', 'Enter the registered Email Address.', 'Enter the account Password.',
    'Use the eye icon if you need to show or hide the password.', 'Select LOGIN.',
    'Wait for the Patient Dashboard to load.'
], 'Web patient login page')
add_procedure('3.7 Sign In on Mobile', 'Open the patient mobile application.', [
    'Open NgitiFy.', 'Enter the registered Email Address and Password.',
    'Use the eye icon to show or hide the password if needed.', 'Tap LOGIN.',
    'On the first successful mobile sign-in, complete the privacy acknowledgement and onboarding steps described in Section 5.1.'
], 'Mobile patient login screen')
add_procedure('3.8 Reset a Forgotten Password', 'Set a new password when the current password is unavailable.', [
    'From Login, select Forgot Password?', 'Enter the registered email address and select Send Code or Send Verification Code.',
    'Open the email from NgitiFy and find the 6-digit verification code.',
    'Enter the complete code and select Enter or Verify Code. Use Resend after the cooldown if needed.',
    'Enter and confirm a new password that meets every displayed password requirement.',
    'Select Enter or Reset Password.', 'After the success message appears, return to Login and sign in with the new password.'
], 'Forgot-password email, verification code, and new-password screens', note=('Privacy', 'The recovery flow may continue even if an address is not recognized. Check the inbox associated with the registered patient account.'))

# 4 WEB
doc.add_heading('4. Patient Web Application', level=1)
doc.add_heading('4.1 Web Portal Navigation', level=2)
doc.add_paragraph('The left sidebar contains Dashboard, My Appointments, Electronic Medical Record, Oral Health Management, Dental Health Education, Notifications, My Profile, Settings, and Logout. Select the sidebar arrow to expand or collapse the labels. A notification badge shows unread items. The floating NgitiBot button opens the patient AI companion.')
add_screenshot('Expanded patient web sidebar and floating NgitiBot button')
add_procedure('4.2 Use the Patient Dashboard', 'Review the day’s schedule, oral-health status, and next-care timing.', [
    'Select Dashboard from the sidebar.', 'Use the week strip or month calendar to choose a date.',
    'Review Selected-Day Schedule. Select Details on an appointment to see its information.',
    'Use My Appointments to open the full appointment hub.',
    'Review Today’s Oral Health Management and select Log Entry or Edit Entry.',
    'Review Next care timing, including Dentist Suggested Next Visit, Current Visit Guidance, source, and reason.',
    'Select Why am I seeing this? for the recommendation explanation.'
], 'Web Patient Dashboard')
add_procedure('4.3 View Appointments and History', 'Review upcoming, completed, and cancelled visits.', [
    'Select My Appointments.', 'Use Upcoming to view pending, confirmed, or in-clinic appointments.',
    'Use History to view completed or cancelled appointments.', 'Select an appointment card to open its details.',
    'Review procedure, dentist, branch, date, time, status, and any notes.'
], 'Web My Appointments - Upcoming and History')
add_procedure('4.4 Book an Appointment in the Web Portal', 'Submit a booking request using the patient’s assigned branch.', [
    'Open My Appointments and select Book New Appointment or the Book Appointment tab.',
    'Confirm the assigned branch and review the one-active-request rule.',
    'Select an available date. Past dates and fully booked dates are disabled.',
    'Select an available time slot.', 'Select a direct-book procedure.',
    'Optionally enter Additional Notes.', 'Review the booking snapshot.',
    'Open the privacy-policy summary and accept the privacy consent for booking use.',
    'Select the confirmation button and wait for Booking Submitted.',
    'Return to My Appointments to see the pending request.'
], 'Web patient appointment booking steps', note=('Booking rule', 'A patient cannot create another booking while an ongoing appointment or active request exists.'))
add_procedure('4.5 Cancel an Eligible Appointment', 'Cancel a pending or confirmed appointment.', [
    'Open My Appointments > Upcoming.', 'Select the eligible appointment.', 'Select Cancel Appointment.',
    'Enter a cancellation reason of at least 20 characters.', 'Select Confirm Cancellation.',
    'Verify that the appointment now appears as Cancelled in appointment history.'
], 'Web appointment cancellation dialog', note=('Availability', 'Cancel and Reschedule are not available for in-clinic, completed, archived, or queue-linked records.', 'warn'))
add_procedure('4.6 Reschedule an Eligible Appointment', 'Move an existing pending or confirmed appointment without creating a second record.', [
    'Open the eligible appointment details.', 'Select Reschedule.', 'Choose a new date.',
    'Choose an available time.', 'Optionally enter a reason.', 'Select Confirm Reschedule.',
    'Verify that the original appointment displays the new schedule.'
], 'Web appointment rescheduling dialog')
add_procedure('4.7 Record a Daily Oral Health Log', 'Save patient-entered symptoms, care habits, and other factors for today or an earlier date.', [
    'Select Oral Health Management.', 'Open Today, or choose a date from the date strip or calendar.',
    'Select at least one symptom, oral-care item, other/risk factor, or enter a note.',
    'For Toothache, Swelling, Jaw Pain, or Mouth Sore, optionally add Severity and Duration, then select Done.',
    'Select Save Entry. If an entry already exists, update it and save again.',
    'Confirm that the selected date is marked as a Patient Log.'
], 'Web Daily Oral Health Log', note=('Important', 'No Symptoms cannot remain selected together with another symptom. Future dates are read-only.', 'warn'))
add_procedure('4.8 Use the Oral Health Calendar', 'Review patient logs, clinic records, appointments, and recommendations by date.', [
    'Open Oral Health Management > Calendar.', 'Use the week selector or select Show Month.',
    'Choose a date with a marker.', 'Review the selected-date sources: Patient Log, Dental Appointment, Clinic Record, or Dentist Recommendation.',
    'Edit only the patient-entered log. Treat clinic records and dentist recommendations as read-only context.'
], 'Web Oral Health Management calendar and selected-date summary')
add_procedure('4.9 Review Trends and Health Factors', 'Review counts from saved Patient Logs and maintain background factors.', [
    'Open the Trends tab.', 'Review the 7-day and 30-day summaries for symptoms, daily care, and other/risk factors.',
    'If Not enough information yet appears, continue completing Daily Oral Health Logs.',
    'Scroll to Health factors.', 'Select or clear applicable factors, then select Save Factors.'
], 'Web oral-health trends and Health factors', note=('Interpretation', 'Trend counts are not a clinical oral-health score and do not provide a diagnosis.', 'warn'))
add_procedure('4.10 Review the Recommended Visit Window', 'Understand current visit timing and its source.', [
    'Open Oral Health Management or review the dashboard Next care timing card.',
    'Read the Recommended Visit Window and Current Visit Guidance.',
    'Review Based on and Source / Reason.', 'Select Why am I seeing this? to compare dentist and system sources.',
    'Use Book Predictive Visit when you are ready to request an available appointment.'
], 'Web Recommended Visit Window explanation', note=('Urgent concerns', 'Contact the clinic directly for persistent, worsening, severe, or concerning symptoms. Do not wait for an automated recommendation.', 'danger'))
add_procedure('4.11 Browse Dental Health Education', 'Find approved oral-health education and personalized topic matches.', [
    'Select Dental Health Education.', 'Review Recommended for You when matching patient-log topics are available.',
    'Enter a keyword in Search topics, or select a category.', 'Select an article to open its title, summary, full content, and What you can do guidance.',
    'Select Show All Topics to clear filters.', 'Select Open NgitiBot if you want an explanation of an educational topic.'
], 'Web Dental Health Education library and article')
add_procedure('4.12 Use NgitiBot', 'Ask supported dental-health questions or request explanations of existing care context.', [
    'Select the floating NgitiBot button.', 'Choose a suggested prompt or enter a dental-health question in Message NgitiBot.',
    'Select Send and wait for the explanation.', 'Use the chat-history button to open saved conversations.',
    'Select New conversation to start a separate chat.', 'Open conversation options to Rename, Pin/Unpin, Archive/Unarchive, or Delete.',
    'Unarchive a conversation before sending more messages to it.', 'Use the information button to review NgitiBot’s scope and limitations.'
], 'Web NgitiBot chat and saved conversations', note=('Scope', 'NgitiBot can explain dental education, appointments, the Recommended Visit Window, Oral Health Management context, and dentist-approved radiograph information. It cannot diagnose or replace the dentist.', 'warn'))
add_procedure('4.13 View the Electronic Medical Record', 'Review authorized clinical information.', [
    'Select Electronic Medical Record.', 'Review Patient Profile and identification information.',
    'Open Medical and Dental History.', 'Open Treatment History; use View to expand a record.',
    'Open Odontogram and review tooth statuses and history. Patient access is read-only.',
    'Open Radiograph Images and select an available image to review dentist-approved information.',
    'Use Export PDF to open the formal patient-record preview, then print or choose Save as PDF if permitted.'
], 'Web Electronic Medical Record tabs', note=('Record corrections', 'Contact the clinic if clinical information appears incomplete or inaccurate. Do not use Patient Logs to overwrite clinic records.'))
add_procedure('4.14 Manage Notifications', 'Find notifications and change their read status.', [
    'Select Notifications.', 'Use Search title or message when needed.',
    'Filter by Type, Read Status, and date range.', 'Select a notification to open its details.',
    'Use Mark as Read or Mark as Unread in the detail view.', 'Select Mark All as Read to clear all unread items.',
    'Use Previous and Next to move between result pages when available.'
], 'Web notification inbox and detail view')
add_procedure('4.15 View and Edit the Patient Profile', 'Review or update patient demographic and contact information.', [
    'Select My Profile to review account and patient information.', 'Select Edit Profile.',
    'Optionally select CHANGE to replace the profile image.',
    'Update Personal Information, Contact Details, Emergency Contact and Guardian, or Home Address.',
    'Use the format 9xxxxxxxxx for mobile numbers.', 'Select Save Changes, then select Yes in the confirmation dialog.',
    'Wait for the profile-updated confirmation.'
], 'Web Edit Profile form', note=('Reset', 'Select Reset to discard unsaved form changes or Cancel to return to the profile.'))
add_procedure('4.16 Change the Password', 'Replace the signed-in patient’s password.', [
    'Open Settings > Account Security.', 'Enter the Current Password and select Verify.',
    'Enter a New Password that meets all displayed requirements and differs from the current password.',
    'Re-enter it in Confirm New Password.', 'Select Update Password and wait for confirmation.'
], 'Web Settings - Account Security')
add_procedure('4.17 Change the Email Address', 'Request a verified change to the sign-in email.', [
    'Open Settings > Account Actions.', 'Select Change Email.', 'Enter the New Email Address.',
    'Enter the Current Password and select Verify.', 'Select Send Verification Link.',
    'Select OK, Log Me Out.', 'Open the verification link sent to the new email address, then sign in again.'
], 'Web change-email verification flow')
add_procedure('4.18 Set Notification and Privacy Preferences', 'Control patient alerts and personalized education.', [
    'Open Settings > Notifications.', 'Turn Appointment Alerts, Recommended Visit Window Reminders, Daily Oral Health Management Reminder, Symptom Follow-Up Reminders, and Dental Health Education / Dental Health Tips on or off.',
    'When the daily reminder is on, choose a reminder time.', 'Open Privacy and Data.',
    'Turn Personalized Dental Health Education on or off. Changes are shared with the mobile app.'
], 'Web notification and privacy settings')
add_procedure('4.19 Review Activity Logs', 'View the read-only history of actions associated with the account.', [
    'Open Settings > Account Actions and select Activity Logs, or open the page from My Profile when available.',
    'Search by action, detail, or category.', 'Apply date and category filters.',
    'Select View Details for the complete recorded entry.', 'Use Export CSV when a local copy is needed and permitted.',
    'Use Previous and Next to move through pages.'
], 'Web patient Activity Logs')
add_procedure('4.20 Log Out of the Web Portal', 'End the current browser session.', [
    'Select Logout in the sidebar.', 'Review the Confirm Logout message.',
    'Select Yes, Logout. Select Cancel to remain signed in.'
], 'Web logout confirmation')

# 5 MOBILE
doc.add_heading('5. Patient Mobile Application', level=1)
add_procedure('5.1 Complete First-Login Privacy and Onboarding', 'Configure the patient’s initial mobile experience.', [
    'After the first successful sign-in, review the Privacy Policy and tap I Agree and Continue.',
    'On Welcome to NgitiFy, tap Get Started.', 'Enter the preferred name or nickname and continue.',
    'Select all topics under What would you like help with?, or tap Skip for now.',
    'Choose brushing and flossing routine options, or skip.',
    'Turn Oral Health Management, Dental Health Education, Visit Recommendations, and Appointment Updates preferences on or off.',
    'Choose the useful notification reminders.', 'Review Your information matters and open Read Privacy Information if desired.',
    'On You’re all set, tap Go to Dashboard or Log Today’s Oral Health.'
], 'Mobile onboarding sequence', note=('Resume support', 'If onboarding is interrupted, NgitiFy saves completed steps and resumes from the appropriate screen after sign-in.'))
doc.add_heading('5.2 Mobile Navigation', level=2)
doc.add_paragraph('The bottom navigation contains Home, Health, Visits, Records, and Profile. The floating NgitiBot button is available above the bottom navigation. The Home notification bell opens Notifications; the profile avatar opens a quick account menu.')
add_screenshot('Mobile bottom navigation, notification bell, profile button, and NgitiBot launcher')
add_procedure('5.3 Use the Mobile Home Screen', 'Review the next visit, recommendation, and common care tools.', [
    'Tap Home.', 'Review Next Visit, including dentist, procedure, status, date, time, and branch.',
    'Tap View visits or Book now as needed.', 'Review the Recommended Visit Window card and tap it for full care context.',
    'Use Care Tools to open Appointments, Electronic Medical Record, Oral Health Management, or Settings.',
    'Pull down to refresh the dashboard.'
], 'Mobile patient Home screen')
add_procedure('5.4 View, Book, Cancel, or Reschedule Visits', 'Manage patient appointments from the Visits tab.', [
    'Tap Visits.', 'Review Upcoming Visits and Visit History.', 'Tap Book Appointment to start a new request.',
    'Confirm the assigned branch, select a procedure, select an available date, and select an available time.',
    'Enter optional notes, review the privacy information, accept consent, and submit the booking.',
    'To cancel, tap Cancel on an eligible appointment, enter at least 20 characters, and tap Confirm Cancellation.',
    'To reschedule, tap Reschedule, choose a new date and available time, enter an optional reason, and tap Confirm Reschedule.'
], 'Mobile Visits screen and booking flow', note=('Booking limits', 'Only pending or confirmed appointments can be changed by the patient. One ongoing appointment or active request is allowed at a time.', 'warn'))
add_procedure('5.5 Create or Update a Daily Oral Health Log', 'Record daily oral-health information from the Health tab.', [
    'Tap Health > Today.', 'Choose today or an earlier date.',
    'Select symptoms, oral-care habits, and other/risk factors that apply.',
    'For supported symptoms, add optional severity and duration information.', 'Optionally enter a note.',
    'Tap Save Entry and confirm the saved-log message.', 'Open the same date later to update the patient-entered entry.'
], 'Mobile Health - Today and Daily Oral Health Log')
add_procedure('5.6 Use Calendar and Oral Health Management History', 'Review information recorded for a selected date.', [
    'Tap Health > Calendar.', 'Choose a date from the week strip or month calendar.',
    'Review markers and the selected-date information.', 'Distinguish Patient Log from Dental Appointment, Clinic Record, and Dentist Recommendation.',
    'Edit only patient-entered information.'
], 'Mobile oral-health calendar and history')
add_procedure('5.7 Review Predictive Visit Guidance', 'View the Recommended Visit Window and explanation.', [
    'Tap Health > Predictive Visit.', 'Review the current visit window or No visit window is available yet.',
    'Read the source, reason, timing, and dentist recommendation when present.',
    'Use Book Appointment if you want to request an available visit.',
    'Contact the clinic directly if the guidance says to make contact sooner or symptoms are concerning.'
], 'Mobile Predictive Visit and Recommended Visit Window')
add_procedure('5.8 Review Trends and Health Factors', 'Understand recent self-tracking patterns.', [
    'Tap Health > Oral Health Management.', 'Switch between 7 days and 30 days.',
    'Review counts for symptoms, oral-care actions, and other/risk factors.',
    'If Not enough information yet appears, continue recording Daily Oral Health Logs.',
    'Review and update Health factors when the option is presented.'
], 'Mobile Oral Health Management trends')
add_procedure('5.9 Browse Dental Health Education', 'Read approved general or context-matched information.', [
    'Tap Health > Dental Health Education.', 'Review Recommended for You when available.',
    'Choose a category or enter a search keyword.', 'Tap an article to read its summary, content, and What you can do guidance.',
    'Clear the search or choose All Topics to return to the full library.'
], 'Mobile Dental Health Education library and article')
add_procedure('5.10 View the Electronic Medical Record', 'Review read-only clinical information from the Records tab.', [
    'Tap Records.', 'Open Medical and Dental History.', 'Open Treatment History and expand a treatment when available.',
    'Open Odontogram and review the tooth-status legend and recorded updates.',
    'Open Radiograph Images and tap a radiograph.', 'Review the image, date, number, dentist-approved summary, and recorded findings.',
    'Tap Explain this record to ask NgitiBot for an educational explanation.'
], 'Mobile Records tabs and X-Ray Viewer', note=('Read-only', 'Patients cannot edit treatment records, odontograms, or dentist-approved radiograph information.'))
add_procedure('5.11 Use NgitiBot on Mobile', 'Ask supported dental-health questions and manage saved chats.', [
    'Tap the floating NgitiBot button.', 'Choose a suggested prompt or type a supported dental-health question.',
    'Tap Send and wait for the response.', 'Open chat history to view saved conversations.',
    'Tap New conversation to start another chat.', 'Open a conversation’s options to Rename, Pin/Unpin, Archive/Unarchive, or Delete.',
    'Unarchive an archived chat before continuing it.', 'Close the sheet to return to the current mobile screen.'
], 'Mobile NgitiBot chat and conversation history')
add_procedure('5.12 Manage Notifications', 'Review patient alerts and follow linked content.', [
    'From Home, tap the notification bell, or open Notifications from the profile menu.',
    'Review the unread count and notification list.', 'Tap an unread notification to mark it read and open the related screen when supported.',
    'Tap Mark all as read to clear all unread notifications.', 'Pull down to refresh the inbox.'
], 'Mobile Notifications inbox')
add_procedure('5.13 View the Profile and Details', 'Review patient profile, quick information, and account shortcuts.', [
    'Tap Profile.', 'Review the patient name, email, age, blood type, and Health Notes count.',
    'Use the action cards to open My Details, Edit Profile, Activity Logs, Settings, or NgitiBot.',
    'Tap My Details to review profile sections.', 'Tap Edit Profile if corrections are needed.'
], 'Mobile My Profile and My Details')
add_procedure('5.14 Edit the Mobile Profile', 'Update patient information from a phone.', [
    'From Profile, tap Edit Profile.', 'Update Personal Information and Contact Details.',
    'Update Emergency Contact and Guardian information where applicable.', 'Update Home Address.',
    'Use 9xxxxxxxxx for mobile numbers.', 'Tap Save, review the confirmation, and wait for Profile Updated.'
], 'Mobile Edit Profile sections')
add_procedure('5.15 Change the Mobile Password', 'Replace the account password from Settings.', [
    'Open Profile > Settings > Change Password.', 'Enter the Current Password and complete identity verification.',
    'Enter a New Password that meets every displayed rule.', 'Re-enter it in Confirm Password.',
    'Tap the password-update button and wait for confirmation.'
], 'Mobile Settings - Change Password')
add_procedure('5.16 Change the Mobile Email Address', 'Request a verified change to the patient email.', [
    'Open Profile > Settings > Change Email.', 'Enter the new email address.',
    'Enter and verify the Current Password.', 'Send the verification link.',
    'Open the link in the new email inbox, then sign in again with the verified address.'
], 'Mobile Settings - Change Email')
add_procedure('5.17 Set Mobile Notification and Privacy Preferences', 'Control alerts shared across mobile and web.', [
    'Open Profile > Settings.', 'Under Notifications, turn Appointment Alerts, Recommended Visit Window Reminders, Daily Oral Health Management Reminder, Symptom Follow-Up Reminders, and Dental Health Education / Dental Health Tips on or off.',
    'If the daily reminder is on, enter a valid 24-hour reminder time such as 20:00.',
    'Under Privacy & Data, turn Personalized Dental Health Education on or off.',
    'Wait for each setting to save before leaving the screen.'
], 'Mobile notification and privacy settings')
add_procedure('5.18 Review Mobile Activity Logs', 'View the read-only account action history.', [
    'Open Profile > Activity Logs.', 'Review each action, category, details, and time.',
    'Scroll to load more entries when available.', 'Pull down to refresh.'
], 'Mobile Activity Logs')
add_procedure('5.19 Log Out of the Mobile Application', 'End the current mobile session.', [
    'Open the profile quick menu or Profile > Settings.', 'Tap Log Out.',
    'Confirm the logout request. Tap Cancel to remain signed in.'
], 'Mobile logout confirmation')

# 6
doc.add_heading('6. Shared Rules and Important Notes', level=1)
add_bullets([
    ('Assigned branch. ', 'Authenticated booking is limited to the branch assigned to the patient account.'),
    ('One active request. ', 'A patient with an ongoing appointment or active request cannot create another until the current record is resolved.'),
    ('Appointment changes. ', 'Only eligible pending or confirmed appointments can be cancelled or rescheduled by the patient.'),
    ('Cancellation reason. ', 'A patient cancellation requires a meaningful reason of at least 20 characters.'),
    ('Patient vs. clinic information. ', 'Patient Logs are self-reported and editable; clinic records and dentist recommendations are read-only.'),
    ('Recommendation scope. ', 'The Recommended Visit Window is guidance, not a diagnosis.'),
    ('NgitiBot scope. ', 'NgitiBot explains supported dental topics and existing authorized context; it does not diagnose or override the dentist.'),
    ('Preference synchronization. ', 'Patient notification and personalized-education settings are shared by the web and mobile applications.'),
    ('Privacy. ', 'Use secure devices, do not share credentials, and log out from shared devices.'),
])

# 7
doc.add_heading('7. Troubleshooting', level=1)
add_matrix(['Problem', 'What to Do'], [
    ('Cannot sign in', 'Confirm the registered email and password, check capitalization, then use Forgot Password if needed.'),
    ('Verification code did not arrive', 'Check spam/junk folders, wait for the resend timer, request another code, and confirm the email address.'),
    ('No appointment time is available', 'Choose another date. The selected date may be full or the remaining times may have passed.'),
    ('Cannot book another appointment', 'Review My Appointments. An ongoing appointment or active request may already exist.'),
    ('Cannot cancel or reschedule', 'The appointment may be in clinic, completed, archived, or otherwise not eligible for patient changes.'),
    ('No visit window yet', 'The account may not contain enough supported clinic history or a dentist-suggested next visit.'),
    ('No trends yet', 'At least multiple recent Daily Oral Health Logs are needed before useful trend counts can be shown.'),
    ('Radiograph image unavailable', 'Try again on a stable connection, or contact the clinic for a copy or clarification.'),
    ('Information appears incorrect', 'Contact Dentime Dental Clinic. Clinical records cannot be edited by the patient.'),
    ('Page or screen will not load', 'Check internet access, refresh or pull to refresh, restart the app, then contact the clinic if the issue continues.'),
], [2.2, 4.3])
add_note('Support details to complete', 'Clinic phone: [INSERT]  |  Clinic email: [INSERT]  |  Support hours: [INSERT]  |  Privacy contact: [INSERT]')

# Appendices
doc.add_heading('Appendix A. Screenshot Checklist', level=1)
doc.add_paragraph('Capture the following images with fictional patient data. Use the same account, branch, dates, and naming style throughout the manual.')
check_rows = []
for idx, caption in enumerate([
    'Public website home/navigation', 'Public appointment request and success', 'Pre-registration and consent review',
    'Activation/password setup', 'Web login and recovery', 'Web sidebar', 'Web dashboard',
    'Web appointments and booking', 'Web cancel/reschedule', 'Web daily oral-health log',
    'Web calendar and trends', 'Web Recommended Visit Window', 'Web education library/article',
    'Web NgitiBot/history', 'Web EMR tabs', 'Web notifications', 'Web profile/edit',
    'Web settings', 'Web activity logs/logout', 'Android APK download/installation/login',
    'iOS TestFlight public-link installation/login', 'Mobile onboarding',
    'Mobile navigation/home', 'Mobile visits/booking', 'Mobile Health tabs', 'Mobile records/X-Ray',
    'Mobile NgitiBot', 'Mobile notifications', 'Mobile profile/edit', 'Mobile settings/activity/logout'
], 1):
    check_rows.append((str(idx), caption, '☐'))
add_matrix(['No.', 'Screenshot', 'Done'], check_rows, [0.55, 5.35, 0.6])

doc.add_heading('Appendix B. Quick Reference', level=1)
add_matrix(['Task', 'Web Path', 'Mobile Path'], [
    ('Book appointment', 'My Appointments > Book Appointment', 'Visits > Book Appointment'),
    ('Record daily log', 'Oral Health Management > Today', 'Health > Today'),
    ('Review visit guidance', 'Dashboard or Oral Health Management', 'Home or Health > Predictive Visit'),
    ('Read education', 'Dental Health Education', 'Health > Dental Health Education'),
    ('View records', 'Electronic Medical Record', 'Records'),
    ('Open NgitiBot', 'Floating NgitiBot button', 'Floating NgitiBot button'),
    ('Change password', 'Settings > Account Security', 'Profile > Settings > Change Password'),
    ('Change email', 'Settings > Account Actions > Change Email', 'Profile > Settings > Change Email'),
    ('Review activity', 'Settings > Account Actions > Activity Logs', 'Profile > Activity Logs'),
    ('Log out', 'Sidebar > Logout', 'Profile menu or Settings > Log Out'),
], [1.5, 2.5, 2.5])

doc.add_paragraph()
end = doc.add_paragraph('END OF PATIENT USER MANUAL')
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
end.runs[0].bold = True
end.runs[0].font.color.rgb = RGBColor.from_string(BLUE)

doc.core_properties.title = 'NgitiFy Patient User Manual - Web and Mobile Applications'
doc.core_properties.subject = 'Screenshot-ready patient user manual'
doc.core_properties.author = 'Dentime Dental Clinic'
doc.core_properties.keywords = 'NgitiFy, Dentime, patient, user manual, web, mobile'
doc.save(OUT)
print(OUT.resolve())
