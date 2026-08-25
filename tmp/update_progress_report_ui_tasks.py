from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


SOURCE = Path(
    r"C:\Users\Administrator\Desktop\NGITIFY DENTIME\.codex-docx-inspect-9486bfeb7990437097e4dbd1548a1969\source.docx"
)
OUTPUT = Path(
    r"C:\Users\Administrator\Desktop\NGITIFY DENTIME\output\SOFTWARE PROGRESS REPORT - UI Tasks Included.docx"
)


WEB_UI_TASKS = {
    "Website Online Booking Module": (
        "Present the booking form with branch, service, dentist, date, and time selectors, "
        "inline validation, a loading state, and clear confirmation or error feedback"
    ),
    "Registration Module": (
        "Present the pre-registration and account-activation interfaces with clear required-field, "
        "consent, guardian, password-strength, and success or error feedback"
    ),
    "Login Module": (
        "Provide the login, Forgot Password, OTP, and new-password interfaces with show or hide "
        "password controls, inline errors, loading states, and navigation links"
    ),
    "Account Settings Module": (
        "Provide profile, email, password, and notification-settings forms with editable fields, "
        "preference controls, confirmation dialogs, and save-result feedback"
    ),
    "Staffs Management Module": (
        "Provide role tabs, searchable and filterable staff tables, status badges, pagination, "
        "add or edit forms, record views, and lifecycle confirmation dialogs"
    ),
    "Patients Management Module": (
        "Provide a searchable and filterable patient table, status badges, pagination, add or edit "
        "forms, record views, branch-transfer controls, and lifecycle confirmation dialogs"
    ),
    "Branch Management Module": (
        "Provide branch records with operating-status indicators, add or edit forms, activation "
        "confirmation dialogs, and analytics cards or charts"
    ),
    "Schedule Management Module": (
        "Provide calendar and table schedule views, filters, schedule forms, a detail panel, status "
        "badges, and rescheduling or status-update dialogs"
    ),
    "Patient EMR Module": (
        "Provide organized EMR sections, history and treatment forms, interactive odontogram "
        "controls, a radiograph review panel, image previews, approval actions, and the patient explanation view"
    ),
    "Inventory Management Module": (
        "Provide stock summary cards, searchable and filterable inventory and batch tables, item and "
        "stock forms, low-stock indicators, delete confirmations, and the material-usage interface"
    ),
    "System Configuration Module": (
        "Provide grouped configuration and website-content forms, media previews, validation messages, "
        "and clear save or reset feedback"
    ),
    "Archive Review Module": (
        "Provide a searchable and filterable archive table, record details, restore and delete action "
        "controls, and confirmation dialogs"
    ),
    "Database Backup Module": (
        "Provide backup status cards, a history table, automatic-backup schedule controls, progress "
        "indicators, and download or verification-result feedback"
    ),
    "Integrity Tools Module": (
        "Provide check-summary cards, selectable checks, affected-record detail tables, progress and "
        "status indicators, and Safe Auto-Fix confirmation and result dialogs"
    ),
    "Activity Logs Module": (
        "Provide a searchable and filterable activity-log table with pagination, date and action "
        "filters, a complete detail modal, and loading, empty, and error states"
    ),
    "System Audit Logs Module": (
        "Provide a searchable and filterable audit-log table with pagination, event filters, a "
        "complete detail modal, and loading, empty, and error states"
    ),
    "AI Patient Engagement Module": (
        "Provide a responsive chat panel with conversation history, suggested-prompt chips, a message "
        "composer, send and loading states, and error or retry feedback"
    ),
    "AI Patient Care Companion Module": (
        "Provide patient-care dashboard cards for the visit window, oral-care summary, education "
        "recommendations, source and reason details, and loading or empty states"
    ),
    "Patient Appointments Module": (
        "Provide upcoming and previous appointment views, appointment cards and status badges, a "
        "booking form, cancellation and rescheduling dialogs, and result feedback"
    ),
    "Oral Health Management Module": (
        "Provide a daily-log form, calendar history, 7-day and 30-day trend cards or charts, a visit "
        "recommendation card, and saved, loading, and empty states"
    ),
    "Dental Health Education Module": (
        "Provide contextual topic cards, search and filter controls, an article list or grid, an "
        "article detail view, and loading or empty states"
    ),
    "Patient Notifications Module": (
        "Provide a notification list with unread styling, category and time details, an unread badge "
        "and count, mark-as-read interactions, and loading or empty states"
    ),
    "My EMR Module": (
        "Provide overview and record tabs, history and treatment lists, the interactive odontogram "
        "display, an X-ray gallery and viewer, and loading or empty states"
    ),
}


MOBILE_UI_TASKS = {
    "Login Module": (
        "Provide keyboard-aware mobile login, Forgot Password, OTP, and new-password screens with "
        "show or hide password controls, inline errors, loading states, and clear navigation"
    ),
    "Account Settings Module": (
        "Provide mobile profile and settings screens with section cards, editable fields, notification "
        "switches, confirmation modals, and save-success or error feedback"
    ),
    "AI Patient Engagement Module": (
        "Provide a mobile chat interface with message bubbles, suggested-prompt chips, a keyboard-aware "
        "composer, typing or loading feedback, and error or retry states"
    ),
    "AI Patient Care Companion Module": (
        "Provide mobile-friendly care-summary cards, recommendation banners, source and reason details, "
        "action links, and refresh, loading, or empty states"
    ),
    "Patient Appointments Module": (
        "Provide upcoming and previous tabs, appointment cards and status chips, mobile booking pickers, "
        "cancel or reschedule modals, and success or error feedback"
    ),
    "Oral Health Management Module": (
        "Provide a mobile daily-log form, calendar selector, 7-day and 30-day trend cards or charts, "
        "recommendation cards, and scroll, loading, saved, or empty states"
    ),
    "Dental Health Education Module": (
        "Provide mobile topic cards, a search bar, filter chips, article list and detail screens, and "
        "loading or empty states"
    ),
    "Patient Notifications Module": (
        "Provide a mobile notification list with unread highlights, an unread badge, tap-to-view and "
        "mark-as-read interactions, pull-to-refresh, and loading or empty states"
    ),
    "My EMR Module": (
        "Provide mobile tabbed record views, history and treatment lists, an interactive odontogram, "
        "an X-ray viewer with image controls, and loading or empty states"
    ),
}


def append_like(template_paragraph, task_text: str):
    new_p = deepcopy(template_paragraph._p)
    text_nodes = new_p.findall(".//" + qn("w:t"))
    if not text_nodes:
        raise RuntimeError("The task paragraph template contains no text node")
    text_nodes[0].text = f"☑ {task_text}"
    for node in text_nodes[1:]:
        node.text = ""
    template_paragraph._p.addnext(new_p)


def update_table(table, tasks):
    seen = set()
    for row in table.rows[1:]:
        module_name = row.cells[0].text.strip()
        if module_name not in tasks:
            raise KeyError(f"No UI task defined for {module_name!r}")

        description_cell = row.cells[1]
        existing = [p for p in description_cell.paragraphs if p.text.strip().startswith("☑")]
        if not existing:
            raise RuntimeError(f"No task paragraph found for {module_name}")

        ui_line = f"☑ {tasks[module_name]}"
        if any(p.text.strip() == ui_line for p in description_cell.paragraphs):
            raise RuntimeError(f"UI task already exists for {module_name}")

        append_like(existing[-1], tasks[module_name])
        seen.add(module_name)

    missing = set(tasks) - seen
    if missing:
        raise RuntimeError(f"Modules missing from report table: {sorted(missing)}")


def main():
    doc = Document(SOURCE)
    if len(doc.tables) != 3:
        raise RuntimeError(f"Expected 3 tables, found {len(doc.tables)}")

    update_table(doc.tables[1], WEB_UI_TASKS)
    update_table(doc.tables[2], MOBILE_UI_TASKS)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Saved {OUTPUT}")
    print(f"Added {len(WEB_UI_TASKS)} web UI tasks and {len(MOBILE_UI_TASKS)} mobile UI tasks")


if __name__ == "__main__":
    main()
